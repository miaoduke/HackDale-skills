#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
bilinote.py — Agent skill CLI that reproduces BiliNote's AI video-note pipeline.

Pipeline (subtitle-priority, no RAG):
    fetch       download audio + subtitles for a video URL / local file
    transcribe  audio -> timestamped transcript (bcut / kuaishou; online only)
    summarize   transcript + title/tags -> structured Markdown note
    run         fetch -> (transcribe if no subtitle) -> summarize

Faithful to BiliNote (JefferyHcool/BiliNote, MIT): same subtitle-priority logic,
same note-generation prompt, same time-marker / screenshot / AI-summary options.
The RAG chat/QA feature is intentionally excluded.
"""

import argparse
import json
import os
import sys
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

SKILL_DIR = Path(__file__).resolve().parent.parent
REFERENCES_DIR = SKILL_DIR / "references"
PROMPT_FILE = REFERENCES_DIR / "note_prompt.md"
STYLE_PROMPTS_FILE = REFERENCES_DIR / "style_prompts.md"

# --------------------------------------------------------------------------
# Conditional prompt snippets (mirrors BiliNote/backend/app/gpt/prompt.py)
# --------------------------------------------------------------------------
LINK = (
    "\n9. **Add time markers**: THIS IS IMPORTANT For every main heading (`##`), "
    "append the starting time of that segment using the format ,start with "
    "*Content ,eg: `*Content-[mm:ss]`.\n"
)
AI_SUM = (
    "\n🧠 Final Touch:\n"
    "At the end of the notes, add a professional **AI Summary** in Chinese – "
    "a brief conclusion summarizing the whole video.\n"
)
SCREENSHOT = (
    "\n8. **Screenshot placeholders**: If a section involves **visual demonstrations, "
    "code walkthroughs, UI interactions**, or any content where visuals aid "
    "understanding, insert a screenshot cue at the end of that section:\n"
    "   - Format: `*Screenshot-[mm:ss]`\n"
    "   - Only use it when truly helpful.\n"
)

# --------------------------------------------------------------------------
# Report styles for the summarize step (user-selectable via --style)
# --------------------------------------------------------------------------
# Report styles for the summarize step (user-selectable via --style).
# Mirrors BiliNote's 9 built-in note styles (官方 9 种预设风格) + bilinote
# platform extensions (wechat / bilibili) whose detailed specs live in
# references/style_prompts.md and are injected at prompt-build time.
# https://docs.bilinote.app/guide/basic —— 控制 AI 笔记的「语气 / 表达 / 侧重点」，
# 与「格式」(目录/截图等) 相互独立。默认=精简风格（上游默认值）。
DEFAULT_STYLE = "concise"
STYLES = {
    "concise": ("精简风格",
        "采用【精简风格】：提炼视频核心要点，内容简洁明了，省略冗余铺垫与细枝末节，"
        "保留关键结论与可执行信息，便于快速预览与效率阅读。"),
    "detailed": ("详细风格",
        "采用【详细风格】：尽可能完整记录视频中的细节、例子与论证过程，全面呈现内容，"
        "适合深度理解与学习总结，不要过度删减。"),
    "tutorial": ("教程风格",
        "采用【教程风格】：针对教学内容优化，按清晰步骤 / 操作流程组织，突出「怎么做」，"
        "适合教程类视频，可用步骤编号或分步说明。"),
    "academic": ("学术风格",
        "采用【学术风格】：结构严谨、用词正式，适当引用理论与概念，避免口语化，"
        "适合讲座、报告、科研类内容。"),
    "xiaohongshu": ("小红书风格",
        "采用【小红书风格】：写结构化种草笔记——爆款标题公式 + 痛点→方案→体验→推荐四段式"
        "正文 + 标签策略，语气亲切有感染力、适当用 emoji；详细规范见下方「平台风格详细规范」。"),
    "wechat": ("公众号风格",
        "采用【公众号风格】：写平台级长文——强标题钩子 + 清晰分级 + 手机原生短段落 + "
        "文末引导关注，遵循 60/30/10 内容比例；详细规范见下方「平台风格详细规范」。"),
    "bilibili": ("B站脚本风格",
        "采用【B站脚本风格】：把内容改写成可拍摄的 B站视频脚本文案——黄金30秒价值预告 + "
        "分论点 + 弹幕预埋点 + 自然三连引导；详细规范见下方「平台风格详细规范」。"),
    "lifestyle": ("生活向风格",
        "采用【生活向风格】：日常化语气、带情感色彩，像朋友聊天般自然，"
        "适合 vlog、纪录片类内容。"),
    "task": ("任务导向风格",
        "采用【任务导向风格】：侧重待办与操作清单，把内容转化为清晰的步骤 / 检查项 / 行动项，"
        "适合工作流与项目管理。"),
    "business": ("商业风格",
        "采用【商业风格】：商业逻辑清晰、表达精准，突出数据、决策与价值，"
        "适合财经、市场分析类内容。"),
    "meeting": ("会议纪要风格",
        "采用【会议纪要风格】：结构标准、语言正式，按议题 / 决议 / 待办组织，"
        "适合会议总结与线上讨论。"),
    # ---- 风格库扩展（2026-07-12）：5 个新底风格，补丁见 style_prompts.md ----
    "zhihu": ("知乎问答体",
        "采用【知乎问答体】：先抛结论→论点分层论证→预判反驳→收束，半正式书面语，"
        "区分事实与观点，数据给来源；详细规范见下方「平台风格详细规范」。"),
    "feynman": ("费曼讲解体",
        "采用【费曼讲解体】：用最简单的话讲清复杂事——具体例子开场、短句锚定结论、"
        "类比讲原理并标注失效边界、主动说不知道；详细规范见下方「平台风格详细规范」。"),
    "debate-digest": ("热议盘点体",
        "采用【热议盘点体】：以评论区高能精选为内容骨架，按争议主题分节引用高赞评论，"
        "中立旁观不站队；详细规范见下方「平台风格详细规范」。"),
    "research-report": ("投资研报体",
        "采用【投资研报体】：BLUF 先给判断→行业背景→主体分析→关键数据表→机会与风险→"
        "决策启示，去情绪化、数据给来源；详细规范见下方「平台风格详细规范」。"),
    "cheatsheet": ("速查表",
        "采用【速查表】：表格化、分类索引，每条≤1 行只留「是什么+怎么用」，"
        "删掉所有解释性铺垫，专为高频查阅设计；详细规范见下方「平台风格详细规范」。"),
}
STYLE_HELP = ("报告风格（" + " / ".join(f"{k}({v[0]})" for k, v in STYLES.items())
             + f"，默认 {DEFAULT_STYLE}）")


def load_style_prompts() -> dict:
    """Parse references/style_prompts.md into {style_key: detailed_spec}.

    Sections are delimited by a top-level `## <style_key>` header (exactly two
    hashes). `###` sub-headings inside a section are kept verbatim. The leading
    `# 平台风格提示词补丁` (one hash) is ignored. Returns {} if the file is
    absent so the caller can degrade gracefully."""
    if not STYLE_PROMPTS_FILE.exists():
        return {}
    out, cur, buf = {}, None, []
    for line in STYLE_PROMPTS_FILE.read_text(encoding="utf-8").splitlines():
        m = re.match(r"^##\s+([a-z0-9_-]+)\s*$", line)
        if m:
            if cur is not None:
                out[cur] = "\n".join(buf).strip()
            cur, buf = m.group(1), []
        elif cur is not None:
            buf.append(line)
    if cur is not None:
        out[cur] = "\n".join(buf).strip()
    return out


# Cache once at import time; re-parsed only if the file changes on disk.
_STYLE_PROMPTS_CACHE = load_style_prompts()


# --------------------------------------------------------------------------
# Plan K — 自定义风格库 + 偏好记忆
#   用户可在 references/user_styles/*.md 或 ~/.bilinote_styles/*.md 放自定义风格
#   （同样 ## name 段格式），用 --style @my-brand 引用。wizard 会记住上次选择。
# --------------------------------------------------------------------------
USER_STYLE_DIRS = [
    SKILL_DIR / "references" / "user_styles",
    Path.home() / ".bilinote_styles",
]


def load_user_styles() -> dict:
    """加载用户自定义风格（## name 段格式，与 style_prompts.md 同范式）。
    扫描 USER_STYLE_DIRS 下所有 .md 文件，返回 {name: spec}。"""
    out = {}
    for d in USER_STYLE_DIRS:
        if not d.exists():
            continue
        for f in sorted(d.glob("*.md")):
            try:
                txt = f.read_text(encoding="utf-8")
            except Exception:  # noqa: BLE001
                continue
            cur, buf = None, []
            for line in txt.splitlines():
                m = re.match(r"^##\s+([a-z0-9_-]+)\s*$", line)
                if m:
                    if cur is not None:
                        out[cur] = "\n".join(buf).strip()
                    cur, buf = m.group(1), []
                elif cur is not None:
                    buf.append(line)
            if cur is not None:
                out[cur] = "\n".join(buf).strip()
    return out


_USER_STYLES_CACHE = load_user_styles()


def resolve_style(name: str) -> str:
    """解析 --style 值：@前缀去掉后查自定义风格库，返回真实风格键。
    内置风格原样返回；@unknown 回退 DEFAULT_STYLE。"""
    if not name:
        return DEFAULT_STYLE
    if name.startswith("@"):
        key = name[1:]
        if key in _USER_STYLES_CACHE:
            return key
        log(f"未知自定义风格 @{key}，回退 {DEFAULT_STYLE}")
        return DEFAULT_STYLE
    return name


def get_style_spec(name: str) -> str:
    """取风格的详细补丁：先查内置 style_prompts，再查用户自定义。"""
    key = resolve_style(name)
    return _STYLE_PROMPTS_CACHE.get(key) or _USER_STYLES_CACHE.get(key) or ""


def _prefs_path() -> Path:
    return Path.home() / ".bilinote_prefs.json"


def save_prefs(prefs: dict) -> None:
    """记住用户上次的选择（风格/富化开关/画质/导出）。失败安全。"""
    try:
        import json
        _prefs_path().write_text(
            json.dumps(prefs, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:  # noqa: BLE001
        pass


def load_prefs() -> dict:
    """读取上次偏好；无则返回 {}。"""
    try:
        import json
        p = _prefs_path()
        if p.exists():
            return json.loads(p.read_text(encoding="utf-8"))
    except Exception:  # noqa: BLE001
        pass
    return {}


# --------------------------------------------------------------------------
# Universal input classification + non-video ingestion
# --------------------------------------------------------------------------
# Supported input kinds (decided up front, then routed to the right flow):
#   video_url   - a video platform URL (Bilibili/YouTube/...) -> fetch+transcribe
#   media_url   - a direct media file URL (.mp4/.mp3/...)      -> fetch+transcribe
#   media_file  - a local video/audio file                    -> fetch_local+transcribe
#   article_url - any other http(s) URL                       -> scrape -> summarize
#   doc_file    - a local .txt/.md/.pdf/.docx/.doc/... file   -> read -> summarize
#   text        - raw pasted text                             -> summarize directly
URL_RE = re.compile(r"^https?://", re.IGNORECASE)
VIDEO_HOSTS = ("bilibili.com", "b23.tv", "youtube.com", "youtu.be", "vimeo.com",
               "twitter.com", "x.com", "douyin.com", "tiktok.com", "ixigua.com",
               "weibo.com", "kuaishou.com", "yinyuetai.com")
VIDEO_EXT = {".mp4", ".mkv", ".webm", ".mov", ".avi", ".flv"}
AUDIO_EXT = {".mp3", ".m4a", ".wav", ".flac", ".aac", ".ogg", ".oga"}
DOC_TEXT_EXT = {".txt", ".md", ".markdown", ".text", ".json", ".csv", ".log",
                ".html", ".htm"}
DOC_BINARY_EXT = {".pdf", ".docx", ".doc", ".rtf", ".epub"}
SUBTITLE_EXT = {".vtt", ".srt", ".ass", ".ssa", ".sub"}


def classify_source(raw: str) -> dict:
    """Classify a raw user input into one of the kinds above.

    Returns a dict with at least ``kind`` and ``source``; media kinds also
    carry ``local`` (bool) and ``is_video`` (bool)."""
    s = (raw or "").strip()
    if not s:
        raise ValueError("输入为空，无法识别内容类型。")

    # 1) URL
    if URL_RE.match(s):
        low = s.lower()
        ext = os.path.splitext(low)[1]
        if ext in VIDEO_EXT:
            return {"kind": "media_url", "source": s, "local": False, "is_video": True}
        if ext in AUDIO_EXT:
            return {"kind": "media_url", "source": s, "local": False, "is_video": False}
        host = low
        try:
            from urllib.parse import urlparse
            host = urlparse(s).netloc.lower()
        except Exception:
            pass
        if any(h in host for h in VIDEO_HOSTS):
            return {"kind": "video_url", "source": s, "local": False, "is_video": True}
        return {"kind": "article_url", "source": s}

    # 2) existing local file
    p = Path(s)
    if p.exists() and p.is_file():
        ext = p.suffix.lower()
        if ext in VIDEO_EXT:
            return {"kind": "media_file", "source": s, "local": True, "is_video": True}
        if ext in AUDIO_EXT:
            return {"kind": "media_file", "source": s, "local": True, "is_video": False}
        if ext in SUBTITLE_EXT:
            return {"kind": "subtitle_file", "source": s, "local": True, "is_video": False}
        if ext in DOC_TEXT_EXT or ext in DOC_BINARY_EXT:
            return {"kind": "doc_file", "source": s}
        # existing local file with a non-empty unknown extension -> reject
        # (pasted text / URLs handled in other branches; extensionless files
        #  kept as text, e.g. a README with no suffix)
        if ext:
            raise ValueError(
                f"不支持的文件类型：{ext}（本地文件扩展名不在已知集合内，"
                f"无法识别为笔记原料）。支持：视频/音频/字幕/文档文本类。")
        return {"kind": "doc_file", "source": s}

    # 3) raw pasted text
    return {"kind": "text", "source": s}


def _cls_hit(cls, needles) -> bool:
    """True if any needle appears in the element's class list/string."""
    if not cls:
        return False
    s = " ".join(cls) if isinstance(cls, (list, tuple)) else str(cls)
    s = s.lower()
    return any(k in s for k in needles)


def scrape_article(url: str) -> tuple[str, str, list]:
    """Fetch a web page and extract its main article text + main images.

    Returns (title, clean_text, images) where `images` is a list of absolute
    content-image URLs found in the article body (best-effort; may be empty).
    Uses trafilatura when available, with a BeautifulSoup fallback (fed raw
    *bytes* so the real charset is auto-detected from the <meta> tag — avoids
    mojibake when the server omits a charset), then a last-resort regex tag
    strip. Raises if no usable text is found."""
    import requests
    from urllib.parse import urljoin
    headers = {"User-Agent": "Mozilla/5.0 (compatible; BiliNoteSkill/1.0)"}
    resp = requests.get(url, headers=headers, timeout=25)
    resp.raise_for_status()
    raw = resp.content  # bytes: let the parser detect UTF-8 from <meta>
    title, text = "", ""

    # 1) trafilatura — best at isolating the main article body
    try:
        import trafilatura
        extracted = trafilatura.extract(raw, include_comments=False)
        if extracted:
            text = extracted
        try:
            meta = trafilatura.extract_metadata(raw)
            if meta and getattr(meta, "title", None):
                title = meta.title
        except Exception:
            pass
    except Exception:
        pass

    # 2) BeautifulSoup fallback (operates on bytes -> correct decoding)
    if not text:
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(raw, "html.parser")
            for tag in soup(["script", "style", "noscript", "header",
                            "footer", "nav", "aside", "form", "iframe"]):
                tag.decompose()
            CONTENT_KEYS = ("article", "artibody", "content", "detail",
                            "post", "post_body", "news_text", "main", "text")
            NOISE = ("comment", "related", "recommend", "share", "sidebar",
                     "footer", "banner", "hot", "rank", "ad-", "tool",
                     "aside", "crumb", "bread")
            main = (soup.find("article")
                    or soup.find("main")
                    or soup.find(class_=lambda c: _cls_hit(c, CONTENT_KEYS))
                    or soup.body)
            if main is None:
                main = soup
            for junk in main.find_all(
                    class_=lambda c: _cls_hit(c, NOISE)):
                junk.decompose()
            if not title:
                t = soup.find("title")
                if t:
                    title = t.get_text(strip=True)
            if not title:
                h1 = main.find("h1")
                if h1:
                    title = h1.get_text(strip=True)
            text = main.get_text(separator="\n")
        except Exception:
            import re as _re
            text = _re.sub(r"<[^>]+>", "", raw.decode("utf-8", "ignore"))

    # 3) normalize: trim lines, drop blanks
    text = "\n".join(ln.strip() for ln in text.splitlines() if ln.strip())
    if not text:
        raise RuntimeError("无法从网页提取正文，可能被反爬或需要登录。")

    # 4) best-effort: collect content image URLs from the article
    images = _extract_article_images(raw, url)
    return (title or url, text, images)


_IMG_NOISE = ("icon", "logo", "avatar", "banner", "ad", "pixel", "spacer",
              "qr", "qrcode", "weixin", "wechat", "tracking", "analytics",
              "count", "1x1", "blank", "s.gif", "loading", "placeholder",
              "emoji", "face", "gif", "login", "thumb", "default", "app",
              "sinafinance", "weibo", "fysqfnf", "wx", "qianniu", "tb")


def _extract_article_images(raw: bytes, page_url: str) -> list:
    """Return up to ~8 absolute content-image URLs found in article HTML
    (lazy-load `data-src`/`data-original` supported; icons/ads filtered)."""
    out = []
    try:
        from bs4 import BeautifulSoup
        from urllib.parse import urljoin
        soup = BeautifulSoup(raw, "html.parser")
        for img in soup.find_all("img"):
            src = (img.get("data-src") or img.get("data-original")
                   or img.get("data-lazy-src") or img.get("src") or "").strip()
            if not src:
                continue
            if any(c in src for c in ("@@", "{", "}")):
                continue  # skip template placeholders like @@pic1@@
            low = src.lower()
            if not low.startswith("http"):
                try:
                    src = urljoin(page_url, src)
                    low = src.lower()
                except Exception:
                    continue
            if any(k in low for k in _IMG_NOISE):
                continue
            if src not in out:
                out.append(src)
            if len(out) >= 8:
                break
    except Exception:
        pass
    return out


def read_document(path: str) -> tuple[str, str]:
    """Read a local document into (title, text).

    Supports .txt/.md/.json/.csv/.html and .pdf/.docx; .doc is attempted via
    textract when available."""
    p = Path(path)
    ext = p.suffix.lower()
    title = p.stem

    if ext in (".txt", ".md", ".markdown", ".text", ".json", ".csv", ".log"):
        return (title, p.read_text(encoding="utf-8", errors="ignore"))

    if ext in (".html", ".htm"):
        html = p.read_text(encoding="utf-8", errors="ignore")
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()
            text = soup.get_text(separator="\n")
        except Exception:
            import re as _re
            text = _re.sub(r"<[^>]+>", "", html)
        text = "\n".join(l.strip() for l in text.splitlines() if l.strip())
        return (title, text)

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            from PyPDF2 import PdfReader
        reader = PdfReader(str(p))
        parts = []
        doc_title = ""
        try:
            if reader.metadata and getattr(reader.metadata, "title", None):
                doc_title = reader.metadata.title
        except Exception:
            pass
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return (doc_title or title, "\n".join(parts))

    if ext == ".docx":
        from docx import Document
        doc = Document(str(p))
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        doc_title = ""
        try:
            doc_title = doc.core_properties.title or ""
        except Exception:
            pass
        return (doc_title or title, "\n".join(parts))

    if ext == ".doc":
        try:
            import textract
            text = textract.process(str(p)).decode("utf-8", "ignore")
            return (title, text)
        except Exception as e:
            raise RuntimeError(
                "暂不支持旧版 .doc（需 textract/antiword）。请改用 .docx，"
                "或先另存为 .txt / .pdf。") from e

    raise RuntimeError(f"不支持的文档类型：{ext}")


def log(msg: str) -> None:
    print(f"[bilinote] {msg}", file=sys.stderr, flush=True)


# Plan F: F1–F10 失败恢复运行时标识（与 SKILL.md「失败恢复」表一一对应）。
# 触发任一恢复分支时调用 log_recovery("F3", "..."), 便于运行日志检索与自动化追踪。
FAILURE_CODES = {
    "F1": "fetch 完全失败",
    "F2": "字幕全空/乱码 → 转写",
    "F3": "主引擎语种错配 → 切引擎",
    "F4": "转写返回空 → 换引擎",
    "F5": "embed-screenshots 无 ffmpeg",
    "F6": "单图下载失败 → 跳过",
    "F7": "URL 追踪/重定向 404",
    "F8": "classify_source=unknown → 当纯文本",
    "F9": "native 未生成 brief → 重跑",
    "F10": "原生模式无 API key 依赖 → 所有生成走简报交 agent 当前 LLM",
}


def log_recovery(code: str, detail: str = "") -> None:
    """记录一次失败恢复事件，带 F 编号标识，便于运行时追踪。"""
    desc = FAILURE_CODES.get(code, "")
    tail = f"：{detail}" if detail else ""
    print(f"[bilinote][{code}] {desc}{tail}", file=sys.stderr, flush=True)


def load_prompt_template() -> str:
    if not PROMPT_FILE.exists():
        raise FileNotFoundError(f"Missing prompt template: {PROMPT_FILE}")
    return PROMPT_FILE.read_text(encoding="utf-8")


def fmt_time(seconds: float) -> str:
    seconds = int(round(seconds))
    m, s = divmod(seconds, 60)
    h, m = divmod(m, 60)
    if h:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


# --------------------------------------------------------------------------
# fetch
# --------------------------------------------------------------------------
def _require(tool: str) -> str:
    path = shutil.which(tool)
    if not path:
        raise RuntimeError(
            f"Required tool '{tool}' not found on PATH. "
            f"Install it (e.g. ffmpeg / yt-dlp) and retry."
        )
    return path


def _ytdlp_cmd() -> list[str]:
    """Resolve the yt-dlp command. Prefer the in-environment module
    (`python -m yt_dlp`) so it works without polluting PATH; fall back to a
    PATH lookup."""
    try:
        import importlib.util
        if importlib.util.find_spec("yt_dlp"):
            return [sys.executable, "-m", "yt_dlp"]
    except Exception:
        pass
    p = shutil.which("yt-dlp")
    if not p:
        raise RuntimeError(
            "Required tool 'yt-dlp' not found. Install it: "
            "pip install yt-dlp")
    return [p]


def _ytdlp_extra_opts() -> list[str]:
    """Extra yt-dlp options sourced from the environment, mirroring upstream
    BiliNote's browser-cookie sync + global proxy handling:
      - BILINOTE_YTDLP_COOKIES : browser name to read cookies from
        (e.g. "chrome", "edge", "firefox"); enables login-state CDN mirrors
        that are often reachable when anonymous ones are blocked.
      - BILINOTE_PROXY / HTTPS_PROXY / HTTP_PROXY : outgoing proxy for
        downloads (helps when the CDN is geo/network restricted).
    Returns [] when nothing is configured, so default behaviour is unchanged.
    """
    opts: list[str] = []
    cookies = (os.environ.get("BILINOTE_YTDLP_COOKIES") or "").strip()
    if cookies:
        opts += ["--cookies-from-browser", cookies]
    proxy = (os.environ.get("BILINOTE_PROXY")
             or os.environ.get("HTTPS_PROXY")
             or os.environ.get("HTTP_PROXY") or "").strip()
    if proxy:
        opts += ["--proxy", proxy]
    return opts


# ==========================================================================
# Plan G — 缓存复用层（content-addressed cache）
#   同源换风格/换派生零重抓：以 source+lang+transcriber 等哈希为键缓存
#   transcript/title/tags/danmaku/comments/ai_summary。命中即返回，--no-cache
#   强制刷新。缓存目录 ~/.bilinote_cache/（可 BILINOTE_CACHE_DIR 配置）。
#   全部失败安全：缓存读写异常不影响主链路。
# ==========================================================================
CACHE_ENABLED = True  # 模块级开关，run/parser 的 --no-cache 会置 False


def _cache_dir() -> Path:
    d = os.environ.get("BILINOTE_CACHE_DIR", "~/.bilinote_cache")
    p = Path(d).expanduser()
    try:
        p.mkdir(parents=True, exist_ok=True)
    except Exception:  # noqa: BLE001
        pass
    return p


def _cache_key(*parts) -> str:
    """把若干字符串 parts 拼成稳定的 sha1 键。"""
    import hashlib
    raw = "|".join(str(x) for x in parts if x is not None)
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()[:16]


def _cache_get(category: str, key: str):
    """读缓存；返回 (True, value) 命中或 (False, None) 未命中/出错。"""
    if not CACHE_ENABLED:
        return False, None
    try:
        f = _cache_dir() / category / f"{key}.json"
        if f.exists():
            import json
            return True, json.loads(f.read_text(encoding="utf-8"))
    except Exception:  # noqa: BLE001
        pass
    return False, None


def _cache_set(category: str, key: str, value) -> None:
    """写缓存；失败安全。"""
    if not CACHE_ENABLED:
        return
    try:
        import json
        d = _cache_dir() / category
        d.mkdir(parents=True, exist_ok=True)
        (d / f"{key}.json").write_text(
            json.dumps(value, ensure_ascii=False), encoding="utf-8")
    except Exception:  # noqa: BLE001
        pass


def cache_list() -> list:
    """列出缓存目录下各 category 的条目数与总大小。"""
    out = []
    try:
        root = _cache_dir()
        for cat in sorted(root.iterdir()):
            if cat.is_dir():
                files = list(cat.glob("*.json"))
                if not files:
                    continue  # 跳过空目录
                size = sum(f.stat().st_size for f in files)
                out.append({"category": cat.name, "count": len(files),
                            "size_kb": round(size / 1024, 1)})
    except Exception:  # noqa: BLE001
        pass
    return out


def cache_clear(category: str = "") -> int:
    """清空缓存；category 为空清全部，否则只清该类。返回删除条目数。"""
    n = 0
    try:
        root = _cache_dir()
        cats = [root / category] if category else [
            d for d in root.iterdir() if d.is_dir()]
        for cat in cats:
            if cat.is_dir():
                for f in cat.glob("*.json"):
                    f.unlink(missing_ok=True)
                    n += 1
    except Exception:  # noqa: BLE001
        pass
    return n


# ==========================================================================
# Plan L — 产出度量与可观测性
#   run/review/produce/batch 产出时写 ~/.bilinote_metrics.jsonl（每条一行 JSON）。
#   metrics 子命令聚合输出仪表盘 Markdown（事件分布/风格分布/质量评分趋势）。
# ==========================================================================
def _metrics_path() -> Path:
    return Path(os.environ.get(
        "BILINOTE_METRICS", "~/.bilinote_metrics.jsonl")).expanduser()


def record_metric(event: str, **data) -> None:
    """追加一条度量记录（失败安全）。event ∈ {run, review, produce, batch}。"""
    try:
        import json, time
        rec = {"ts": int(time.time()), "event": event}
        rec.update(data)
        with open(_metrics_path(), "a", encoding="utf-8") as f:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    except Exception:  # noqa: BLE001
        pass


def metrics_summary() -> str:
    """聚合度量记录，返回仪表盘 Markdown。"""
    import json, collections
    p = _metrics_path()
    if not p.exists():
        return f"暂无度量记录（{p}）。跑 run/review/produce/batch 后自动采集。"
    events, styles, scores = collections.Counter(), collections.Counter(), []
    total = 0
    try:
        for line in p.read_text(encoding="utf-8").splitlines():
            if not line.strip():
                continue
            try:
                r = json.loads(line)
            except Exception:  # noqa: BLE001
                continue
            total += 1
            events[r.get("event", "?")] += 1
            if r.get("style"):
                styles[r["style"]] += 1
            if r.get("score") is not None:
                scores.append(r["score"])
    except Exception:  # noqa: BLE001
        pass
    lines = [f"# 产出度量仪表盘\n\n共 {total} 条记录（来源：{p}）\n\n",
             "## 事件分布\n"]
    for ev, n in events.most_common():
        lines.append(f"- {ev}: {n} 次")
    if styles:
        lines.append("\n## 风格产出分布\n")
        for s, n in styles.most_common():
            lines.append(f"- {s}: {n} 次")
    if scores:
        avg = round(sum(scores) / len(scores), 2)
        lines.append(f"\n## 质量评分趋势\n- 共 {len(scores)} 次评分，均值 {avg}/5"
                     f"，最高 {max(scores)}，最低 {min(scores)}")
    return "\n".join(lines)


# ==========================================================================
# Video-enrichment capabilities (borrowed ideas from Bilibili clients like
# PiliPlus / bilibili-API-collect). All are OPTIONAL enrichers: each returns
# a safe empty/failed result so the core subtitle-priority pipeline still
# works unchanged when they are unavailable.
#   #0 fetch_bilibili_ai_summary  — official server-side AI summary
#                                   (summary + timestamped outline + AI subtitle)
#   #1 fetch_bilibili_danmaku     — danmaku as a "second subtitle" context
#   #2 suggest_screenshot_timestamps — high-energy / chapter-driven frame picks
#   #3 linkify_timestamps         — clickable [MM:SS] -> video?t=seconds
#   #4 detect_sponsor_segments_from_danmaku + clean_transcript
#                                   — 弹幕驱动的赞助/广告段检测 + 物理净化
#                                     (替代外网 SponsorBlock，原生适配 B 站)
# ==========================================================================

_BILI_BV_RE = re.compile(r"(BV[0-9A-Za-z]{10})")
_BILI_AV_RE = re.compile(r"av(\d+)", re.I)

# WBI signature mixin-key reorder table (bilibili-API-collect, misc/sign/wbi)
_WBI_MIXIN_TAB = [
    46, 47, 18, 2, 53, 8, 23, 32, 15, 50, 10, 31, 58, 3, 45, 35, 27, 43, 5, 49,
    33, 9, 42, 19, 29, 28, 14, 39, 12, 38, 41, 13, 37, 48, 7, 16, 24, 55, 40,
    61, 26, 17, 0, 1, 60, 51, 30, 4, 22, 25, 54, 21, 56, 59, 6, 63, 57, 62, 11,
    36, 20, 34, 44, 52,
]


def _is_bilibili(url: str) -> bool:
    return "bilibili.com" in (url or "") or bool(_BILI_BV_RE.search(url or ""))


def _bili_extract_id(url: str) -> dict:
    m = _BILI_BV_RE.search(url or "")
    if m:
        return {"bvid": m.group(1)}
    m = _BILI_AV_RE.search(url or "")
    if m:
        return {"aid": int(m.group(1))}
    return {}


def _bili_headers(cookie: str = "") -> dict:
    h = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                      "AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/120.0.0.0 Safari/537.36",
        "Referer": "https://www.bilibili.com/",
    }
    if cookie:
        h["Cookie"] = cookie
    return h


def _wbi_get_mixin_key(img_key: str, sub_key: str) -> str:
    raw = img_key + sub_key
    return "".join(raw[i] for i in _WBI_MIXIN_TAB if i < len(raw))[:32]


def _wbi_sign(params: dict, img_key: str, sub_key: str) -> dict:
    """Sign params with WBI (adds wts + w_rid). Pure function — unit-testable."""
    import time as _t
    import hashlib
    from urllib.parse import urlencode
    mixin = _wbi_get_mixin_key(img_key, sub_key)
    params = dict(params)
    params["wts"] = int(_t.time())
    filtered = {k: "".join(ch for ch in str(v) if ch not in "!'()*")
                for k, v in sorted(params.items())}
    query = urlencode(filtered)
    params["w_rid"] = hashlib.md5((query + mixin).encode()).hexdigest()
    return params


def _bili_get_wbi_keys(cookie: str = "") -> tuple:
    import requests
    r = requests.get("https://api.bilibili.com/x/web-interface/nav",
                     headers=_bili_headers(cookie), timeout=15)
    wbi = (r.json().get("data", {}) or {}).get("wbi_img", {}) or {}
    img = wbi.get("img_url", "") or ""
    sub = wbi.get("sub_url", "") or ""
    img_key = img.rsplit("/", 1)[-1].split(".")[0] if img else ""
    sub_key = sub.rsplit("/", 1)[-1].split(".")[0] if sub else ""
    return img_key, sub_key


def _bili_get_view(ids: dict, cookie: str = "") -> dict:
    """Resolve cid/up_mid/aid/bvid/duration/title via the view API."""
    import requests
    params = {}
    if "bvid" in ids:
        params["bvid"] = ids["bvid"]
    elif "aid" in ids:
        params["aid"] = ids["aid"]
    r = requests.get("https://api.bilibili.com/x/web-interface/view",
                     params=params, headers=_bili_headers(cookie), timeout=15)
    d = r.json().get("data", {}) or {}
    return {
        "cid": d.get("cid"),
        "up_mid": (d.get("owner", {}) or {}).get("mid"),
        "aid": d.get("aid"),
        "bvid": d.get("bvid"),
        "duration": d.get("duration"),
        "title": d.get("title"),
    }


def fetch_bilibili_ai_summary(url: str, cookie: str = "") -> dict:
    """#0 Call Bilibili's official AI summary (conclusion/get).

    Returns {"ok": True, "summary", "outline":[{title,timestamp,points}],
             "subtitle":[{start,end,content}], "cid","bvid","aid","duration"}
    or {"ok": False, "reason": ...} on any failure so callers can fall back to
    the normal subtitle-priority -> transcribe pipeline.
    Requires SESSDATA cookie for full results; without login the API usually
    returns code=-101 (not logged in).
    """
    import requests
    ids = _bili_extract_id(url)
    if not ids:
        return {"ok": False, "reason": "not a bilibili video id"}
    # Plan G：缓存命中（键 = 视频id + cookie 哈希）
    ck = _cache_key("ai", ids.get("bvid") or ids.get("aid"), cookie)
    hit, cached = _cache_get("ai_summary", ck)
    if hit and isinstance(cached, dict):
        log("[cache] ai_summary 命中")
        return cached
    try:
        view = _bili_get_view(ids, cookie)
        cid = view.get("cid")
        if not cid:
            return {"ok": False, "reason": "could not resolve cid"}
        img_key, sub_key = _bili_get_wbi_keys(cookie)
        if not img_key or not sub_key:
            return {"ok": False, "reason": "could not get wbi keys"}
        params = {"cid": cid}
        if view.get("bvid"):
            params["bvid"] = view["bvid"]
        elif view.get("aid"):
            params["aid"] = view["aid"]
        if view.get("up_mid"):
            params["up_mid"] = view["up_mid"]
        signed = _wbi_sign(params, img_key, sub_key)
        r = requests.get(
            "https://api.bilibili.com/x/web-interface/view/conclusion/get",
            params=signed, headers=_bili_headers(cookie), timeout=20)
        j = r.json()
        if j.get("code") != 0:
            return {"ok": False,
                    "reason": f"api code={j.get('code')} {j.get('message')}"}
        data = j.get("data", {}) or {}
        inner = data.get("code")
        if inner == -1:
            return {"ok": False, "reason": "AI summary unsupported (sensitive)"}
        if inner == 1:
            return {"ok": False, "reason": "no summary (voice not recognized)"}
        mr = data.get("model_result", {}) or {}
        outline = [{
            "title": seg.get("title", ""),
            "timestamp": seg.get("timestamp", 0),
            "points": [{"t": p.get("timestamp", 0), "content": p.get("content", "")}
                       for p in (seg.get("part_outline", []) or [])],
        } for seg in (mr.get("outline", []) or [])]
        subs = []
        for sblk in (mr.get("subtitle", []) or []):
            for ps in (sblk.get("part_subtitle", []) or []):
                subs.append({"start": ps.get("start_timestamp", 0),
                             "end": ps.get("end_timestamp", 0),
                             "content": ps.get("content", "")})
        result = {"ok": True, "summary": mr.get("summary", ""),
                  "outline": outline, "subtitle": subs,
                  "result_type": mr.get("result_type", 0),
                  "cid": cid, "bvid": view.get("bvid"), "aid": view.get("aid"),
                  "duration": view.get("duration"), "title": view.get("title")}
        _cache_set("ai_summary", ck, result)  # Plan G：成功才缓存
        return result
    except Exception as e:
        return {"ok": False, "reason": f"exception: {e}"}


def bili_ai_summary_to_transcript(ai: dict) -> str:
    """Build a [MM:SS]-prefixed transcript from the official AI subtitle list."""
    return "\n".join(f"[{fmt_time(s.get('start', 0))}] {s.get('content', '')}"
                     for s in (ai.get("subtitle", []) or []))


def bili_ai_summary_to_markdown(ai: dict, with_screenshots: bool = False) -> str:
    """Render the official AI summary + timestamped outline as a MD scaffold."""
    out = []
    if ai.get("summary"):
        out.append("## AI 总结（B 站官方）\n")
        out.append(ai["summary"] + "\n")
    for seg in (ai.get("outline", []) or []):
        ts = fmt_time(seg.get("timestamp", 0))
        out.append(f"\n## {seg.get('title', '')} *Content-[{ts}]*")
        for p in seg.get("points", []):
            out.append(f"- [{fmt_time(p.get('t', 0))}] {p.get('content', '')}")
        if with_screenshots:
            out.append(f"\n*Screenshot-[{ts}]*")
    return "\n".join(out) + "\n"


def _pb_read_varint(buf: bytes, i: int):
    """Read a protobuf varint from buf at offset i -> (value, new_offset)."""
    shift = 0
    result = 0
    while i < len(buf):
        b = buf[i]
        result |= (b & 0x7F) << shift
        i += 1
        if not (b & 0x80):
            break
        shift += 7
    return result, i


def _pb_iter_fields(buf: bytes):
    """Yield (field_number, wire_type, value) for a protobuf message.
    value is int for varint (wt 0) or bytes for length-delimited (wt 2)."""
    i = 0
    while i < len(buf):
        key, i = _pb_read_varint(buf, i)
        field, wt = key >> 3, key & 0x07
        if wt == 0:
            val, i = _pb_read_varint(buf, i)
            yield field, wt, val
        elif wt == 2:
            ln, i = _pb_read_varint(buf, i)
            yield field, wt, buf[i:i + ln]
            i += ln
        elif wt == 5:
            yield field, wt, buf[i:i + 4]
            i += 4
        elif wt == 1:
            yield field, wt, buf[i:i + 8]
            i += 8
        else:  # unknown wire type -> stop to avoid runaway
            break


def _parse_dm_pb_segment(data: bytes) -> list:
    """Parse a DmSegMobileReply protobuf: repeated DanmakuElem elem = 1.
    Each DanmakuElem: progress(ms)=field 2 (varint), content=field 7 (string).
    Dependency-free wire-format parser. Returns [{t, text}]."""
    items = []
    for field, wt, val in _pb_iter_fields(data):
        if field == 1 and wt == 2:  # a DanmakuElem
            progress_ms = 0
            content = ""
            for f2, w2, v2 in _pb_iter_fields(val):
                if f2 == 2 and w2 == 0:
                    progress_ms = v2
                elif f2 == 7 and w2 == 2:
                    try:
                        content = v2.decode("utf-8", "ignore")
                    except Exception:
                        content = ""
            if content:
                items.append({"t": progress_ms / 1000.0, "text": content})
    return items


def fetch_bilibili_danmaku(cid, cookie: str = "", limit: int = 0,
                           duration: float = None) -> list:
    """#1 Fetch danmaku as a "second subtitle". Returns [{t,text}] sorted by t.

    Strategy: prefer the modern protobuf segmented endpoint
    (x/v2/dm/web/seg.so, 6-min segments) which covers current videos; fall
    back to the legacy XML endpoint (comment.bilibili.com/{cid}.xml) for old
    videos. Both are best-effort and return [] on any failure so callers can
    proceed without danmaku.
    """
    import requests
    if not cid:
        return []
    # Plan G：缓存命中（键 = cid + limit + duration）
    ck = _cache_key("dm", cid, limit, int(duration or 0))
    hit, cached = _cache_get("danmaku", ck)
    if hit and isinstance(cached, list):
        log("[cache] danmaku 命中")
        return cached
    items = []
    # --- modern protobuf segmented endpoint ---
    try:
        n_seg = int((duration or 0) // 360) + 1 if duration else 20  # cap 2h
        for seg in range(1, n_seg + 1):
            r = requests.get(
                "https://api.bilibili.com/x/v2/dm/web/seg.so",
                params={"type": 1, "oid": cid, "segment_index": seg},
                headers=_bili_headers(cookie), timeout=20)
            if r.status_code != 200 or not r.content:
                break
            ct = r.headers.get("content-type", "")
            if "json" in ct:  # error payload (e.g. code -352 / -101)
                break
            seg_items = _parse_dm_pb_segment(r.content)
            if not seg_items:
                break
            items.extend(seg_items)
    except Exception:
        items = []
    # --- legacy XML fallback ---
    if not items:
        try:
            import xml.etree.ElementTree as ET
            r = requests.get(f"https://comment.bilibili.com/{cid}.xml",
                             headers=_bili_headers(cookie), timeout=20)
            r.encoding = "utf-8"
            root = ET.fromstring(r.text)
            for d in root.findall("d"):
                p = (d.get("p") or "").split(",")
                if not p or not d.text:
                    continue
                try:
                    items.append({"t": float(p[0]), "text": d.text})
                except ValueError:
                    continue
        except Exception:
            pass
    items.sort(key=lambda x: x["t"])
    out = items[:limit] if limit else items
    _cache_set("danmaku", ck, out)  # Plan G：成功才缓存
    return out


def danmaku_as_context(danmaku: list, max_chars: int = 4000) -> str:
    """Condense danmaku into a compact [MM:SS] context block for the LLM."""
    if not danmaku:
        return ""
    text = "\n".join(f"[{fmt_time(d['t'])}] {d['text']}" for d in danmaku)
    if len(text) > max_chars:
        text = text[:max_chars] + "\n…(弹幕过多已截断)"
    return text


# --------------------------------------------------------------------------
# #C 评论区高能精选（comment-area highlights）
#    抓取 B 站评论区高赞评论，作为「观众态度」参考或可选加入报告正文。
#    与弹幕同源：都用 WBI 签名；oid = aid。评论本身不带视频时间戳，
#    故「高能」以「高赞」为核心信号（mode=3 按热度 + 本地再按赞排序）。
# --------------------------------------------------------------------------
_NOISE_CM = ["前排", "沙发", "板凳", "地板", "围观", "mark", "MARK",
             "留名", "路过", "打卡", "催更", "占楼", "镇楼", "前排占座"]


def _is_noise_comment(txt: str) -> bool:
    """Drop 纯表情/符号、过短、以及「前排/沙发」等无信息噪声评论。"""
    if not txt:
        return True
    s = txt.strip()
    if len(s) < 5:
        return True
    # 纯符号/表情（去掉常见标点后没有字母数字）
    stripped = "".join(ch for ch in s if ch not in "，。、！？…~·…—“”‘’()（）[]【】 ")
    if stripped and not any(ch.isalnum() for ch in stripped) and len(s) <= 14:
        return True
    if any(k in s for k in _NOISE_CM):
        return True
    return False


def fetch_bilibili_comments(oid, cookie: str = "", limit: int = 60,
                            sort: str = "hot") -> list:
    """#C 抓取 B 站评论区（WBI 签名的 reply/wbi/main）。

    返回 [{'rpid','text','likes','user','ctime','sub':[...]}]，按赞数降序。
    oid = aid（视频 av 号）；type=1。最多抓 3 页再做本地按赞排序。
    Best-effort：任何失败都返回 []，调用方可在无评论时继续。
    """
    import requests
    import json as _json
    if not oid:
        return []
    # Plan G：缓存命中（键 = oid + limit + sort）
    ck = _cache_key("cm", oid, limit, sort)
    hit, cached = _cache_get("comments", ck)
    if hit and isinstance(cached, list):
        log("[cache] comments 命中")
        return cached
    mode = 3 if sort == "hot" else 2  # 3=按热度 2=按时间；本地再按赞排序兜底
    collected = []
    pagination_str = "{}"
    headers = _bili_headers(cookie)
    try:
        img_key, sub_key = _bili_get_wbi_keys(cookie)
        if not img_key or not sub_key:
            return []
        for _ in range(3):
            params = {"oid": oid, "type": 1, "mode": mode,
                      "web_location": "333.1007", "pagination_str": pagination_str}
            signed = _wbi_sign(params, img_key, sub_key)
            r = requests.get("https://api.bilibili.com/x/v2/reply/wbi/main",
                             params=signed, headers=headers, timeout=20)
            j = r.json()
            if j.get("code") != 0:
                break
            data = j.get("data", {}) or {}
            replies = data.get("replies") or []
            if not replies:
                break
            for c in replies:
                member = (c.get("member", {}) or {})
                content = (c.get("content", {}) or {})
                txt = (content.get("message") or "").strip()
                if not txt:
                    continue
                collected.append({
                    "rpid": c.get("rpid"),
                    "text": txt,
                    "likes": int(c.get("like", 0) or 0),
                    "user": member.get("uname", "") or "匿名",
                    "ctime": c.get("ctime", 0),
                    "sub": [(s.get("content", {}) or {}).get("message", "")
                            for s in (c.get("replies") or [])
                            if (s.get("content", {}) or {}).get("message")][:3],
                })
            cur = data.get("cursor", {}) or {}
            if cur.get("is_end"):
                break
            pagination_str = cur.get("pagination_str") or "{}"
            if pagination_str == "{}":
                break
        collected.sort(key=lambda x: x["likes"], reverse=True)
        out = collected[:limit] if limit else collected
        _cache_set("comments", ck, out)  # Plan G：成功才缓存
        return out
    except Exception:
        return []


def select_highlight_comments(comments: list, top_n: int = 8,
                              min_likes: int = 10) -> list:
    """『高能精选』：滤掉噪声/表情/前排，按赞数取前 top_n（赞数需 >= min_likes）。"""
    cand = [c for c in (comments or []) if not _is_noise_comment(c.get("text", ""))]
    cand.sort(key=lambda x: x.get("likes", 0), reverse=True)
    out = []
    for c in cand:
        if len(out) >= top_n:
            break
        if c.get("likes", 0) < min_likes:
            continue
        out.append(c)
    return out


def comments_as_context(highlights: list, max_chars: int = 2000) -> str:
    """安静上下文块：提示 LLM 观众态度/争议点，不强制进正文（类似弹幕上下文）。"""
    if not highlights:
        return ""
    lines = [f"[@{c['user']} · ❤{c['likes']}] {c['text']}" for c in highlights]
    text = "\n".join(lines)
    if len(text) > max_chars:
        text = text[:max_chars] + "\n…(评论过多已截断)"
    return text


def comments_as_report_section(highlights: list) -> str:
    """报告正文段：## 评论区高能精选（逐条高赞评论 + @用户 + 赞数）。"""
    if not highlights:
        return ""
    lines = ["## 评论区高能精选\n",
             "> 以下为评论区高赞（高能）观点，由 B 站评论区抓取并按赞排序，"
             "仅供辅助理解观众态度，**不代表笔记立场**：\n"]
    for i, c in enumerate(highlights, 1):
        lines.append(f"{i}. {c['text']}")
        lines.append(f"   — @{c['user']} · ❤ {c['likes']} 赞")
        for s in (c.get("sub") or []):
            if s and s.strip():
                lines.append(f"   - ↳ {s.strip()}")
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


# --------------------------------------------------------------------------
# #4 content cleaning: drop sponsor/ad + intro/outro (Bilibili-native,
#     replaces the external SponsorBlock API which has no Bilibili data)
# --------------------------------------------------------------------------
# #4 内容净化 —— 亦称「弹幕恰饭检测」（danmaku sponsor / 恰饭 detection）。
# 弹幕里观众对"恰饭/广告"最敏感，这些词是 B 站更可靠的赞助信号：
_SPONSOR_KW = ["恰饭", "赞助", "广告", "推广", "种草", "商务合作", "带货",
               "优惠券", "领券", "下单", "链接在简介", "小黄车", "橱窗",
               "购买链接", "好物", "福利", "限时优惠", "活动价", "品宣"]


def detect_sponsor_segments_from_danmaku(danmaku: list, keywords=None,
                                         window: int = 25, min_hits: int = 2) -> list:
    """#4 内容净化 / 弹幕恰饭检测：从弹幕关键词聚簇出疑似赞助/广告时间区间（秒）
    -> [[start, end], ...]。

    B 站观众会在恰饭段刷"恰饭/广告/种草"等词，比外网 SponsorBlock 的
    YouTube 众包库更适配 B 站。把命中时间戳按 `window` 秒聚成连续区间。
    """
    kws = keywords or _SPONSOR_KW
    hits = [float(d["t"]) for d in danmaku
            if any(k in (d.get("text") or "") for k in kws)]
    if not hits:
        return []
    hits.sort()
    segs = []
    cur = [hits[0], hits[0]]
    for t in hits[1:]:
        if t - cur[1] <= window:
            cur[1] = t
        else:
            segs.append(cur)
            cur = [t, t]
    segs.append(cur)
    result = []
    for s in segs:
        if (s[1] - s[0]) >= window or len(hits) >= min_hits:
            result.append([max(0.0, s[0] - 5.0), s[1] + 10.0])
    return result


_LINE_TS_RE = re.compile(
    r"^(?:\[(\d{1,2}:\d{2}(?::\d{2})?)\]|(\d{1,2}:\d{2}(?::\d{2})?)\s*-\s*)\s*(.*)$")


def _line_seconds(line: str):
    """从转录行提取秒数。支持 'mm:ss - text'（字幕/ASR）与 '[mm:ss] text'
    （官方 AI 字幕）两种格式。无时间戳返回 (None, line)。"""
    m = _LINE_TS_RE.match(line)
    if not m:
        return None, line
    ts = m.group(1) or m.group(2)
    parts = [int(x) for x in ts.split(":")]
    weights = [3600, 60, 1][-len(parts):]
    secs = sum(p * w for p, w in zip(parts, weights))
    return secs, m.group(3)


def clean_transcript(transcript: str, drop_ranges=None, intro_skip: int = 0,
                     outro_skip: int = 0, duration: float = None) -> tuple:
    """#4 物理去除转录行：落在 drop_ranges（[start,end] 秒）、位于前 intro_skip 秒、
    或位于后 outro_skip 秒（需知 duration）。用于（clean 子命令）硬去除赞助/广告文本。
    返回 (cleaned_text, dropped_count)。无时间戳的行一律保留。
    """
    drop_ranges = drop_ranges or []
    kept, dropped = [], 0
    for line in transcript.splitlines():
        secs, _ = _line_seconds(line)
        if secs is None:
            kept.append(line)
            continue
        skip = False
        if intro_skip and secs < intro_skip:
            skip = True
        if outro_skip and duration and secs > (duration - outro_skip):
            skip = True
        for s, e in drop_ranges:
            if s <= secs <= e:
                skip = True
                break
        if skip:
            dropped += 1
        else:
            kept.append(line)
    return "\n".join(kept), dropped


def build_clean_instruction(clean_ranges: list = None) -> str:
    """#4 生成"内容净化"规则文本（用于原生/brief 模式，交给 agent LLM 执行跳过）。
    要求模型基于内容判断剔除片头口播、片尾 CTA、赞助/广告口播；若有弹幕检出的
    疑似赞助区间，一并列出供参考。
    """
    lines = [
        "**内容净化（请主动剔除，不要写入正文）**：整理笔记时请跳过以下非干货内容：",
        "  - 片头：开场寒暄、频道口播、\"欢迎来到/关注我\"、节目/栏目预告等；",
        "  - 片尾：求点赞/投币/收藏/三连、关注引导、下期预告、\"记得订阅\"等 CTA；",
        "  - 赞助/广告：品牌口播、商品推广、优惠券/领券/下单引导、\"恰饭\"段等。",
        "判断依据：以转录内容本身的明显口播套话为主，结合下方弹幕/信号为辅。"
        "宁可漏删也不要误删正常讲解；若某段疑似赞助但拿不准，可保留一句"
        "「（此处为赞助口播，已略）」的标注，不要展开。",
    ]
    if clean_ranges:
        seg = "\n".join(
            f"  - 约 {fmt_time(s)}–{fmt_time(e)}（弹幕出现\"恰饭/广告\"等信号）"
            for s, e in clean_ranges)
        lines.append("弹幕信号提示以下疑似赞助/广告时间区间，优先跳过：\n" + seg)
    return "\n".join(lines)


def suggest_screenshot_timestamps(outline: list = None, danmaku: list = None,
                                  duration: float = None, n: int = 6) -> list:
    """#2 Suggest screenshot times (seconds). Priority:
       1) official AI outline chapter starts (most reliable);
       2) danmaku-density peaks (high-energy moments);
       3) uniform sampling as a last resort.
    """
    if outline:
        ts = sorted({int(seg.get("timestamp", 0))
                     for seg in outline if seg.get("timestamp", 0) >= 0})
        if ts:
            return ts[:n] if n else ts
    if danmaku and duration:
        bucket = 15
        counts = {}
        for d in danmaku:
            b = int(d["t"] // bucket)
            counts[b] = counts.get(b, 0) + 1
        if counts:
            top = sorted(counts.items(), key=lambda kv: kv[1], reverse=True)[:n]
            return sorted(int(b * bucket + bucket / 2) for b, _ in top)
    if duration and n:
        step = duration / (n + 1)
        return [int(step * (i + 1)) for i in range(n)]
    return []


def markers_from_timestamps(seconds_list: list) -> str:
    """Turn a list of seconds into *Screenshot-[MM:SS]* marker lines."""
    return "\n".join(f"*Screenshot-[{fmt_time(s)}]*" for s in seconds_list)


def _ts_seconds(ts: str) -> int:
    parts = [int(x) for x in ts.split(":")]
    if len(parts) == 3:
        return parts[0] * 3600 + parts[1] * 60 + parts[2]
    return parts[0] * 60 + parts[1]


def _video_url_at(base_url: str, seconds: int) -> str:
    from urllib.parse import urlparse, urlencode, parse_qsl, urlunparse
    u = urlparse(base_url)
    q = dict(parse_qsl(u.query))
    if "youtube.com" in u.netloc or "youtu.be" in u.netloc:
        q["t"] = f"{seconds}s"
    else:  # bilibili and most others accept ?t=seconds
        q["t"] = str(seconds)
    return urlunparse(u._replace(query=urlencode(q)))


_CONTENT_MARK_RE = re.compile(r"\*Content-\[(\d{1,2}:\d{2}(?::\d{2})?)\]\*?")
_BARE_TS_RE = re.compile(r"(?<![-\w\[])\[(\d{1,2}:\d{2}(?::\d{2})?)\](?!\()")


def linkify_timestamps(md_text: str, video_url: str) -> str:
    """#3 Convert *Content-[MM:SS]* markers and bare [MM:SS] tokens into
    clickable links (video?t=seconds). Leaves *Screenshot-* markers untouched
    so screenshot embedding still works if run afterwards."""
    def _mk(ts):
        return f"[{ts}]({_video_url_at(video_url, _ts_seconds(ts))})"
    md_text = _CONTENT_MARK_RE.sub(lambda m: "*Content-" + _mk(m.group(1)) + "*",
                                   md_text)
    md_text = _BARE_TS_RE.sub(lambda m: _mk(m.group(1)), md_text)
    return md_text


def linkify_note(note_path: str, video_url: str, out_path: str = None) -> str:
    """Read a Markdown note, make its timestamps clickable, write it back.
    Returns the output path. Thin file wrapper around linkify_timestamps (#3)."""
    src = Path(note_path)
    dst = Path(out_path) if out_path else src
    text = src.read_text(encoding="utf-8")
    dst.write_text(linkify_timestamps(text, video_url), encoding="utf-8")
    log(f"时间戳已转为可点击跳转链接 -> {dst}")
    return str(dst)


def _build_video_format(quality: str) -> str:
    """Map a quality preset to a yt-dlp -f selector for the screenshot/embedding
    video download. 720 is the default (clear enough for screenshots, smaller
    footprint); 1080/best only resolve when a logged-in cookie is supplied
    (Bilibili serves >=1080p only to authenticated sessions)."""
    q = (quality or "720").strip().lower()
    if q in ("best", "max"):
        return "bestvideo[ext=mp4]/bestvideo/best"
    if q in ("1080", "1080p"):
        return "bestvideo[height<=1080][ext=mp4]/bestvideo[ext=mp4]/bestvideo/best"
    # default 720
    return "bestvideo[height<=720][ext=mp4]/bestvideo[ext=mp4]/bestvideo/best"


def fetch(url: str, workdir: Path, lang: str = "zh",
          with_video: bool = False, video_quality: str = "720") -> dict:
    """Download best audio + subtitles. Returns dict with paths and metadata.
    When `with_video` is True, also downloads the best video stream to
    `video.mp4` so frames can later be extracted for screenshot embedding."""
    workdir.mkdir(parents=True, exist_ok=True)
    ytdlp = _ytdlp_cmd()
    ffmpeg = _require("ffmpeg")

    # 1) metadata
    log("Fetching video metadata ...")
    meta = subprocess.run(
        ytdlp + _ytdlp_extra_opts() + ["-J", "--no-warnings", url],
        capture_output=True, text=True, check=True,
    )
    info = json.loads(meta.stdout)
    title = info.get("title", "unknown")
    tags = ", ".join(info.get("tags", [])[:10]) if info.get("tags") else ""
    log(f"Title: {title}")

    audio_path = workdir / "audio.mp3"
    sub_path = workdir / "subtitle.vtt"
    video_path = workdir / "video.mp4"

    # 2) audio (bestaudio -> mp3)
    log("Downloading audio ...")
    subprocess.run(
        ytdlp + _ytdlp_extra_opts() + ["-x", "--audio-format", "mp3",
                 "--audio-quality", "5", "-o", str(audio_path),
                 "--no-warnings", url],
        check=True,
    )

    # 3) subtitles (prefer manual, fall back to auto) — subtitle priority
    log("Trying subtitles (subtitle-priority) ...")
    for attempt in (
        ["--write-subs", "--sub-langs", lang, "--sub-format", "vtt"],
        ["--write-auto-subs", "--sub-langs", lang, "--sub-format", "vtt"],
    ):
        try:
                           subprocess.run(
                ytdlp + _ytdlp_extra_opts() + attempt
                + ["--skip-download",
                                   "-o", str(workdir / "subtitle.%(ext)s"),
                                   "--no-warnings", url],
                check=True, capture_output=True, text=True,
            )
        except subprocess.CalledProcessError:
            continue
        found = sorted(workdir.glob("subtitle*.vtt"))
        if found:
            shutil.move(str(found[0]), str(sub_path))
            # rename any extras
            for extra in sorted(workdir.glob("subtitle*.vtt")):
                if extra != sub_path:
                    extra.unlink()
            break

    used_subtitle = sub_path.exists()
    log("Subtitle found." if used_subtitle else "No subtitle; will transcribe audio.")

    video_file = None
    if with_video:
        log("Downloading video (for screenshot embedding) ...")
        try:
            subprocess.run(
                ytdlp + _ytdlp_extra_opts() + ["-f", _build_video_format(video_quality),
                         "-o", str(video_path), "--no-warnings", url],
                check=True, capture_output=True, text=True,
            )
            if video_path.exists():
                video_file = str(video_path)
                log(f"Video saved: {video_file}")
            else:
                log("Video download produced no file; screenshot embedding "
                    "will be unavailable.")
        except subprocess.CalledProcessError as e:
            log(f"Video download failed: {e}; screenshot embedding "
                "will be unavailable.")

    return {
        "title": title,
        "tags": tags,
        "audio": str(audio_path),
        "subtitle": str(sub_path) if used_subtitle else None,
        "video": video_file,
        "used_subtitle": used_subtitle,
    }


def fetch_local(video_path: str, workdir: Path,
                with_video: bool = False) -> dict:
    """Local video file -> extract audio; no subtitles (transcribe instead).
    When `with_video` is True, the original video path is kept for screenshot
    embedding."""
    workdir.mkdir(parents=True, exist_ok=True)
    ffmpeg = _require("ffmpeg")
    video_path = Path(video_path)
    title = video_path.stem
    audio_path = workdir / "audio.mp3"
    log(f"Extracting audio from local file: {video_path.name}")
    subprocess.run(
        [ffmpeg, "-y", "-i", str(video_path), "-vn", "-ac", "1",
         "-ar", "16000", "-b:a", "64k", str(audio_path)],
        check=True, capture_output=True, text=True,
    )
    return {"title": title, "tags": "", "audio": str(audio_path),
            "subtitle": None, "used_subtitle": False,
            "video": str(video_path) if with_video else None}


# --------------------------------------------------------------------------
# transcript parsing / formatting
# --------------------------------------------------------------------------
def parse_vtt(path: Path) -> str:
    """Parse a .vtt subtitle file into 'mm:ss - text' lines."""
    text = path.read_text(encoding="utf-8", errors="ignore")
    blocks = re.split(r"\n\s*\n", text)
    out = []
    ts_re = re.compile(r"(\d{1,2}:\d{2}:\d{2})\.?\d*\s*-->\s*\d")
    for block in blocks:
        lines = [l.strip() for l in block.splitlines() if l.strip()]
        if not lines:
            continue
        m = ts_re.search(lines[0])
        if not m:
            continue
        start = sum(int(x) * f for x, f in zip(m.group(1).split(":"), [3600, 60, 1]))
        content = " ".join(lines[1:]) if len(lines) > 1 else ""
        content = re.sub(r"<[^>]+>", "", content).strip()
        if content:
            out.append(f"{fmt_time(start)} - {content}")
    return "\n".join(out)


def transcribe(audio_path: str, backend: str, **kwargs) -> str:
    """Transcribe audio -> 'mm:ss - text' transcript. Online engines only
    (local faster-whisper/mlx-whisper were removed)."""
    if backend in ("bcut", "biji"):
        return _transcribe_bcut(audio_path, **kwargs)
    if backend == "kuaishou":
        return _transcribe_kuaishou(audio_path, **kwargs)
    raise ValueError(
        f"不支持的转写引擎 '{backend}'。在线引擎仅支持："
        f"bcut(必剪, 免费无密钥) / kuaishou(快手, 免费无密钥)。"
        f"本地 fast-whisper/mlx-whisper 已移除。")


# --------------------------------------------------------------------------
# subtitle + ASR cross-calibration
# --------------------------------------------------------------------------
_TS_RE = re.compile(r"^(\d{1,2}:\d{2}:\d{2}|\d{1,2}:\d{2})\s*-\s*(.*)$")


def _parse_ts(ts: str) -> float:
    parts = [int(x) for x in ts.split(":")]
    if len(parts) == 3:
        return parts[0] * 3600 + parts[1] * 60 + parts[2]
    return parts[0] * 60 + parts[1]


def parse_transcript_to_segments(transcript: str) -> list:
    """Parse a 'mm:ss - text' transcript into a list of (start_sec, text) tuples."""
    segs = []
    for line in transcript.splitlines():
        line = line.strip()
        if not line:
            continue
        m = _TS_RE.match(line)
        if not m:
            continue
        ts = _parse_ts(m.group(1))
        text = m.group(2).strip()
        if text:
            segs.append((ts, text))
    return segs


# Online, keyless transcribers. The language-aware router
# (route_transcribe) picks the primary/fallback order per specified language
# (zh/en -> bcut first; other langs -> kuaishou first). biji is an alias of bcut.
ONLINE_FREE_ENGINES = ("bcut", "kuaishou")


def transcribe_online_free(audio_path: str, preferred: str = "bcut", **kwargs) -> str:
    """Try online-free transcribers in order, auto-switching on failure."""
    pool = list(ONLINE_FREE_ENGINES)
    if preferred in pool:
        order = [preferred] + [e for e in pool if e != preferred]
    else:
        order = list(pool)
    errors = []
    for eng in order:
        try:
            log(f"Transcribing via online-free engine: {eng}")
            text = transcribe(audio_path, eng, **kwargs)
            if text and text.strip():
                return text
            log(f"{eng} returned empty transcript; trying next engine.")
        except Exception as e:  # noqa: BLE001 - fall back to the next engine
            log(f"{eng} failed: {e}")
            errors.append(f"{eng}: {e}")
    raise RuntimeError(
        "All online-free transcribers failed:\n" + "\n".join(errors))


def dispatch_transcribe(audio_path: str, backend: str, **kwargs) -> str:
    """Pick single or fallback strategy based on the chosen backend."""
    if backend == "biji":
        backend = "bcut"  # 必剪(在线) == BCut; joins the online-free fallback pool
    if backend == "auto" or backend in ONLINE_FREE_ENGINES:
        preferred = ONLINE_FREE_ENGINES[0] if backend == "auto" else backend
        return transcribe_online_free(audio_path, preferred=preferred, **kwargs)
    return transcribe(audio_path, backend, **kwargs)# Languages where bcut (必剪) is reliable; everything else goes to kuaishou
# first (Kuaishou transcribed ja/ko correctly in testing).
BCUT_LANGS = {"zh", "en"}


def route_transcribe(audio_path: str, lang: str = None, **kwargs) -> str:
    """Language-aware transcription router used by `run` (auto mode).

    1) Use `lang` when supplied; otherwise default to Chinese (bcut-first).
    2) zh/en -> bcut first, fallback kuaishou; other langs -> kuaishou first,
       fallback bcut.
    3) If both online engines fail, raise RuntimeError carrying advice.
    """
    detected = lang  # 本地 whisper 已移除，按 --lang 或默认中文路由
    if detected:
        log(f"Routing transcript for language: {detected}")
    else:
        log("未指定语种；按中文(bcut 优先) 路由。")
    # Plan G：缓存命中（键 = 音频路径+大小+mtime+语种，廉价指纹避免重读文件）
    try:
        st = os.stat(audio_path)
        ck = _cache_key("tr", audio_path, st.st_size, int(st.st_mtime), detected or "")
        hit, cached = _cache_get("transcribe", ck)
        if hit and isinstance(cached, str) and cached.strip():
            log("[cache] transcribe 命中")
            return cached
    except Exception:  # noqa: BLE001
        ck = None

    primary, fallback = ("bcut", "kuaishou") \
        if (detected in BCUT_LANGS or not detected) else ("kuaishou", "bcut")

    errors = []
    for eng in (primary, fallback):
        try:
            log(f"Router: transcribing via {eng} (lang={detected or 'auto'})")
            text = transcribe(audio_path, eng, **kwargs)
            if text and text.strip():
                if ck:
                    _cache_set("transcribe", ck, text)  # Plan G：成功才缓存
                return text
            log_recovery("F4", f"{eng} 返回空，切换下一引擎")
        except Exception as e:
            log_recovery("F3", f"{eng} 失败：{e}")
            errors.append(f"{eng}: {e}")
    advice = (
        "在线免费转写引擎（bcut/快手）均失败，该语种暂不被支持或网络异常。\n"
        "请选择后续方案：\n"
        "  1) 手动指定语种重试：--lang <code> （如 --lang ja）\n"
        "调试信息：\n" + "\n".join(errors)
    )
    raise RuntimeError(advice)


def _transcribe_bcut(audio_path: str, **_):
    """Bilibili BCut (必剪) online ASR — keyless, free. Faithful to
    BiliNote/backend/app/transcriber/bcut.py: request upload -> upload parts
    via boss -> commit -> create task -> poll until done -> parse utterances
    (timestamps are milliseconds)."""
    import requests, json as _json, time as _time
    audio = Path(audio_path)
    API_BASE = "https://member.bilibili.com/x/bcut/rubick-interface"
    REQ_UPLOAD = API_BASE + "/resource/create"
    COMMIT_UPLOAD = API_BASE + "/resource/create/complete"
    CREATE_TASK = API_BASE + "/task"
    QUERY_RESULT = API_BASE + "/task/result"
    headers = {
        "User-Agent": "Bilibili/1.0.0 (https://www.bilibili.com)",
        "Content-Type": "application/json",
    }
    session = requests.Session()
    fb = audio.read_bytes()
    size = len(fb)

    # 1) request upload
    resp = session.post(REQ_UPLOAD,
        data=_json.dumps({"type": 2, "name": "audio.mp3", "size": size,
                          "ResourceFileType": "mp3", "model_id": "8"}),
        headers=headers, timeout=30)
    resp.raise_for_status()
    d = resp.json()["data"]
    in_boss_key = d["in_boss_key"]; resource_id = d["resource_id"]
    upload_id = d["upload_id"]; upload_urls = d["upload_urls"]
    per_size = d["per_size"]; clips = len(upload_urls)

    # 2) upload parts
    etags = []
    for clip in range(clips):
        start = clip * per_size
        end = min((clip + 1) * per_size, size)
        r = session.put(upload_urls[clip], data=fb[start:end],
                        headers={"Content-Type": "application/octet-stream"},
                        timeout=120)
        r.raise_for_status()
        etags.append(r.headers.get("Etag", "").strip('"'))

    # 3) commit upload
    resp = session.post(COMMIT_UPLOAD,
        data=_json.dumps({"InBossKey": in_boss_key, "ResourceId": resource_id,
                          "Etags": ",".join(etags), "UploadId": upload_id,
                          "model_id": "8"}),
        headers=headers, timeout=30)
    resp.raise_for_status()
    cd = resp.json()
    if cd.get("code") != 0:
        raise RuntimeError(f"BCut commit failed: {cd.get('message', 'unknown')}")
    download_url = cd["data"]["download_url"]

    # 4) create task
    resp = session.post(CREATE_TASK,
        json={"resource": download_url, "model_id": "8"}, headers=headers,
        timeout=30)
    resp.raise_for_status()
    td = resp.json()
    if td.get("code") != 0:
        raise RuntimeError(
            f"BCut task create failed: {td.get('message', 'unknown')}")
    task_id = td["data"]["task_id"]

    # 5) poll result
    task_resp = None
    for _ in range(500):
        resp = session.get(QUERY_RESULT,
            params={"model_id": 7, "task_id": task_id}, headers=headers,
            timeout=30)
        resp.raise_for_status()
        task_resp = resp.json()["data"]
        if task_resp["state"] == 4:    # done
            break
        if task_resp["state"] == 3:    # failed
            raise RuntimeError("BCut ASR task failed (state=3).")
        _time.sleep(1)
    if not task_resp or task_resp["state"] != 4:
        raise RuntimeError("BCut ASR did not complete.")

    result_json = _json.loads(task_resp["result"])
    out = []
    for u in result_json.get("utterances", []):
        text = (u.get("transcript") or "").strip()
        if not text:
            continue
        start = float(u.get("start_time", 0)) / 1000.0
        out.append(f"{fmt_time(start)} - {text}")
    if not out:
        raise RuntimeError("BCut returned no utterances.")
    return "\n".join(out)


def _md5(path: Path) -> str:
    import hashlib
    h = hashlib.md5()
    h.update(path.read_bytes())
    return h.hexdigest()


def _transcribe_kuaishou(audio_path: str, **_):
    """Kuaishou (快手) online ASR — keyless, free, single synchronous request.
    Faithful to BiliNote/backend/app/transcriber/kuaishou.py: a multipart POST
    with form field typeId=1 returns the full subtitle list (no polling).
    start/end times are in seconds. Best for non-Chinese audio (ja/ko/...);
    zh/en are also supported. If the server returns a non-zero business code
    (e.g. 501 "效果subtitle_generate禁用") this raises so the caller can
    fall back to another engine."""
    import requests
    audio = Path(audio_path)
    api = "https://ai.kuaishou.com/api/effects/subtitle_generate"
    file_binary = audio.read_bytes()
    files = [("file", (audio.name, file_binary, "audio/mpeg"))]
    resp = requests.post(api, data={"typeId": "1"}, files=files, timeout=300)
    resp.raise_for_status()
    result = resp.json()
    if result.get("code", 0) != 0 or "data" not in result:
        raise RuntimeError(
            f"Kuaishou ASR error (code={result.get('code')}): "
            f"{result.get('msg') or result.get('message') or 'unknown'}")
    segments = []
    for u in result.get("data", {}).get("text", []):
        text = (u.get("text") or "").strip()
        if not text:
            continue
        start = float(u.get("start_time", 0))
        segments.append(f"{fmt_time(start)} - {text}")
    if not segments:
        raise RuntimeError("Kuaishou ASR returned no text segments.")
    return "\n".join(segments)


# --------------------------------------------------------------------------
# summarize
# --------------------------------------------------------------------------
def build_user_prompt(title: str, tags: str, segment_text: str,
                      time_markers: bool, screenshots: bool, ai_summary: bool,
                      style: str = DEFAULT_STYLE, media: bool = True,
                      images: list = None, include_images: bool = False,
                      multimodal_frames: list = None) -> str:
    tpl = load_prompt_template()
    body = tpl.format(video_title=title, tags=tags or "（无）",
                      segment_text=segment_text)
    extra = ""
    if not media:
        extra += (
            "\n\n> 注意：本次输入为文章 / 文本 / 文档（非视频），请直接基于"
            "内容本身整理笔记；不要使用视频时间标记 `*Content-[mm:ss]*` "
            "或截图占位 `*Screenshot-[mm:ss]*`。\n"
        )
    if screenshots:
        extra += SCREENSHOT
    if time_markers:
        extra += LINK
    if ai_summary:
        extra += AI_SUM
    if include_images and images:
        img_list = "\n".join(f"  {i+1}. {u}" for i, u in enumerate(images))
        extra += (
            "\n\n## 原文配图（可选嵌入）\n"
            f"原文包含 {len(images)} 张配图，URL 如下（编号 1..{len(images)}）：\n"
            f"{img_list}\n"
            "若使用者选择「嵌入原文图片」，请在笔记中合适的段落之后插入 "
            "`*Image-1*`、`*Image-2*` …… 等标记来引用对应图片（编号与上面一致）。"
            "这些标记稍后会被自动下载并替换为真实图片。\n"
        )
    if multimodal_frames:
        frame_list = "\n".join(f"  {i+1}. {p}"
                               for i, p in enumerate(multimodal_frames))
        extra += (
            "\n\n## 视频关键帧（多模态视觉输入）\n"
            f"以下 {len(multimodal_frames)} 帧是按时间均匀抽取的视频关键帧（编号 1..{len(multimodal_frames)}）：\n"
            f"{frame_list}\n"
            "这些帧与本片段转录**同时输入给你**。请结合**帧画面中的文字、图表、操作演示、PPT 内容**"
            "与转录文本一起理解视频内容并在笔记中体现。例如：\n"
            "- 若帧中出现关键数据 / 图表 / PPT 要点，请在相关小节中写出；\n"
            "- 若帧中展示了 UI 操作或演示画面，请在笔记中描述步骤；\n"
            "- 在对应小节末尾插入 `*Screenshot-[mm:ss]*` 占位（会被替换为实帧）。\n"
        )
    label, instruction = STYLES.get(style, STYLES[DEFAULT_STYLE])
    extra += (
        f"\n\n## 报告风格要求（必须严格遵循）\n"
        f"请按以下风格生成笔记：{instruction}\n"
    )
    detail = get_style_spec(style)
    if detail:
        extra += (
            f"\n\n## 平台风格详细规范（{label}）\n"
            f"严格按以下平台写法执行（结构化模板，直接套用）：\n\n{detail}\n"
        )
    return body + extra


def build_outline_prompt(transcript: str, title: str, style: str,
                         media: bool = True) -> str:
    """Plan J：生成「大纲前置」简报——先出论点树供使用者确认，再填正文。

    返回一段自足 Markdown，指导 agent 从转录抽取 H1+H2+要点+时间锚的论点树。"""
    label = STYLES.get(style, STYLES[DEFAULT_STYLE])[0]
    src = "视频转录" if media else "原文内容"
    return (
        f"# BiliNote 大纲前置简报（Plan J · 论点树）\n\n"
        f"为《{title}》先生成一份**论点树大纲**，供使用者确认后再据大纲填充正文。\n\n"
        f"## 要求\n"
        f"1. 从下方{src}中识别全部主要论点，按逻辑分组。\n"
        f"2. 输出结构：`# 标题` + `## 各节`（每节 2-3 句要点）+ 关键时间锚"
        f"{' `*Content-[mm:ss]*`' if media else ''}。\n"
        f"3. 风格：{label}（与最终笔记一致）。\n"
        f"4. 只列要点不展开正文；覆盖全部论点避免遗漏。\n"
        f"5. 写入 `<stem>.outline.md` 后告知使用者确认；确认后再写正文 note.md。\n\n"
        f"## {src}\n---\n\n{transcript}\n"
    )


def _write_brief(transcript: str, title: str, tags: str, time_markers: bool,
                 screenshots: bool, ai_summary: bool, out: str,
                 style: str = DEFAULT_STYLE, video_path: str = None,
                 media: bool = True, images: list = None,
                 include_images: bool = False,
                 multimodal_frames: list = None,
                 enrich: dict = None,
                 outline_first: bool = False) -> str:
    """agent-native mode: emit a self-contained brief the agent's own
    LLM uses to produce the note. No external API key required.
    When media=False (article/text/document input), no time markers or
    screenshots are expected. `images`/`include_images` enable embedding the
    article's original pictures (non-video input only).
    `multimodal_frames` is a list of local frame paths that will be passed
    to the agent LLM as visual input for multimodal video understanding."""
    p = Path(out)
    brief = p.parent / (p.stem + ".brief.md")
    rules = [
        "仅返回 Markdown 正文，不要包裹在代码块中。",
        "编号标题用 `1\\. **内容**` 或 `## 1. 内容` 防误渲染。",
        "使用中文撰写，专有名词/术语保留英文。",
        "去除广告、填充词、问候语；保留关键事实、示例、结论、建议。",
        "数学公式以 LaTeX 语法呈现。",
    ]
    if media:
        rules.append("开启时间标记：相关内容后加 `*Content-[mm:ss]*`。")
    else:
        rules.append("本次输入为文章/文本/文档（非视频）：无需时间标记与截图"
                     "占位，直接按内容整理笔记。")
    if screenshots:
        rules.append("开启截图占位：相关内容后加 `*Screenshot-[mm:ss]*`。")
    if ai_summary:
        rules.append("笔记开头用一两句话给出 AI 总结。")
    label, instruction = STYLES.get(style, STYLES[DEFAULT_STYLE])
    style_list = " / ".join(f"{k}({v[0]})" for k, v in STYLES.items())
    rules.append(f"**报告风格 = {label}**：{instruction}；可选风格：{style_list}。")
    detail = get_style_spec(style)
    if detail:
        rules.append(
            f"**平台风格详细规范（{label}）**：严格按以下平台写法执行"
            f"（结构化模板，直接套用）：\n\n{detail}")
    if include_images and images:
        img_block = "\n".join(f"  {i+1}. {u}" for i, u in enumerate(images))
        rules.append(
            f"原文包含 {len(images)} 张配图，URL 如下（编号 1..{len(images)}）：\n"
            f"{img_block}\n"
            "若使用者选择嵌入图片，请在笔记合适段落之后插入 `*Image-1*`、"
            "`*Image-2*` …… 标记引用对应图片（编号与上面一致）；生成后运行 "
            "`embed-article-images` 将其下载并替换为真实图片，图片存于笔记同目录 `images/`。"
        )
    if multimodal_frames:
        frame_block = "\n".join(f"  {i+1}. {p}"
                                for i, p in enumerate(multimodal_frames))
        rules.append(
            f"多模态视频理解：以下 {len(multimodal_frames)} 帧是按时间均匀抽取的视频关键帧：\n"
            f"{frame_block}\n"
            "这些帧已作为**视觉输入**与本简报同时提供给当前 LLM。请结合帧画面中的文字、"
            "图表、PPT、操作演示等视觉信息与转录内容一起生成笔记；在相关小节末尾插入 "
            "`*Screenshot-[mm:ss]*` 占位符（会被替换为实帧）。"
        )
    # ---- borrowed enrichers (#0 AI summary / #1 danmaku / #2 frames /
    #      #4 sponsor / #5 translate) as extra context + rules ----
    enrich = enrich or {}
    if enrich.get("translate"):
        rules.append(
            "**外语原声翻译**：本视频为外语内容，转录为原文。请在整理笔记时"
            "把要点翻译成简体中文（专有名词保留英文），必要时可在关键句后用"
            "括号附上原文；最终笔记正文以中文为主。")
    if enrich.get("ai_summary_md"):
        rules.append(
            "已附「B 站官方 AI 总结」（见下方参考区）：其 summary 可作为开头总结的"
            "重要参考，其时间轴大纲(*Content-[mm:ss]*)可用于组织小节与定位；"
            "请与转录内容交叉印证，不要照搬，以转录为准、AI 总结为辅。")
    if enrich.get("danmaku_ctx"):
        rules.append(
            "已附「弹幕上下文」（见下方参考区）：可用于判断观众关注的高能片段、"
            "梗与高赞观点，辅助你标注重点小节；勿把弹幕原文直接抄进正文。")
    if enrich.get("comments_ctx"):
        if enrich.get("comments_in_report"):
            rules.append(
                "已附「评论区高能精选」（见下方参考区）：请在笔记末尾新增一节 "
                "「## 评论区高能精选」，按参考区给出的高赞评论逐条列出（保留 "
                "@用户名 与赞数），作为观众态度补充；不要改写原意，并标注"
                "「不代表笔记立场」。")
        else:
            rules.append(
                "已附「评论区精选」（见下方参考区）：可用于感知观众态度与争议点，"
                "辅助你把握语气；勿把评论原文直接抄进正文。")
    if enrich.get("suggest_ts") and screenshots:
        marker_lines = markers_from_timestamps(enrich["suggest_ts"])
        rules.append(
            "已根据官方大纲/弹幕高能点自动推荐以下截图时间点，请优先在对应小节末尾"
            "放置这些 `*Screenshot-[mm:ss]*` 占位（可按需增减）：\n" + marker_lines)
    if enrich.get("clean_instruction"):
        rules.append(enrich["clean_instruction"])
    if enrich.get("linkify") and enrich.get("video_url") and media:
        rules.append(
            "生成后可让时间标记可点击跳转：对该笔记运行 "
            f"`python scripts/bilinote.py linkify --note {out} "
            f"--url \"{enrich['video_url']}\"`，会把 `*Content-[mm:ss]*` 与裸 "
            "`[mm:ss]` 转成指向视频对应时间点的链接。")
    # Native-mode interaction: confirm style + image preference with the user
    img_opt = ("是否抽取并嵌入视频截图（--screenshots）"
               if media else "是否下载并嵌入原文配图（--images）")
    rules.append(
        "交互确认：笔记将由 agent 当前 LLM 生成。若调用 run 时未显式指定 "
        f"--style、{img_opt}、--export，请在生成前先用 AskUserQuestion 向使用者确认：\n"
        f"(1) 报告风格 {len(STYLES)} 选 1；\n"
        "(2) " + img_opt + "；\n"
        "(3) 输出格式多选（可勾选多个）：pdf / docx / 思维导图(mindmap)。\n"
        "得到明确答复后，把对应参数（--style X --images/--screenshots --export pdf,docx,mindmap）"
        "传回 run 调用再生成并嵌入。若使用者选不出默认给 pdf+docx。"
    )
    if screenshots and video_path:
        rules.append(
            "视频已下载并保存在：\n"
            f"  `{video_path}`\n"
            "若使用者希望把截图**真正嵌入**笔记（抽取视频帧），生成含 "
            "`*Screenshot-[mm:ss]*` 标记的笔记后，运行：\n"
            f"  `python scripts/bilinote.py embed-screenshots --note {out} "
            f"--video {video_path} --out {out}`\n"
            "抽取出的图片会保存在笔记同目录 `screenshots/` 下。"
        )
    if outline_first:
        rules.append(
            "**大纲前置（Plan J）**：在写正文前，先生成一份论点树大纲"
            "（H1 标题 + H2 各节 + 每节 2-3 句要点 + 关键时间锚 `*Content-[mm:ss]*`），"
            "写入 `<stem>.outline.md` 供使用者确认后再据大纲填充正文。"
            "大纲应覆盖转录全部主要论点，避免遗漏；使用者未确认前不写正文。"
        )
    rules_text = ("## 生成规则（必须严格遵循）\n"
                  "完整规则见技能 references/note_prompt.md。关键要求：\n"
                  + "\n".join(f"{i+1}. {r}" for i, r in enumerate(rules)) + "\n")
    if media:
        intro = ("请使用**当前对话的 LLM** 生成结构化视频笔记，"
                 "无需任何外部 API Key。\n\n")
        section_header = "## 视频转录\n---\n"
    else:
        intro = ("请使用**当前对话的 LLM** 生成结构化笔记"
                 "（本次输入为文章/文本/文档，非视频），无需任何外部 API Key。\n\n")
        section_header = "## 原文内容\n---\n"
    ref_blocks = ""
    if enrich.get("ai_summary_md"):
        ref_blocks += ("## 参考区 · B 站官方 AI 总结\n---\n"
                       + enrich["ai_summary_md"] + "\n---\n\n")
    if enrich.get("danmaku_ctx"):
        ref_blocks += ("## 参考区 · 弹幕上下文（按时间）\n---\n"
                       + enrich["danmaku_ctx"] + "\n---\n\n")
    if enrich.get("comments_ctx"):
        ref_blocks += ("## 参考区 · 评论区高能精选\n---\n"
                       + enrich["comments_ctx"] + "\n---\n\n")
    content = (
        "# BiliNote 笔记生成简报（agent 原生模式）\n\n"
        + intro +
        "## 元数据\n"
        f"- 标题: {title}\n"
        f"- 标签: {tags or '（无）'}\n"
        f"- 输出文件: {out}\n\n"
        f"{rules_text}\n"
        + ref_blocks
        + section_header +
        f"{transcript}\n"
        "---\n\n"
        f"生成后将该笔记写入：{out}\n"
    )
    brief.write_text(content, encoding="utf-8")
    # Plan J：大纲前置 —— 额外写一份 outline-brief 供 agent 先出论点树
    if outline_first:
        ob = p.parent / (p.stem + ".outline-brief.md")
        ob.write_text(build_outline_prompt(transcript, title, style, media),
                      encoding="utf-8")
        log(f"大纲前置：已生成 outline 简报 {ob}，请先确认论点树再填正文。")
    return str(brief)


def summarize(transcript: str, title: str, tags: str = "",
             time_markers: bool = True, screenshots: bool = False,
             ai_summary: bool = True, out: str = "./note.md",
             style: str = DEFAULT_STYLE,
             video_path: str = None, media: bool = True,
             images: list = None, include_images: bool = False,
             multimodal_frames: list = None,
             enrich: dict = None,
             outline_first: bool = False) -> tuple[str, str]:
    """Return ('brief', payload). payload is the path to a self-contained brief
    for the agent's OWN LLM to finalize the note. agent-native mode is the
    ONLY mode (no external LLM API key is used anywhere in this skill).
    `media` is False for article/text/document input (no timestamps/screenshots).
    `images`/`include_images` enable embedding the article's original pictures.
    `multimodal_frames` is a list of local frame paths for multimodal video
    understanding (passed as visual input alongside the transcript)."""
    label, _ = STYLES.get(style, STYLES[DEFAULT_STYLE])
    # agent-native mode: delegate final generation to the agent LLM.
    Path(out).parent.mkdir(parents=True, exist_ok=True)
    brief = _write_brief(transcript, title, tags, time_markers,
                         screenshots, ai_summary, out, style=style,
                         video_path=video_path, media=media,
                         images=images, include_images=include_images,
                         multimodal_frames=multimodal_frames,
                         enrich=enrich,
                         outline_first=outline_first)
    log(f"agent 原生模式（风格={label}）：已生成简报 {brief}，由 agent 当前 LLM 生成笔记。")
    return "brief", brief


# --------------------------------------------------------------------------
# Plan C: 质量闸（review）—— 派生内容前，先校验笔记对源的「覆盖度 + 一致性」，
# 堵住源头遗漏与幻觉，避免错误被放大到所有平台产物。
# --------------------------------------------------------------------------
REVIEW_INSTRUCTION = (
    "请对照【源转录】审阅【笔记】，输出一份 Markdown 质检报告，严格包含以下小节：\n\n"
    "## 笔记类型判定\n"
    "- 先判定本篇类型（金融/研报 · 学术/严肃 · 配图类 · 通用），据此套用下方对应《反例黑名单》做 pre-flight。\n\n"
    "## 反例黑名单 pre-flight\n"
    "- 按上一步类型，逐条核对黑名单，列出命中/未命中；命中的即为必须修的问题：\n"
    "  · 金融/研报：①单一数据源未交叉验证 ②未标注数据时间戳 ③基准标签未闭环 ④仅凭视频口述数据未检索核实\n"
    "  · 学术/严肃：①引用无出处 ②连续30字与源文本雷同未改写 ③编造URL\n"
    "  · 配图类：①相似卡片/盒子超3个未抽象 ②单一模块体积超画面20% ③缺原文指纹\n"
    "  · 通用：①破折号含全角—— ②外链不可点 ③直接引用超150字\n\n"
    "## 覆盖度检查\n"
    "- 列出源转录出现、但笔记遗漏或明显弱化的**关键论点/数据/结论**（逐条，附源转录依据），"
    "每条用 `[原文]` 标注可回溯性。若无遗漏写「无重大遗漏」。\n\n"
    "## 一致性 / 疑似幻觉\n"
    "- 列出笔记中**源转录并未支撑**的说法/数字/名称/因果，逐条标注 `[原文]`（可回溯）"
    "或 `[推断]`（推理无直接依据，权重自动×0.5）。若无写「未发现明显幻觉」。\n\n"
    "## 结构与可读性\n"
    "- 指出层级混乱、重复、逻辑跳跃等问题（可选，简明）。\n\n"
    "## 修订建议\n"
    "- 给出 3–7 条可执行的具体修改建议，每条标注 `[原文]`/`[检索]`/`[推断]`"
    "（`[推断]`项权重自动×0.5）。\n\n"
    "## 来源追溯小结\n"
    "- 汇总：`[原文]` X项 · `[检索]` Y项 · `[推断]` Z项；"
    "若 `[推断]`≥3 项，或金融类数据无 `[检索]` 验证，标记「🔴可信度不足」。\n\n"
    "## 质量评分\n"
    "- 覆盖度 / 准确性 / 结构 各给 1–5 分，并一句话总评。\n\n"
    "要求：只输出报告本身；引用证据时简短；不要重写整篇笔记；"
    "来源标签用方括号原样标注（如 `[原文]`）。"
)


def build_review_prompt(note_md: str, transcript: str) -> str:
    """拼接质检 user prompt（源转录在前、笔记在后，均可能被截断以控长度）。"""
    t = transcript.strip()
    n = note_md.strip()
    # 控制超长转录：保留头尾（论点常在首尾），中间省略
    max_t = 12000
    if len(t) > max_t:
        head = t[: max_t * 2 // 3]
        tail = t[-max_t // 3:]
        t = head + "\n\n……（中间转录省略）……\n\n" + tail
    return (
        f"{REVIEW_INSTRUCTION}\n\n"
        f"===== 源转录（事实基准）=====\n{t}\n\n"
        f"===== 待审阅笔记 =====\n{n}\n"
    )


def review_note(note_path: str, transcript_path: str = "", transcript: str = "",
                out: str = "") -> tuple[str, str]:
    """质量闸主函数（agent 原生模式，唯一模式）。返回 ('brief', payload)：
    payload 是质检简报路径，交 agent 当前 LLM 审阅后写出质检报告。"""
    note_md = Path(note_path).read_text(encoding="utf-8")
    if not transcript and transcript_path:
        transcript = Path(transcript_path).read_text(encoding="utf-8")
    if not transcript:
        raise ValueError("review 需要源转录：请传 --transcript 或 transcript 文本。")

    out = out or str(Path(note_path).with_suffix("").as_posix() + ".review.md")
    # agent-native mode: write a self-contained brief for the agent LLM.
    brief_path = str(Path(note_path).with_suffix("").as_posix()
                     + ".review-brief.md")
    content = (
        "# 质检简报（agent 原生模式）\n\n"
        "> 你（agent 当前 LLM）来充当质检员。请阅读下方【源转录】与【笔记】，"
        f"按要求输出质检报告，并写入：`{out}`。\n\n"
        f"{REVIEW_INSTRUCTION}\n\n"
        "---\n\n"
        "## 源转录（事实基准）\n\n"
        f"{transcript.strip()}\n\n"
        "---\n\n"
        "## 待审阅笔记\n\n"
        f"{note_md.strip()}\n"
    )
    Path(brief_path).write_text(content, encoding="utf-8")
    log(f"质量闸原生模式：已生成质检简报 {brief_path}，由 agent 当前 LLM 审阅。")
    return "brief", brief_path


def _maybe_review(note_path: str, transcript: str) -> None:
    """run 流程内的质检钩子（--review）：失败不阻断主流程。"""
    try:
        kind, payload = review_note(note_path, transcript=transcript)
        log(f"质量闸：已生成质检简报 -> {payload}（由 agent 当前 LLM 审阅）")
    except Exception as e:
        log(f"质量闸跳过（{e}）。")


# --------------------------------------------------------------------------
# Plan I: 质量闭环（review → auto-fix → re-review）
#   review 出报告后，自动定位问题段并回修，再复检直到达标或轮次上限。
# --------------------------------------------------------------------------


def parse_review_score(report: str):
    """从质检报告的「质量评分」段抽取分数，返回平均分(0-5)或 None。"""
    nums = [int(x) for x in re.findall(r"(\d)\s*[/／]\s*5", report)]
    return round(sum(nums) / len(nums), 1) if nums else None


def build_autofix_prompt(note_md: str, report: str, transcript: str) -> str:
    """拼接 auto-fix user prompt：只修报告指出的问题，保持其余不变。"""
    t = transcript.strip()
    if len(t) > 10000:
        t = t[:6000] + "\n\n……（省略）……\n\n" + t[-3000:]
    return (
        f"===== 源转录（事实基准，仅供核对）=====\n{t}\n\n"
        f"===== 质检报告（指出的问题）=====\n{report.strip()}\n\n"
        f"===== 原笔记（待修复）=====\n{note_md.strip()}\n\n"
        "请输出**修复后的完整笔记**（Markdown 正文，不要代码块包裹）。"
        "规则：①只修报告指出的遗漏/幻觉/错误；②其余内容原样保留；"
        "③不新增报告未提及的事实；④保持原风格与结构。"
    )


def auto_fix_note(note_path: str, report: str, transcript: str,
                  out: str = "") -> tuple[str, str]:
    """Plan I：据质检报告回修笔记（agent 原生模式，唯一模式）。
    返回 ('brief', payload)：payload 是修复简报路径，交 agent 当前 LLM 回修。"""
    note_md = Path(note_path).read_text(encoding="utf-8")
    out = out or str(Path(note_path).with_suffix("").as_posix() + ".autofixed.md")
    brief_path = str(Path(note_path).with_suffix("").as_posix()
                     + ".auto-fix-brief.md")
    content = (
        "# 自动修复简报（Plan I · agent 原生模式）\n\n"
        "> 你（agent 当前 LLM）来充当修复员。按下方【质检报告】指出的问题，"
        f"只修复问题、保持其余不变，把修复后的完整笔记写入：`{out}`。\n\n"
        f"{build_autofix_prompt(note_md, report, transcript)}\n"
    )
    Path(brief_path).write_text(content, encoding="utf-8")
    log(f"auto-fix 原生模式：已生成修复简报 {brief_path}")
    return "brief", brief_path


def review_note_with_autofix(note_path: str, transcript: str,
                             auto_fix: bool = False,
                             target_score: float = 4.5,
                             max_rounds: int = 2) -> dict:
    """Plan I 质量闭环（agent 原生模式，唯一模式）：review → auto-fix 简报。
    返回 {rounds, final_score, review_path, fixed_path, stopped_reason}。
    原生模式下 review 与 auto-fix 均产出简报交 agent 当前 LLM 执行，不做自动循环。"""
    result = {"rounds": 0, "final_score": None,
              "review_path": "", "fixed_path": "",
              "stopped_reason": ""}
    cur_note = note_path
    for r in range(1, max_rounds + 1):
        kind, payload = review_note(cur_note, transcript=transcript)
        result["review_path"] = payload
        result["rounds"] = r
        if kind != "review":
            result["stopped_reason"] = "native 模式：仅出 review+autofix 简报，不做循环"
            if auto_fix:
                fk, fp = auto_fix_note(cur_note, Path(payload).read_text(
                    encoding="utf-8"), transcript)
                result["fixed_path"] = fp
            record_metric("review", auto_fix=auto_fix, score=result["final_score"],
                          rounds=result["rounds"])
            return result
        report = Path(payload).read_text(encoding="utf-8")
        score = parse_review_score(report)
        result["final_score"] = score
        if not auto_fix:
            result["stopped_reason"] = "未启用 auto_fix"
            record_metric("review", auto_fix=auto_fix, score=result["final_score"],
                          rounds=result["rounds"])
            return result
        if score is not None and score >= target_score:
            result["stopped_reason"] = f"达标 {score}>={target_score}"
            record_metric("review", auto_fix=auto_fix, score=result["final_score"],
                          rounds=result["rounds"])
            return result
        if r >= max_rounds:
            result["stopped_reason"] = f"达轮次上限 {max_rounds}（末轮评分 {score}）"
            record_metric("review", auto_fix=auto_fix, score=result["final_score"],
                          rounds=result["rounds"])
            return result
        # 回修并复检
        fk, fp = auto_fix_note(cur_note, report, transcript)
        result["fixed_path"] = fp
        if fk != "note":
            result["stopped_reason"] = "native auto-fix 简报已出，交 agent 修复后复检"
            record_metric("review", auto_fix=auto_fix, score=result["final_score"],
                          rounds=result["rounds"])
            return result
        cur_note = fp  # 用修复稿进入下一轮复检
    record_metric("review", auto_fix=auto_fix, score=result["final_score"],
                  rounds=result["rounds"])
    return result


# --------------------------------------------------------------------------
# screenshot embedding (extract real frames into the note)
# --------------------------------------------------------------------------
_SCREENSHOT_RE = re.compile(r"\*Screenshot-\[(\d{1,2}:\d{2}(?::\d{2})?)\]\*?")


def _ts_to_seconds(ts: str) -> float:
    parts = [int(x) for x in ts.split(":")]
    if len(parts) == 3:
        return parts[0] * 3600 + parts[1] * 60 + parts[2]
    if len(parts) == 2:
        return parts[0] * 60 + parts[1]
    return float(parts[0])


def _extract_frame(video_path: str, ts: str, out_jpg: str) -> None:
    ffmpeg = _require("ffmpeg")
    secs = _ts_to_seconds(ts)
    subprocess.run(
        [ffmpeg, "-y", "-ss", f"{secs:.3f}", "-i", video_path,
         "-frames:v", "1", "-q:v", "2", out_jpg],
        check=True, capture_output=True, text=True,
    )


def embed_screenshots(note_path: str, video_path: str,
                      out_path: str = None) -> list[str]:
    """Find every `*Screenshot-[mm:ss]*` marker in a Markdown note, extract a
    frame from `video_path` at that timestamp, and replace the marker with a
    Markdown image reference. Images are written next to the note under
    `screenshots/`. Returns the list of extracted image paths (empty if none)."""
    note = Path(note_path)
    out = Path(out_path) if out_path else note
    text = note.read_text(encoding="utf-8")
    if not _SCREENSHOT_RE.search(text):
        log("No *Screenshot-[mm:ss]* markers found; nothing to embed.")
        return []
    img_dir = out.parent / "screenshots"
    img_dir.mkdir(parents=True, exist_ok=True)
    seen: dict[str, Path] = {}

    def repl(m):
        ts = m.group(1)
        if ts not in seen:
            idx = len(seen) + 1
            jpg = img_dir / f"{out.stem}_{idx:02d}.jpg"
            try:
                _extract_frame(video_path, ts, str(jpg))
                seen[ts] = jpg
            except Exception as e:  # keep marker if extraction fails
                log(f"Screenshot at {ts} failed: {e}")
                return m.group(0)
        rel = os.path.relpath(str(seen[ts]), str(out.parent)).replace("\\", "/")
        return f"![screenshot at {ts}]({rel})"

    new_text = _SCREENSHOT_RE.sub(repl, text)
    out.write_text(new_text, encoding="utf-8")
    paths = [str(p) for p in seen.values()]
    log(f"Embedded {len(paths)} screenshot(s) into {out}")
    return paths


# --------------------------------------------------------------------------
# multimodal video understanding: sample frames as visual input for the LLM
# --------------------------------------------------------------------------
def _get_video_duration(video_path: str) -> float:
    """Get video duration in seconds via ffprobe (returns 0.0 on failure)."""
    ffprobe = shutil.which("ffprobe")
    if not ffprobe:
        return 0.0
    try:
        r = subprocess.run(
            [ffprobe, "-v", "quiet", "-show_entries", "format=duration",
             "-of", "csv=p=0", video_path],
            check=True, capture_output=True, text=True)
        return float(r.stdout.strip())
    except Exception as e:
        log(f"Could not get video duration: {e}")
        return 0.0


def sample_frames(video_path: str, n_frames: int, work_dir: Path) -> list:
    """Extract n_frames evenly-spaced frames from a video for multimodal
    understanding. Returns list of absolute frame paths (jpg). Frames are
    named frame_01.jpg, frame_02.jpg... Skips silently on failure."""
    duration = _get_video_duration(video_path)
    if duration <= 0:
        log("Cannot sample frames: unknown duration.")
        return []
    work_dir.mkdir(parents=True, exist_ok=True)
    interval = duration / (n_frames + 1)
    frames = []
    for i in range(n_frames):
        ts = interval * (i + 1)
        mm, ss = divmod(int(ts), 60)
        hh, mm = divmod(mm, 60)
        ts_str = f"{hh:02d}:{mm:02d}:{ss:02d}"
        jpg = work_dir / f"frame_{i+1:02d}.jpg"
        try:
            _extract_frame(video_path, ts_str, str(jpg))
            frames.append(str(jpg))
        except Exception as e:
            log(f"Frame at {ts_str} failed: {e}")
    log(f"Sampled {len(frames)} frame(s) from {video_path}")
    return frames


_IMG_RE = re.compile(r"\*Image-(\d+)\*")


def _download_image(url: str, out_path: Path, timeout: int = 20) -> None:
    import requests
    from urllib.parse import urlparse
    headers = {"User-Agent": "Mozilla/5.0 (compatible; BiliNoteSkill/1.0)"}
    try:
        origin = "{0.scheme}://{0.netloc}".format(urlparse(url))
        headers["Referer"] = origin  # bypass simple hotlink protection
    except Exception:
        pass
    r = requests.get(url, headers=headers, timeout=timeout)
    r.raise_for_status()
    out_path.write_bytes(r.content)


def embed_article_images(note_path: str, images: list, out_path: str = None) -> list[str]:
    """Find every `*Image-N*` marker in a Markdown note, download the N-th
    image URL from `images` and replace the marker with a Markdown image
    reference. Images are written next to the note under `images/`. Returns
    the list of embedded image paths (empty if none / all failed)."""
    note = Path(note_path)
    out = Path(out_path) if out_path else note
    text = note.read_text(encoding="utf-8")
    if not _IMG_RE.search(text):
        log("No *Image-N* markers found; nothing to embed.")
        return []
    img_dir = out.parent / "images"
    img_dir.mkdir(parents=True, exist_ok=True)
    used: dict[int, Path] = {}
    n = len(images)

    def repl(m):
        idx = int(m.group(1))
        if idx < 1 or idx > n:
            return m.group(0)
        if idx not in used:
            jpg = img_dir / f"img_{idx:02d}.jpg"
            try:
                _download_image(images[idx - 1], jpg)
                used[idx] = jpg
            except Exception as e:  # keep marker if download fails
                log(f"Image {idx} download failed: {e}")
                return m.group(0)
        rel = os.path.relpath(str(used[idx]), str(out.parent)).replace("\\", "/")
        return f"![图{idx}]({rel})"

    new_text = _IMG_RE.sub(repl, text)
    out.write_text(new_text, encoding="utf-8")
    paths = [str(used[i]) for i in sorted(used)]
    log(f"Embedded {len(paths)} article image(s) into {out}")
    return paths


# --------------------------------------------------------------------------
# export (Markdown note -> PDF / DOCX / Mermaid mindmap)
# --------------------------------------------------------------------------
_MD_IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def generate_mindmap(md_text: str) -> str:
    """Generate a Mermaid mindmap from a Markdown note's heading structure."""
    lines = md_text.splitlines()
    root = "笔记"
    nodes = []
    stack = []
    heading_re = re.compile(r"^(#{1,6})\s+(.+?)\s*\**\s*$")
    for line in lines:
        m = heading_re.match(line)
        if not m:
            continue
        level = len(m.group(1))
        text = m.group(2).strip()
        text = re.sub(r"\*Content-\[(\d{2}:\d{2})\]*", r" [\1]", text)
        text = re.sub(r"\*Screenshot-\[(\d{2}:\d{2})\]*", "", text).strip()
        if not text:
            continue
        if not stack:
            root = text
            stack = [(0, text)]
            nodes.append(f'  {root}')
            continue
        stack = stack[:level]
        indent = "  " * (level + 1)
        nodes.append(f'{indent}{text}')
        stack.append((level, text))
    if not nodes:
        return "# 思维导图 (mindmap)\n\n_(笔记中没有标题结构，无法生成思维导图)_\n"
    body = "\n".join(nodes)
    return (
        "# Mermaid 思维导图 (mindmap)\n\n"
        "可从 Markdown 标题自动提取。在支持 Mermaid 的阅读器中打开即可查看 "
        "(如 GitHub / Notion / Obsidian / https://mermaid.live)：\n\n"
        "```mermaid\nmindmap\n" + body + "\n```\n")


def _slice_tall_image(src: str, work_dir: Path, max_ratio: float = 1.4) -> list:
    """Split a very tall image (e.g. an infographic long-strip) into
    page-friendly vertical chunks so PDF/DOCX renderers don't crush or clip
    it. Returns absolute paths; a single-element list means no slicing."""
    from PIL import Image
    im = Image.open(src)
    w, h = im.size
    if w <= 0 or h <= w * max_ratio:
        return [src]
    work_dir.mkdir(parents=True, exist_ok=True)
    chunk_h = max(1, int(w * max_ratio))
    im = im.convert("RGB")
    stem = Path(src).stem
    parts, y, idx = [], 0, 0
    while y < h:
        box = (0, y, w, min(y + chunk_h, h))
        p = work_dir / f"{stem}_p{idx:02d}.jpg"
        im.crop(box).save(p, "JPEG", quality=90)
        parts.append(str(p))
        y += chunk_h
        idx += 1
    return parts


def _prepare_export_images(md_text: str, base_dir: Path, work_dir: Path) -> str:
    """Rewrite Markdown image refs for export: resolve local paths relative to
    the note directory, and slice over-tall images into page-sized chunks.
    Remote (http/data) images are left untouched. Chunk files land in
    `work_dir` and are referenced relative to `base_dir`."""
    base_dir = Path(base_dir)

    def repl(m):
        alt, path = m.group(1), m.group(2).strip()
        if path.startswith(("http://", "https://", "data:", "//")):
            return m.group(0)
        src = (base_dir / path)
        if not src.exists():
            return m.group(0)
        try:
            parts = _slice_tall_image(str(src.resolve()), work_dir)
        except Exception as e:
            log(f"Image prepare skipped for {path}: {e}")
            return m.group(0)
        if len(parts) == 1:
            return m.group(0)  # keep original ref (already valid)
        refs = []
        for i, p in enumerate(parts):
            rel = os.path.relpath(p, str(base_dir)).replace("\\", "/")
            refs.append(f"![{alt} ({i+1}/{len(parts)})]({rel})")
        return "\n\n".join(refs)

    return _MD_IMG_RE.sub(repl, md_text)


def _add_inline(paragraph, text: str) -> None:
    """Render **bold**, *italic*, _italic_, `code` as styled runs."""
    pat = re.compile(r"(\*\*([^*]+)\*\*|\*([^*]+)\*|_([^_]+)_|`([^`]+)`)")
    pos = 0
    for m in pat.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(2) is not None:
            r = paragraph.add_run(m.group(2)); r.bold = True
        elif m.group(3) is not None:
            r = paragraph.add_run(m.group(3)); r.italic = True
        elif m.group(4) is not None:
            r = paragraph.add_run(m.group(4)); r.italic = True
        elif m.group(5) is not None:
            r = paragraph.add_run(m.group(5)); r.font.name = "Courier New"
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx(md_text: str, out_path: str, base_dir: str = ".") -> None:
    """Convert Markdown to a .docx file (headings, lists, bold/italic/code,
    images). Image paths are resolved relative to `base_dir`."""
    from docx import Document
    from docx.shared import Emu
    doc = Document()
    # usable page width (page width minus L/R margins) for image scaling
    sec = doc.sections[0]
    avail_w = sec.page_width - sec.left_margin - sec.right_margin
    base = Path(base_dir)
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        im = _MD_IMG_RE.fullmatch(stripped)
        if im:
            path = im.group(2).strip()
            if not path.startswith(("http://", "https://", "data:", "//")):
                src = (base / path)
                if src.exists():
                    try:
                        from PIL import Image as _PILImage
                        w, h = _PILImage.open(str(src)).size
                        pic_w = avail_w
                        # cap height so one picture never exceeds ~a page
                        if w > 0 and (h / w) * avail_w > Emu(int(9.0 * 914400)):
                            pic_w = int(Emu(int(9.0 * 914400)) * w / h)
                        doc.add_picture(str(src.resolve()), width=Emu(int(pic_w)))
                    except Exception as e:
                        log(f"DOCX image skipped ({path}): {e}")
                        p = doc.add_paragraph()
                        _add_inline(p, im.group(1) or path)
            i += 1
            continue
        if stripped.startswith("```"):
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(buf))
            run.font.name = "Courier New"
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            h = doc.add_heading(level=min(len(m.group(1)), 4))
            _add_inline(h, m.group(2).strip())
            i += 1
            continue
        if re.match(r"^[-*_]{3,}\s*$", stripped):
            doc.add_paragraph("─" * 24)
            i += 1
            continue
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, stripped[2:].strip())
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue
        p = doc.add_paragraph()
        _add_inline(p, stripped)
        i += 1
    doc.save(out_path)


def _ensure_h1(md_text: str) -> str:
    """markdown_pdf requires the FIRST heading to be level 1. If the note has
    no top-level (`#`) heading, promote the first `##` to `#`."""
    if re.search(r"(?m)^#\s", md_text):
        return md_text
    return re.sub(r"(?m)^##\s", "# ", md_text, count=1)


def markdown_to_pdf(md_text: str, out_path: str, root: str = ".") -> None:
    """Convert Markdown to a .pdf file (PyMuPDF-based, no LaTeX needed).
    `root` is the base directory used to resolve relative image paths."""
    from markdown_pdf import MarkdownPdf, Section
    md_text = _ensure_h1(md_text)
    pdf = MarkdownPdf()
    pdf.add_section(Section(md_text, root=str(root)))
    pdf.save(out_path)


def export(md_path: str, out_base: Path, formats) -> None:
    """Export a Markdown note to one or more formats (pdf, docx, mindmap).
    Local images are resolved relative to the note's directory and over-tall
    images are sliced so they render fully (not crushed/clipped).
    `mindmap` format emits a Mermaid mindmap (.mindmap.md)."""
    import shutil
    md_path = Path(md_path)
    base_dir = md_path.parent if str(md_path.parent) else Path(".")
    md_text = md_path.read_text(encoding="utf-8")
    out_base = Path(out_base)
    tmp_dir = base_dir / "_export_img"
    prepared = _prepare_export_images(md_text, base_dir, tmp_dir)
    try:
        for fmt in formats:
            fmt = str(fmt).strip().lower()
            if fmt == "pdf":
                out = str(out_base.with_suffix(".pdf"))
                log(f"Exporting PDF -> {out}")
                markdown_to_pdf(prepared, out, root=str(base_dir))
            elif fmt == "docx":
                out = str(out_base.with_suffix(".docx"))
                log(f"Exporting DOCX -> {out}")
                markdown_to_docx(prepared, out, base_dir=str(base_dir))
            elif fmt == "mindmap":
                out = out_base.with_suffix(".mindmap.md")
                log(f"Exporting Mermaid mindmap -> {out}")
                out.write_text(generate_mindmap(md_text), encoding="utf-8")
            else:
                log(f"Unknown export format skipped: {fmt}")
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# --------------------------------------------------------------------------
# run (orchestration)
# --------------------------------------------------------------------------
def _finalize(out: Path, kind: str, payload: str, embed_flag: bool,
              persisted_video, export_formats, images=None,
              include_images: bool = False) -> str:
    """Write / export the result of summarize() and return the note path or
    (in native mode) the brief path."""
    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)
    if kind == "note":
        out.write_text(payload, encoding="utf-8")
        log(f"Note written to {out}")
        if embed_flag and persisted_video:
            try:
                embed_screenshots(str(out), persisted_video)
            except Exception as e:
                log(f"Screenshot embedding failed: {e}")
        if include_images and images:
            try:
                embed_article_images(str(out), images)
            except Exception as e:
                log(f"Article image embedding failed: {e}")
        if export_formats:
            export(str(out), out.with_suffix(""), export_formats)
        record_metric("run", words=len(payload.split()),
                       exports=",".join(export_formats or []))
        return str(out)
    # kind == "brief": agent's own LLM will finalize the note.
    if export_formats:
        log(f"原生模式：笔记将由 agent 生成，稍后对其运行 export 即可（"
            f"bilinote.py export {out} --formats {export_formats}）。")
    return payload  # path to the brief


def run(source: str, out: Path, local: bool = None, lang: str = None,
        transcriber: str = "auto",
        time_markers: bool = True, screenshots: bool = False,
        ai_summary: bool = True,
        export_formats: list = None, style: str = DEFAULT_STYLE,
        embed_screenshots_flag: bool = False,
        include_images: bool = False,
        multimodal: bool = False,
        multimodal_frames: int = 6,
        subtitle_file: str = None,
        bili_ai: bool = True,
        danmaku: bool = False,
        comments: bool = False,
        comments_in_report: bool = False,
        bili_cookie: str = "",
        linkify: bool = False,
        clean: bool = True,
        translate: bool = False,
        review: bool = False,
        video_quality: str = "720",
        outline_first: bool = False) -> str:
    # Embedding requires screenshot markers, which require --screenshots.
    screenshots = screenshots or embed_screenshots_flag
    bili_cookie = bili_cookie or os.environ.get("BILINOTE_BILI_COOKIE", "")
    cls = classify_source(source)
    kind = cls["kind"]
    out = Path(out)
    if kind == "unknown":
        log_recovery("F8", "无法识别输入类型，按纯文本处理")
        kind = "text"
        cls["kind"] = "text"
    log(f"输入识别：{kind} -> 路由到对应流程")

    # ---------- subtitle file (+ optional video) ----------
    if kind == "subtitle_file" or (subtitle_file and Path(subtitle_file).exists()):
        sub_path = Path(source) if kind == "subtitle_file" else Path(subtitle_file)
        log(f"检测到字幕文件：{sub_path}")
        transcript = parse_vtt(sub_path)
        if not transcript.strip():
            raise RuntimeError("字幕文件内容为空，无法生成笔记。")
        title = sub_path.stem
        tags = ""
        log(f"Title (from subtitle filename): {title}")

        # If a video source was also provided (alongside the subtitle),
        # extract its metadata for multimodal / screenshot embedding.
        persisted_video = None
        mm_frames = []
        if kind in ("video_url", "media_url", "media_file"):
            is_local = bool(cls.get("local", False)) if local is None else local
            with tempfile.TemporaryDirectory() as tmp:
                work = Path(tmp)
                if is_local:
                    meta = fetch_local(source, work, with_video=screenshots or multimodal)
                else:
                    meta = fetch(source, work, lang=lang or "zh",
                                  with_video=screenshots or multimodal)
                title = meta["title"]
                tags = meta["tags"]
                log(f"Video metadata merged: {title}")
                if multimodal and meta.get("video"):
                    out.parent.mkdir(parents=True, exist_ok=True)
                    persisted_video = str(out.parent / (out.stem + ".video.mp4"))
                    try:
                        shutil.copy(meta["video"], persisted_video)
                        mm_frames = sample_frames(persisted_video, multimodal_frames,
                                                  out.parent / "multimodal_frames")
                    except Exception as e:
                        log(f"Could not persist video / sample frames: {e}")
                elif screenshots and meta.get("video"):
                    out.parent.mkdir(parents=True, exist_ok=True)
                    persisted_video = str(out.parent / (out.stem + ".video.mp4"))
                    try:
                        shutil.copy(meta["video"], persisted_video)
                    except Exception as e:
                        log(f"Could not persist video: {e}")

        kind_r, payload = summarize(transcript, title, tags,
                                    True, screenshots, ai_summary,
                                    out=str(out),
                                    style=style, video_path=persisted_video,
                                    media=True,
                                    multimodal_frames=mm_frames if mm_frames else None,
                                    outline_first=outline_first)
        return _finalize(out, kind_r, payload, embed_screenshots_flag,
                         persisted_video, export_formats)

    # ---------- media (video / audio) ----------
    if kind in ("video_url", "media_url", "media_file"):
        is_local = bool(cls.get("local", False)) if local is None else local
        with_video = screenshots or multimodal
        persisted_video = None
        is_bili = (not is_local) and _is_bilibili(source)

        # ---- gather borrowed enrichers (bilibili-centric) ----
        enrich = {"video_url": "" if is_local else source,
                  "linkify": linkify, "translate": translate}
        ai = None
        bili_cid = None
        bili_aid = None
        bili_duration = None
        if is_bili and bili_ai:
            log("尝试调用 B 站官方 AI 总结（conclusion/get，#0 最高优先）...")
            ai = fetch_bilibili_ai_summary(source, bili_cookie)
            if ai.get("ok"):
                bili_cid = ai.get("cid")
                bili_aid = ai.get("aid")
                bili_duration = ai.get("duration")
                enrich["ai_summary_md"] = bili_ai_summary_to_markdown(ai)
                log(f"命中官方 AI 总结：{len(ai.get('outline', []))} 段时间轴大纲、"
                    f"{len(ai.get('subtitle', []))} 条 AI 字幕。")
            else:
                log(f"官方 AI 总结不可用：{ai.get('reason')}（回退字幕优先/转写）。"
                    " 提示：设置环境变量 BILINOTE_BILI_COOKIE 或用 --bili-cookie "
                    "传入含 SESSDATA 的 Cookie 可显著提高命中率。")
                ai = None

        # ---- danmaku: context (#1) + content-cleaning sponsor detection (#4) ----
        danmaku_list = []
        clean_ranges = []
        if is_bili and (danmaku or clean):
            if not bili_cid:
                try:
                    _v = _bili_get_view(_bili_extract_id(source), bili_cookie)
                    bili_cid = _v.get("cid")
                    bili_aid = _v.get("aid")
                    bili_duration = bili_duration or _v.get("duration")
                except Exception:
                    bili_cid = None
            if bili_cid:
                danmaku_list = fetch_bilibili_danmaku(bili_cid, bili_cookie,
                                                      duration=bili_duration)
        if danmaku_list:
            if danmaku:
                enrich["danmaku_ctx"] = danmaku_as_context(danmaku_list)
                log(f"弹幕补充：抓取 {len(danmaku_list)} 条弹幕作为二级上下文。")
            if clean:
                clean_ranges = detect_sponsor_segments_from_danmaku(danmaku_list)
                if clean_ranges:
                    log(f"内容净化：弹幕信号检出 {len(clean_ranges)} 段疑似赞助/广告"
                        f"（{', '.join(fmt_time(s)+'–'+fmt_time(e) for s, e in clean_ranges)}）。")
        # 通用内容净化指令：片头口播 / 片尾 CTA / 赞助广告口播由 LLM 内容自判断跳过；
        # 若弹幕检出疑似赞助区间则一并作为参考（不依赖外网 SponsorBlock）。
        if clean:
            enrich["clean_instruction"] = build_clean_instruction(clean_ranges)

        # ---- comment-area highlights (#C): fetch + select, opt-in to report ----
        if is_bili and comments:
            if not bili_aid:
                try:
                    _v = _bili_get_view(_bili_extract_id(source), bili_cookie)
                    bili_aid = _v.get("aid")
                    bili_cid = bili_cid or _v.get("cid")
                    bili_duration = bili_duration or _v.get("duration")
                except Exception:
                    bili_aid = None
            if bili_aid:
                raw_cm = fetch_bilibili_comments(bili_aid, bili_cookie,
                                                 limit=60, sort="hot")
                if raw_cm:
                    hi = select_highlight_comments(raw_cm, top_n=8,
                                                   min_likes=10)
                    if hi:
                        enrich["comments_ctx"] = comments_as_context(hi)
                        enrich["comments_highlights"] = hi
                        enrich["comments_in_report"] = bool(comments_in_report)
                        log(f"评论区精选：抓取 {len(raw_cm)} 条评论，精选 "
                            f"{len(hi)} 条高能评论"
                            + ("（将加入报告正文）。" if comments_in_report
                               else "（仅作上下文参考）。"))

        # ---- transcript selection: official AI subtitle takes priority ----
        transcript = None
        meta_title = ai.get("title") if (ai and ai.get("title")) else None
        meta_tags = ""
        if ai and ai.get("subtitle"):
            transcript = bili_ai_summary_to_transcript(ai)
            log("使用官方 AI 字幕作为转写来源（免下载音频/ASR）。")

        need_download = with_video or (transcript is None)
        mm_frames = []
        if need_download:
            with tempfile.TemporaryDirectory() as tmp:
                work = Path(tmp)
                if is_local:
                    meta = fetch_local(source, work, with_video=with_video)
                else:
                    meta = fetch(source, work, lang=lang or "zh",
                                  with_video=with_video,
                                  video_quality=video_quality)
                meta_title = meta_title or meta["title"]
                meta_tags = meta["tags"]
                bili_duration = bili_duration or meta.get("duration")

                def _pick():
                    if transcriber == "auto":
                        return route_transcribe(meta["audio"], lang=lang)
                    return dispatch_transcribe(meta["audio"], transcriber)

                if transcript is None:
                    if meta["used_subtitle"]:
                        subtitle = parse_vtt(Path(meta["subtitle"]))
                        if not subtitle.strip():
                            log_recovery("F2", "字幕为空，回退音频转写")
                            transcript = _pick()
                        else:
                            transcript = subtitle
                    else:
                        transcript = _pick()

                if multimodal and meta.get("video"):
                    out.parent.mkdir(parents=True, exist_ok=True)
                    persisted_video = str(out.parent / (out.stem + ".video.mp4"))
                    try:
                        shutil.copy(meta["video"], persisted_video)
                        log(f"Video persisted for multimodal: {persisted_video}")
                        mm_frames = sample_frames(persisted_video, multimodal_frames,
                                                  out.parent / "multimodal_frames")
                    except Exception as e:
                        log(f"Could not persist video / sample frames: {e}")
                        persisted_video = meta["video"]
                elif with_video and meta.get("video"):
                    out.parent.mkdir(parents=True, exist_ok=True)
                    persisted_video = str(out.parent / (out.stem + ".video.mp4"))
                    try:
                        shutil.copy(meta["video"], persisted_video)
                        log(f"Video persisted for embedding: {persisted_video}")
                    except Exception as e:
                        log(f"Could not persist video: {e}")
                        persisted_video = meta["video"]

        # ---- high-energy / chapter-driven screenshot suggestions (#2) ----
        if screenshots:
            ts = suggest_screenshot_timestamps(
                outline=ai.get("outline") if ai else None,
                danmaku=danmaku_list or None,
                duration=bili_duration,
                n=multimodal_frames if multimodal else 6)
            if ts:
                enrich["suggest_ts"] = ts
                log(f"已推荐 {len(ts)} 个截图时间点（大纲/弹幕高能优先）。")

        kind_r, payload = summarize(transcript, meta_title or "视频笔记", meta_tags,
                                    time_markers, screenshots, ai_summary,
                                    out=str(out),
                                    style=style, video_path=persisted_video,
                                    media=True,
                                    multimodal_frames=mm_frames if mm_frames else None,
                                    enrich=enrich,
                                    outline_first=outline_first)
        note_path = _finalize(out, kind_r, payload, embed_screenshots_flag,
                              persisted_video, export_formats)
        # #3 make timestamps clickable for note-mode output right away;
        # brief-mode notes are linkified by the agent per the brief rule.
        if linkify and kind_r == "note" and not is_local:
            try:
                linkify_note(note_path, source)
            except Exception as e:
                log(f"linkify 失败：{e}")
        if review and kind_r == "note":
            _maybe_review(note_path, transcript)
        return note_path

    # ---------- text / article / document ----------
    log("非媒体输入：跳过语音识别，直接整理为笔记。")
    article_images = None
    if kind == "article_url":
        title, text, article_images = scrape_article(source)
    elif kind == "doc_file":
        title, text = read_document(source)
    else:  # raw text
        text = source
        title = (source[:50].replace("\n", " ").strip()) or "直接输入的文本"
    if not text or not text.strip():
        raise RuntimeError("未能读取到任何正文内容，无法生成笔记。")
    kind_r, payload = summarize(text, title, tags="", time_markers=False,
                                screenshots=False, ai_summary=ai_summary,
                                out=str(out), style=style, video_path=None,
                                media=False, images=article_images,
                                include_images=include_images,
                                outline_first=outline_first)
    note_path = _finalize(out, kind_r, payload, False, None, export_formats,
                          images=article_images, include_images=include_images)
    if review and kind_r == "note":
        _maybe_review(note_path, text)
    return note_path


# --------------------------------------------------------------------------
# Plan B: run --wizard 交互起飞前清单（绕过 AskUserQuestion 4 选项上限）
# --------------------------------------------------------------------------
def _resolve_wizard_style(inp: str) -> str:
    """把用户输入（编号或键名）解析为合法风格键；非法/空回退默认。"""
    if not inp:
        return DEFAULT_STYLE
    if inp.isdigit() and 1 <= int(inp) <= len(STYLES):
        return list(STYLES.keys())[int(inp) - 1]
    if inp in STYLES:
        return inp
    return DEFAULT_STYLE


def _build_wizard_argv(url, style, danmaku, screenshots, translate,
                       illu_mode, vq, export, out,
                       comments=False, comments_in_report=False):
    """将起飞前清单的回答拼成可复用 run 管线的 argv（不含 --wizard，避免递归）。"""
    argv = ["run", url, "--style", style, "--video-quality", vq, "--out", out]
    if danmaku:
        argv.append("--danmaku")
    if screenshots:
        argv.append("--screenshots")
    if translate:
        argv.append("--translate")
    if comments:
        argv.append("--comments")
        if comments_in_report:
            argv.append("--comments-in-report")
    if export:
        argv += ["--export", export]
    return argv


def run_wizard(url: str, out_default: str = "note.md") -> None:
    """交互式起飞前清单：逐一确认风格/富化开关/画质/导出/手绘模式，
    再复用 run 管线（及可选 illustrate 后处理）。用 input() 读取，非 tty
    时可管道喂入；EOF 回退默认。绕开 AskUserQuestion 每题最多 4 选项的限制，
    把全部 11 风格与 3 个富化开关（弹幕/抽帧/翻译）一次性铺开。"""
    def ask(prompt, default=""):
        try:
            v = input(prompt).strip()
        except EOFError:
            return default
        return v or default

    def yn(prompt):
        return ask(f"{prompt} [y/N] ", "n").lower() in ("y", "yes", "是", "1")

    print("\n=== BiliNote 起飞前清单 (wizard) ===")
    print("可选风格（输入编号或键名）：")
    for i, (k, (label, _)) in enumerate(STYLES.items(), 1):
        print(f"  {i:2}. {k:12} {label}")
    style_in = ask("选择风格 > ", DEFAULT_STYLE)
    style = _resolve_wizard_style(style_in)
    if style != style_in.strip():
        print(f"  → 已锁定风格：{style}")

    danmaku = yn("弹幕补转写 (--danmaku，抓弹幕作第二字幕，绝不抄进正文)?")
    screenshots = yn("智能抽帧截图 (--screenshots，按 大纲/弹幕密度 出 *Screenshot-[mm:ss]* 建议)?")
    translate = yn("外语原声翻译 (--translate)?")
    comments = yn("评论区高能精选 (--comments，抓 B 站高赞评论作观众态度参考)?")
    comments_in_report = False
    if comments:
        comments_in_report = yn("把精选评论加入报告正文 (--comments-in-report，生成 ## 评论区高能精选 节)?")
    illu = ask("手绘配图: 0=不需要 / 1=小黑手绘(illustration) / 2=一页图解(infographic) > ", "0")
    illu_mode = {"1": "illustration", "2": "infographic"}.get(illu, None)
    vq = ask("视频画质 720/1080/best > ", "720")
    if vq not in ("720", "1080", "best"):
        vq = "720"
    export = ask("导出格式 (md/pdf/docx/mindmap，逗号分隔，留空=仅 md) > ", "")
    out = ask("输出笔记路径 > ", out_default)

    argv = _build_wizard_argv(url, style, danmaku, screenshots, translate,
                              illu_mode, vq, export, out,
                              comments=comments,
                              comments_in_report=comments_in_report)
    print(f"\n>>> 执行: bilinote.py {' '.join(argv)}")
    main(argv)

    if illu_mode:
        print(f"\n>>> 追加手绘配图: bilinote.py illustrate {out} --mode {illu_mode}")
        main(["illustrate", out, "--mode", illu_mode])


# --------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------
def main(argv=None):
    p = argparse.ArgumentParser(
        prog="bilinote", description="Generate AI video notes (BiliNote-style, no RAG).")
    sub = p.add_subparsers(dest="cmd", required=True)

    pf = sub.add_parser("fetch", help="Download audio + subtitles")
    pf.add_argument("url")
    pf.add_argument("--workdir", default="./bilinote_work")
    pf.add_argument("--lang", default="zh")

    pt = sub.add_parser("transcribe", help="Audio -> transcript")
    pt.add_argument("audio")
    pt.add_argument("--backend", default=os.environ.get(
        "BILINOTE_TRANSCRIBER", "bcut"),
        help="bcut/必剪 (默认, 免费无密钥), kuaishou/快手 (免费无密钥), "
             "auto (按语种路由)")
    pt.add_argument("--out", default="./transcript.txt")

    ps = sub.add_parser("summarize", help="Transcript -> Markdown note")
    ps.add_argument("--transcript", required=True)
    ps.add_argument("--title", default="unknown")
    ps.add_argument("--tags", default="")
    ps.add_argument("--no-time-markers", action="store_true")
    ps.add_argument("--screenshots", action="store_true")
    ps.add_argument("--no-ai-summary", action="store_true")
    ps.add_argument("--style", default=os.environ.get(
        "BILINOTE_STYLE", DEFAULT_STYLE),
        help=STYLE_HELP + "；也可用 @name 引用自定义风格（见 references/user_styles/）")
    ps.add_argument("--export", default="",
                    help="Export alongside the note: comma-separated pdf,docx,mindmap")
    ps.add_argument("--outline-first", action="store_true",
                    help="Plan J：先生成论点树大纲供确认，再据大纲填正文。")
    ps.add_argument("--out", default="./note.md")

    pr = sub.add_parser("run", help="统一入口：自动识别输入类型并路由")
    pr.add_argument("url", help="输入：视频/文章 URL、本地视频/文档路径，"
                        "或直接粘贴的文本（自动识别类型后进入对应流程）")
    pr.add_argument("--local", action="store_true",
                    help="（通常无需指定）强制按本地文件处理；默认由输入自动识别")
    pr.add_argument("--lang", default=None,
                    help="视频语种 (ISO-639-1, 如 zh/en/ja/ko)。留空则自动侦测语种后再选引擎。")
    pr.add_argument("--transcriber", default=os.environ.get(
        "BILINOTE_TRANSCRIBER", "auto"),
        help="auto (默认): 中文/英文先bcut、其他语种先kuaishou，失败互相反向兜底；"
             "或显式指定 bcut/必剪、kuaishou/快手。本地 fast-whisper/groq 已移除。")
    pr.add_argument("--no-time-markers", action="store_true")
    pr.add_argument("--screenshots", action="store_true")
    pr.add_argument("--video-quality", default="720",
                    choices=["720", "1080", "best"],
                    help="截图/多模态用的视频画质：720(默认,够清晰且省体积) | 1080 "
                         "| best。下 1080p/best 需登录态（设 BILINOTE_YTDLP_COOKIES "
                         "从浏览器读，或传 cookie 文件），否则 B站只返回低清。")
    pr.add_argument("--embed-screenshots", action="store_true",
                    help="把笔记中 *Screenshot-[mm:ss]* 标记处真抽取视频帧并嵌入"
                         "（自动开启 --screenshots；需 ffmpeg，并会下载视频）。")
    pr.add_argument("--no-ai-summary", action="store_true")
    pr.add_argument("--style", default=os.environ.get(
        "BILINOTE_STYLE", DEFAULT_STYLE),
        help=STYLE_HELP + "；也可用 @name 引用自定义风格（见 references/user_styles/）")
    pr.add_argument("--wizard", action="store_true",
                    help="Plan B 交互起飞前清单：逐一确认全部 11 风格 + 弹幕/抽帧/翻译/"
                         "手绘开关 + 画质 + 导出，绕过 AskUserQuestion 4 选项上限，再调用 run 管线。")
    pr.add_argument("--images", action="store_true",
                    help="（文章/网页输入）下载并嵌入原文配图：在笔记中用 "
                         "*Image-N* 标记引用对应图片，生成后运行 "
                         "embed-article-images 下载并替换为真实图片。")
    pr.add_argument("--multimodal", action="store_true",
                    help="多模态视频理解：均匀抽样视频帧作为视觉输入与转录内容"
                         "一起喂给大模型，捕捉画面中的文字/图表/PPT/操作演示"
                         "（v2.3.0+；需要 --embed-screenshots 或先下载视频；"
                         "每帧会被 base64 编码进 prompt，注意 token 成本）。")
    pr.add_argument("--multimodal-frames", type=int, default=6,
                    metavar="N",
                    help="多模态模式抽取的关键帧数（默认 6；每帧约 50–150KB base64，"
                         "帧越多视觉信息越丰富但 token 越多）。")
    pr.add_argument("--export", default="",
                    help="Export alongside the note: comma-separated pdf,docx,mindmap")
    pr.add_argument("--subtitle-file", dest="subtitle_file", default=None,
                    help="字幕文件路径 (.vtt/.srt/.ass)：与视频 URL 或本地视频同时传入，"
                         "跳直接字幕文件读取，跳过 fetch 字幕下载；"
                         "仅传字幕文件时自动进入纯字幕 summarize 流程。")
    pr.add_argument("--no-bili-ai", action="store_true",
                    help="关闭 B 站官方 AI 总结优先（#0）。默认对 B 站视频先尝试官方"
                         "AI 总结（summary + 时间轴大纲 + AI 字幕，免下载），失败自动"
                         "回退字幕优先/ASR 转写。")
    pr.add_argument("--danmaku", action="store_true",
                    help="抓取弹幕作为二级上下文（#1）喂给 LLM，辅助判断高能片段与"
                         "观众观点（仅 B 站；弹幕不会照抄进正文）。")
    pr.add_argument("--comments", action="store_true",
                    help="评论区高能精选（#C）：抓取 B 站高赞评论作观众态度参考"
                         "（仅 B 站）。配合 --comments-in-report 可加入报告正文。")
    pr.add_argument("--comments-in-report", action="store_true",
                    help="把精选高赞评论写入笔记正文（## 评论区高能精选 一节）；"
                         "需与 --comments 同时使用，否则无效。")
    pr.add_argument("--bili-cookie", default="",
                    help="B 站 Cookie（含 SESSDATA），用于官方 AI 总结等需登录接口；"
                         "留空则读环境变量 BILINOTE_BILI_COOKIE。")
    pr.add_argument("--linkify", action="store_true",
                    help="把笔记中的 *Content-[mm:ss]* 与裸 [mm:ss] 时间戳转成可点击"
                         "跳转到视频对应时间点的链接（#3；note 模式即时生效，"
                         "原生/brief 模式由 agent 生成后运行 linkify 子命令）。")
    pr.add_argument("--no-clean", action="store_true",
                    help="关闭内容净化（#4）：默认会剔除赞助/广告/片头片尾等非干货"
                         "内容；用此开关保留原文。")
    pr.add_argument("--translate", action="store_true",
                    help="外语视频原声翻译（#5）：提示 LLM 把要点翻译成简体中文"
                         "（专名保留英文），最终笔记以中文为主。")
    pr.add_argument("--review", action="store_true",
                    help="Plan C 质量闸：note 生成后自动跑一遍质检（源转录 vs 笔记），"
                         "输出 <名>.review.md / .review-brief.md。默认关闭。")
    pr.add_argument("--no-cache", action="store_true",
                    help="Plan G：关闭同源取数缓存，强制重新抓取/转写。")
    pr.add_argument("--outline-first", action="store_true",
                    help="Plan J：先生成论点树大纲供确认，再据大纲填正文（减少遗漏）。")
    pr.add_argument("--out", default="./note.md")

    pe = sub.add_parser("embed-screenshots",
                         help="把笔记里的 *Screenshot-[mm:ss]* 标记替换为真实视频帧")
    pe.add_argument("--note", required=True, help="含 *Screenshot-[mm:ss]* 标记的笔记 .md")
    pe.add_argument("--video", default="", help="本地视频文件路径")
    pe.add_argument("--url", default="", help="或提供视频 URL，脚本先下载视频")
    pe.add_argument("--out", default="",
                    help="输出笔记路径（默认覆盖 --note）")
    pe.add_argument("--workdir", default="./bilinote_work",
                    help="下载视频时的临时目录")
    pe.add_argument("--video-quality", default="720",
                    choices=["720", "1080", "best"],
                    help="仅 --url 下载视频时生效：720(默认) | 1080 | best。"
                         "1080/best 需登录态(BILINOTE_YTDLP_COOKIES)。")

    pi = sub.add_parser("embed-article-images",
                        help="把笔记里的 *Image-N* 标记替换为下载的原文图片")
    pi.add_argument("--note", required=True,
                    help="含 *Image-N* 标记的笔记 .md")
    pi.add_argument("--images", required=True,
                    help="图片 URL 列表：JSON 数组（如 '[\"url1\",\"url2\"]'）"
                         "或逗号分隔的 URL 串")
    pi.add_argument("--out", default="",
                    help="输出笔记路径（默认覆盖 --note）")

    px = sub.add_parser("export", help="Markdown note -> PDF/DOCX")
    px.add_argument("note", help="Source .md note")
    px.add_argument("--formats", default="pdf,docx",
                    help="Comma-separated: pdf,docx,mindmap")
    px.add_argument("--out", default=None,
                    help="Output base path (default: note without .md)")

    # -------- borrowed-feature standalone subcommands --------
    pbs = sub.add_parser("bili-summary",
                         help="#0 拉取 B 站官方 AI 总结（summary+时间轴大纲+AI字幕）")
    pbs.add_argument("url", help="B 站视频 URL 或 BV/av 号")
    pbs.add_argument("--cookie", default="",
                     help="含 SESSDATA 的 Cookie；留空读 BILINOTE_BILI_COOKIE")
    pbs.add_argument("--json", action="store_true", help="输出原始 JSON")
    pbs.add_argument("--out", default="",
                     help="把 summary+大纲渲染成 Markdown 写入该路径")
    pbs.add_argument("--transcript", default="",
                     help="把 AI 字幕导出为 [mm:ss] 转写文本到该路径")

    pdm = sub.add_parser("danmaku", help="#1 抓取弹幕作为二级上下文")
    pdm.add_argument("url", help="B 站视频 URL 或 BV/av 号（自动解析 cid）")
    pdm.add_argument("--cookie", default="")
    pdm.add_argument("--limit", type=int, default=0, help="最多导出条数（0=全部）")
    pdm.add_argument("--out", default="", help="导出为 [mm:ss] 上下文文本到该路径")

    pcm = sub.add_parser("comments", help="#C 评论区高能精选：抓取 B 站高赞评论")
    pcm.add_argument("url", help="B 站视频 URL 或 BV/av 号（自动解析 aid 作为 oid）")
    pcm.add_argument("--cookie", default="")
    pcm.add_argument("--limit", type=int, default=60, help="抓取评论条数上限")
    pcm.add_argument("--top", type=int, default=8, help="精选（按赞）输出条数")
    pcm.add_argument("--min-likes", type=int, default=10,
                     help="入选精选区的最低赞数")
    pcm.add_argument("--report", action="store_true",
                     help="输出为可直接粘贴进报告的 ## 评论区高能精选 段")
    pcm.add_argument("--out", default="", help="导出路径（默认打印到 stdout）")

    plk = sub.add_parser("linkify",
                         help="#3 把笔记里的时间戳转成可点击跳转链接")
    plk.add_argument("--note", required=True, help="含时间戳的笔记 .md")
    plk.add_argument("--url", required=True, help="视频 URL（用于拼接 ?t=秒）")
    plk.add_argument("--out", default="", help="输出路径（默认覆盖 --note）")

    pc = sub.add_parser("clean",
                         help="#4 去除转录/笔记中的赞助/广告/片头片尾段"
                              "（替代外网 SponsorBlock，原生适配 B 站）")
    pc.add_argument("input", help="转录 .txt 或笔记 .md 路径")
    pc.add_argument("--ranges", default="",
                    help="手动指定赞助区间（秒），如 30-45,120-140")
    pc.add_argument("--intro", type=int, default=0,
                    help="跳过片头前 N 秒（按时间戳物理去除）")
    pc.add_argument("--outro", type=int, default=0,
                    help="跳过片尾后 N 秒（需配合 --duration）")
    pc.add_argument("--duration", type=int, default=0,
                    help="视频总时长（秒），用于片尾计算")
    pc.add_argument("--url", default="",
                    help="B 站视频 URL：自动抓弹幕检测赞助段并合并进 --ranges")
    pc.add_argument("--out", default="",
                    help="输出路径（默认覆盖 input）")

    psf = sub.add_parser("suggest-frames",
                         help="#2 依据官方大纲/弹幕高能点推荐截图时间点")
    psf.add_argument("url", help="B 站视频 URL 或 BV/av 号")
    psf.add_argument("--cookie", default="")
    psf.add_argument("--n", type=int, default=6, help="推荐数量（默认 6）")
    psf.add_argument("--markers", action="store_true",
                     help="输出为 *Screenshot-[mm:ss]* 标记行")

    pg = sub.add_parser("gzh",
                        help="E8 将 note.md 排成公众号内联 HTML"
                             "（借鉴 gzh-design-skill 排版理念）")
    pg.add_argument("input", help="输入 note.md 路径")
    pg.add_argument("--out", default="",
                    help="输出 HTML 路径（默认同目录 <名>.gzh.html）")
    pg.add_argument("--title", default="",
                    help="封面大标题（默认取首个 H1 或文件名）")
    pg.add_argument("--theme", default="moyu-green",
                    help="主题：moyu-green/red-white/graphite/zen/ink-blue/olive")
    pg.add_argument("--lint", action="store_true",
                    help="仅对可粘贴正文做禁用标签自检，不写文件")

    pr = sub.add_parser("review",
                        help="Plan C 质量闸：审稿一致性/覆盖度质检（源转录 vs 笔记）")
    pr.add_argument("note", help="待质检的 note.md 路径")
    pr.add_argument("--transcript", default="",
                    help="源转录文本文件路径（用于事实核对）")
    pr.add_argument("--out", default="",
                    help="质检报告输出路径（默认同目录 <名>.review.md）")
    pr.add_argument("--auto-fix", action="store_true",
                    help="Plan I 质量闭环：质检后自动回修问题段并复检"
                         "（默认上限 2 轮，达标 --target-score 即停）")
    pr.add_argument("--target-score", type=float, default=4.5,
                    help="auto-fix 达标分数线（0-5，默认 4.5）")
    pr.add_argument("--max-rounds", type=int, default=2,
                    help="auto-fix 最大轮次（默认 2）")

    psp = sub.add_parser("style-preview",
                         help="Plan D 风格预览：查看某个 --style 会注入笔记的完整提示词"
                              "（风格说明 + 平台/通用详细规范）")
    psp.add_argument("style", nargs="?", default="",
                     help="风格键名（留空则列出全部 11 种风格及一句话说明）")

    pp = sub.add_parser("produce",
                        help="Plan B 一源多产：以 note.md 为单一事实源，"
                             "派生多份风格化/多平台产物")
    pp.add_argument("note", help="单一事实源 note.md 路径")
    pp.add_argument("--targets", default="",
                    help="逗号分隔的衍生目标：风格键(xiaohongshu/wechat/...)"
                         " / gzh / mindmap / pdf / docx；"
                         "默认 xiaohongshu,wechat,bilibili")
    pp.add_argument("--out-dir", default="",
                    help="产物输出目录（默认与 note.md 同目录）")
    pp.add_argument("--preset", default="",
                    help="成品包预设：knowledge-ip / courseware / social-matrix"
                         "（自动展开 targets + 跑扩展简报）")

    # ---- Plan A: E1–E7 扩展闭环固化 ----
    phz = sub.add_parser("humanize",
                         help="Plan A/E6 去 AI 味：本地扫描 10 类 AI 写作痕迹 + native 改写")
    phz.add_argument("note", help="待处理的 note.md 路径")
    phz.add_argument("--scan-only", action="store_true",
                     help="仅本地扫描并打印 AI 痕迹报告，不改写")
    phz.add_argument("--out", default="", help="改写输出路径（默认 <名>.humanized.md）")

    pil = sub.add_parser("illustrate",
                         help="Plan A/E3-E4 配图简报：抽取配图锚点 -> ImageGen 生图简报")
    pil.add_argument("note", help="note.md 路径")
    pil.add_argument("--mode", default="illustration",
                     choices=["illustration", "infographic"],
                     help="illustration=小黑手绘逐张(E3)；infographic=一页图解(E4)")
    pil.add_argument("--points", type=int, default=4, help="配图锚点数量上限")
    pil.add_argument("--out", default="")

    prs = sub.add_parser("research",
                         help="Plan A/E2 研究补充：抽取待证据主张 -> WebSearch 研究简报")
    prs.add_argument("note", help="note.md 路径")
    prs.add_argument("--max", type=int, default=8, help="主张数量上限")
    prs.add_argument("--out", default="")

    psl = sub.add_parser("slides",
                         help="Plan A/E5 网页PPT：抽取叙事弧大纲 -> PPT 简报")
    psl.add_argument("note", help="note.md 路径")
    psl.add_argument("--minutes", type=int, default=15, help="演讲时长(分钟)，决定页数")
    psl.add_argument("--out", default="")

    # ---- Plan E9: 零 key 社媒卡发布（v3.0 迭代 1） ----
    psc = sub.add_parser(
        "social-card",
        help="E9 零 key 社媒卡：note.md -> 小红书 3:4 卡 + 公众号封面（零依赖）")
    psc.add_argument("note", help="note.md 路径")
    psc.add_argument("--out", default="", help="输出目录（默认 <note_dir>/social-cards）")
    psc.add_argument("--platforms", default="xiaohongshu,wechat",
                     help="逗号分隔：xiaohongshu,wechat,zhihu,x")
    psc.add_argument("--png", action="store_true",
                     help="尝试栅格化为 PNG（需 playwright/wkhtmltoimage）")

    # ---- Plan E3: 配图科学化 · 构思卡 + 八项测试（v3.0 迭代 3） ----
    pip = sub.add_parser(
        "illustration-plan",
        help="E3 配图科学化：note.md -> 构思卡+八项测试 简报（交 agent 生图）")
    pip.add_argument("note", help="note.md 路径")
    pip.add_argument("--max-cards", type=int, default=4,
                     help="候选配图位上限（默认 4，宁少勿滥）")
    pip.add_argument("--out", default="", help="输出 plan 简报路径")

    # ---- 迭代2·质量收口：质量门禁 ----
    pqg = sub.add_parser(
        "quality-gate",
        help="迭代2质量门禁：E6评分+残留痕迹+来源可信度 确定性判定")
    pqg.add_argument("note", help="原始/定稿 note.md 路径")
    pqg.add_argument("--humanized", default="",
                     help="E6 改写后笔记路径（含评分卡合计行）")
    pqg.add_argument("--review", default="",
                     help="质检报告 .review.md 路径（解析 [原文]/[检索]/[推断]）")
    pqg.add_argument("--json", action="store_true", help="输出 JSON")

    # ---- Plan E: 批处理与系列化 ----
    pb = sub.add_parser(
        "batch",
        help="Plan E 批处理/系列化：对多源(清单/目录/多参数)批量跑 "
             "run/produce/humanize/… ，失败隔离 + 汇总报告")
    pb.add_argument("sources", nargs="*",
                    help="直接传入的多个源(URL/文件)；也可用 --manifest / --dir")
    pb.add_argument("--manifest", default="",
                    help="清单文件：.json（数组或 {items,defaults}）或 "
                         ".txt（每行一个源，支持 '源 | 标题'，# 注释）")
    pb.add_argument("--dir", dest="dir_", default="",
                    help="扫描目录取源（配合 --glob，默认 *.md），"
                         "常用于对一批 note.md 批量后处理")
    pb.add_argument("--glob", dest="glob_", default="*.md",
                    help="--dir 扫描通配符（默认 *.md）")
    pb.add_argument("--op", default="run", choices=list(BATCH_OPS),
                    help="每个源执行的操作（默认 run 全流程出笔记）")
    pb.add_argument("--then", default="",
                    help="仅 --op run：对每篇产出的真实笔记再链式执行的后处理，"
                         "逗号分隔，如 produce,humanize")
    pb.add_argument("--out-dir", dest="out_dir", default="./batch-out",
                    help="统一输出目录（默认 ./batch-out）")
    pb.add_argument("--jobs", type=int, default=1,
                    help="并发数（默认 1 串行；>1 用线程池，适合下载/LLM I/O）")
    pb.add_argument("--stop-on-error", action="store_true",
                    help="遇错即停（默认失败隔离，继续处理其余项）")
    pb.add_argument("--style", default=DEFAULT_STYLE, choices=list(STYLES.keys()),
                    help="--op run 的笔记风格")
    pb.add_argument("--targets", default="",
                    help="--op produce / --then produce 的目标（见 produce）")
    pb.add_argument("--preset", default="",
                    help="--op produce / --then produce 的成品包预设"
                         "（knowledge-ip / courseware / social-matrix）")
    pb.add_argument("--mode", default="illustration",
                    choices=["illustration", "infographic"],
                    help="--op illustrate 的配图模式")
    pb.add_argument("--points", type=int, default=4)
    pb.add_argument("--max", type=int, default=8)
    pb.add_argument("--minutes", type=int, default=15)
    pb.add_argument("--lang", default=None)
    pb.add_argument("--no-clean", action="store_true")
    pb.add_argument("--review", action="store_true")
    pb.add_argument("--export", default="")

    # ---- Plan G: 缓存管理 ----
    pcache = sub.add_parser(
        "cache",
        help="Plan G 缓存管理：查看/清空同源取数缓存（ls / clear）")
    pcache.add_argument("action", choices=["ls", "clear"],
                        help="ls=列出各类别条目数与大小；clear=清空")
    pcache.add_argument("--category", default="",
                        help="clear 时只清该类别（ai_summary/danmaku/comments/transcribe）；"
                             "默认清全部")

    # ---- Plan L: 度量面板 ----
    pmetrics = sub.add_parser(
        "metrics",
        help="Plan L 产出度量：查看事件分布/风格分布/质量评分趋势")

    # ---- Plan H: 系列聚合 ----
    pseries = sub.add_parser(
        "series",
        help="Plan H 跨笔记聚合：多份 note.md -> 系列总导图+索引页+知识图谱")
    pseries.add_argument("notes", nargs="*", help="待聚合的 note.md 列表")
    pseries.add_argument("--dir", dest="dir_", default="",
                         help="扫描目录取 note.md（配合 --glob，默认 *.md）")
    pseries.add_argument("--glob", dest="glob_", default="*.md")
    pseries.add_argument("--title", default="系列专题", help="系列名称")
    pseries.add_argument("--out-dir", dest="out_dir", default="",
                         help="产物输出目录（默认与第一篇同目录）")

    args = p.parse_args(argv)

    if args.cmd == "cache":
        if args.action == "ls":
            items = cache_list()
            if not items:
                print("缓存为空（目录：{}）".format(_cache_dir()))
            else:
                print(f"缓存目录：{_cache_dir()}")
                for it in items:
                    print(f"  {it['category']}: {it['count']} 条 / {it['size_kb']} KB")
        else:  # clear
            n = cache_clear(args.category)
            print(f"已清空{'类别 ' + args.category if args.category else '全部'}缓存，"
                  f"删除 {n} 条。")
        return

    if args.cmd == "metrics":
        print(metrics_summary())
        return

    if args.cmd == "series":
        notes = list(args.notes)
        if args.dir_:
            notes = [str(p) for p in sorted(Path(args.dir_).glob(args.glob_))]
        if not notes:
            log("series：未提供 note.md（用位置参数或 --dir）")
            return
        res = series_aggregate(notes, out_dir=args.out_dir, title=args.title)
        print(f"系列聚合：{res['count']} 篇")
        print(f"  总导图：{res['mindmap']}")
        print(f"  索引页：{res['index']}")
        print(f"  知识图谱：{res['graph']}")
        return

    if args.cmd == "fetch":
        meta = fetch_local(args.url, Path(args.workdir)) if args.url.endswith(
            (".mp4", ".mkv", ".webm", ".mov", ".avi", ".flv")) else fetch(
            args.url, Path(args.workdir), lang=args.lang)
        print(json.dumps(meta, ensure_ascii=False, indent=2))
        return

    if args.cmd == "transcribe":
        text = dispatch_transcribe(args.audio, args.backend)
        Path(args.out).write_text(text, encoding="utf-8")
        Path(args.out).write_text(text, encoding="utf-8")
        log(f"Transcript -> {args.out}")
        return

    if args.cmd == "summarize":
        transcript = Path(args.transcript).read_text(encoding="utf-8")
        kind, payload = summarize(transcript, args.title, args.tags,
                         not args.no_time_markers, args.screenshots,
                         not args.no_ai_summary, out=args.out,
                         style=args.style,
                         outline_first=getattr(args, "outline_first", False))
        if kind == "note":
            Path(args.out).write_text(payload, encoding="utf-8")
            log(f"Note -> {args.out}")
            if args.export:
                export(args.out, Path(args.out).with_suffix(""),
                       [f for f in args.export.split(",") if f])
        else:
            log(f"简报已生成：{payload}。请使用当前对话 LLM 生成笔记并写入 {args.out}。")
        return

    if args.cmd == "run":
        if getattr(args, "no_cache", False):
            global CACHE_ENABLED
            CACHE_ENABLED = False
            log("[cache] 已关闭（--no-cache），强制重新抓取/转写")
        if getattr(args, "wizard", False):
            run_wizard(args.url, args.out or "note.md")
            return
        run(args.url, Path(args.out), local=None, lang=args.lang,
            transcriber=args.transcriber,
            time_markers=not args.no_time_markers, screenshots=args.screenshots,
            ai_summary=not args.no_ai_summary,
            export_formats=[f for f in args.export.split(",") if f],
            style=args.style,
            embed_screenshots_flag=args.embed_screenshots,
            include_images=args.images,
            multimodal=getattr(args, 'multimodal', False),
            multimodal_frames=getattr(args, 'multimodal_frames', 6),
            subtitle_file=args.subtitle_file,
            bili_ai=not args.no_bili_ai,
            danmaku=args.danmaku,
            comments=args.comments,
            comments_in_report=args.comments_in_report,
            bili_cookie=args.bili_cookie,
            linkify=args.linkify,
            clean=not args.no_clean,
            translate=args.translate,
            review=args.review,
            outline_first=getattr(args, "outline_first", False))
        return

    if args.cmd == "embed-screenshots":
        video = args.video
        if not video and args.url:
            wd = Path(args.workdir)
            wd.mkdir(parents=True, exist_ok=True)
            vp = wd / "video.mp4"
            log(f"Downloading video from {args.url} for frame extraction ...")
            subprocess.run(_ytdlp_cmd() + _ytdlp_extra_opts()
                           + ["-f", _build_video_format(args.video_quality), "-o", str(vp),
                              "--no-warnings", args.url],
                           check=True)
            video = str(vp)
        if not video:
            log("ERROR: provide --video PATH or --url URL.")
            return
        embed_screenshots(args.note, video, args.out or None)
        return

    if args.cmd == "embed-article-images":
        import json as _json
        raw = (args.images or "").strip()
        if raw.startswith("["):
            img_list = _json.loads(raw)
        else:
            img_list = [u.strip() for u in raw.split(",") if u.strip()]
        embed_article_images(args.note, img_list, args.out or None)
        return

    if args.cmd == "export":
        out_base = Path(args.out) if args.out else Path(args.note).with_suffix("")
        export(args.note, out_base, [f for f in args.formats.split(",") if f])
        return

    if args.cmd == "bili-summary":
        cookie = args.cookie or os.environ.get("BILINOTE_BILI_COOKIE", "")
        ai = fetch_bilibili_ai_summary(args.url, cookie)
        if not ai.get("ok"):
            log(f"官方 AI 总结不可用：{ai.get('reason')}")
            if not cookie:
                log("提示：多数视频需登录，请用 --cookie 传入含 SESSDATA 的 Cookie "
                    "或设置环境变量 BILINOTE_BILI_COOKIE。")
            return
        if args.transcript:
            Path(args.transcript).write_text(
                bili_ai_summary_to_transcript(ai), encoding="utf-8")
            log(f"AI 字幕转写 -> {args.transcript}")
        if args.out:
            Path(args.out).write_text(
                bili_ai_summary_to_markdown(ai), encoding="utf-8")
            log(f"AI 总结 Markdown -> {args.out}")
        if args.json or not (args.out or args.transcript):
            print(json.dumps(ai, ensure_ascii=False, indent=2))
        return

    if args.cmd == "danmaku":
        cookie = args.cookie or os.environ.get("BILINOTE_BILI_COOKIE", "")
        cid = None
        dur = None
        try:
            v = _bili_get_view(_bili_extract_id(args.url), cookie)
            cid = v.get("cid")
            dur = v.get("duration")
        except Exception as e:
            log(f"解析 cid 失败：{e}")
        dm = (fetch_bilibili_danmaku(cid, cookie, args.limit, duration=dur)
              if cid else [])
        log(f"抓取弹幕 {len(dm)} 条。")
        ctx = danmaku_as_context(dm)
        if args.out:
            Path(args.out).write_text(ctx, encoding="utf-8")
            log(f"弹幕上下文 -> {args.out}")
        else:
            print(ctx)
        return

    if args.cmd == "comments":
        cookie = args.cookie or os.environ.get("BILINOTE_BILI_COOKIE", "")
        aid = None
        try:
            v = _bili_get_view(_bili_extract_id(args.url), cookie)
            aid = v.get("aid")
        except Exception as e:
            log(f"解析 aid 失败：{e}")
        raw = (fetch_bilibili_comments(aid, cookie, limit=args.limit, sort="hot")
               if aid else [])
        log(f"抓取评论 {len(raw)} 条。")
        hi = select_highlight_comments(raw, top_n=args.top,
                                        min_likes=args.min_likes)
        log(f"精选高能评论 {len(hi)} 条。")
        out_text = (comments_as_report_section(hi) if args.report
                    else comments_as_context(hi))
        if args.out:
            Path(args.out).write_text(out_text, encoding="utf-8")
            log(f"评论精选 -> {args.out}")
        else:
            print(out_text)
        return

    if args.cmd == "linkify":
        linkify_note(args.note, args.url, args.out or None)
        return

    if args.cmd == "clean":
        text = Path(args.input).read_text(encoding="utf-8")
        ranges = []
        for part in [p for p in args.ranges.split(",") if p]:
            a, b = part.split("-")
            ranges.append([float(a), float(b)])
        if args.url:
            try:
                _v = _bili_get_view(_bili_extract_id(args.url))
                _dm = fetch_bilibili_danmaku(_v.get("cid"), duration=_v.get("duration"))
                _r = detect_sponsor_segments_from_danmaku(_dm)
                if _r:
                    ranges += _r
                    log(f"弹幕检测合并 {len(_r)} 段疑似赞助区间。")
            except Exception as e:
                log(f"弹幕检测失败（{e}），仅用手动 --ranges。")
        cleaned, dropped = clean_transcript(
            text, ranges, args.intro, args.outro,
            args.duration or None)
        dst = Path(args.out) if args.out else Path(args.input)
        dst.write_text(cleaned, encoding="utf-8")
        log(f"内容净化完成：去除 {dropped} 行 -> {dst}")
        print(json.dumps(ranges, ensure_ascii=False, indent=2))
        return

    if args.cmd == "suggest-frames":
        cookie = args.cookie or os.environ.get("BILINOTE_BILI_COOKIE", "")
        ai = fetch_bilibili_ai_summary(args.url, cookie)
        outline = ai.get("outline") if ai.get("ok") else None
        duration = ai.get("duration") if ai.get("ok") else None
        cid = ai.get("cid") if ai.get("ok") else None
        if not cid:
            try:
                view = _bili_get_view(_bili_extract_id(args.url), cookie)
                cid = view.get("cid")
                duration = duration or view.get("duration")
            except Exception:
                pass
        dm = fetch_bilibili_danmaku(cid, cookie, duration=duration) if cid else []
        ts = suggest_screenshot_timestamps(outline=outline, danmaku=dm or None,
                                           duration=duration, n=args.n)
        if args.markers:
            print(markers_from_timestamps(ts))
        else:
            print(json.dumps([fmt_time(t) for t in ts], ensure_ascii=False))
        return

    if args.cmd == "gzh":
        try:
            import gzh_format
        except ImportError:
            sys.path.insert(0, str(Path(__file__).resolve().parent))
            import gzh_format
        src = Path(args.input)
        md = src.read_text(encoding="utf-8")
        title = args.title or _gzh_guess_title(md, src.stem)
        if args.lint:
            html = gzh_format.render(md, theme=args.theme, title=title)
            body = gzh_format.extract_body(html)
            issues = gzh_format.lint_gzh_body(body)
            if issues:
                log("公众号正文自检发现问题：")
                for it in issues:
                    log("  - " + it)
                sys.exit(1)
            log("公众号正文自检通过：无禁用标签，全内联。")
            return
        out = Path(args.out) if args.out else src.with_suffix(".gzh.html")
        gzh_format.convert_file(str(src), str(out), theme=args.theme, title=title)
        log(f"公众号 HTML -> {out}")
        log("在浏览器打开后点『复制到公众号』，再粘贴进公众号编辑器即可。")
        return

    if args.cmd == "review":
        if getattr(args, "auto_fix", False):
            transcript = ""
            if args.transcript:
                transcript = Path(args.transcript).read_text(encoding="utf-8")
            res = review_note_with_autofix(
                args.note, transcript,
                auto_fix=True, target_score=args.target_score,
                max_rounds=args.max_rounds)
            log(f"质量闭环：{res['rounds']} 轮，末轮评分 {res['final_score']}，"
                f"停止原因：{res['stopped_reason']}")
            if res.get("fixed_path"):
                log(f"修复稿 -> {res['fixed_path']}")
            return
        kind, payload = review_note(
            args.note, transcript_path=args.transcript, out=args.out)
        if kind == "review":
            log(f"质检报告 -> {payload}")
        else:
            log(f"质检 brief（交由 agent 完成）-> {payload}")
        return

    if args.cmd == "style-preview":
        print(preview_style(getattr(args, "style", "") or ""))
        return

    if args.cmd == "produce":
        results = produce_from_note(
            args.note, targets=args.targets,
            out_dir=args.out_dir, preset=args.preset)
        for kind, payload in results:
            log(f"产物（{kind}）-> {payload}")
        return

    if args.cmd == "humanize":
        note_md = Path(args.note).read_text(encoding="utf-8")
        if args.scan_only:
            print(format_scan_report(scan_ai_patterns(note_md)))
            return
        kind, payload, findings = humanize_note(args.note, out=args.out)
        log(format_scan_report(findings))
        if kind == "note":
            log(f"去 AI 味改写 -> {payload}")
        else:
            log(f"去 AI 味简报（交由 agent 改写）-> {payload}")
        return

    if args.cmd == "social-card":
        from social_card import generate_social_cards
        res = generate_social_cards(args.note, args.out, args.platforms, args.png)
        n = sum(len(v) for k, v in res.items() if not k.startswith("_"))
        log(f"社媒卡生成完成：{n} 个 HTML；索引 {res.get('_index','')}")
        if res.get("_png"):
            log(f"PNG：{len(res['_png'])} 张")
        return

    if args.cmd == "quality-gate":
        quality_gate(args.note, args.humanized, args.review, args.json)
        return

    if args.cmd == "illustration-plan":
        _, p, cards = illustration_plan(args.note, args.max_cards, args.out)
        n_pass = sum(1 for c in cards if c["passed"])
        log(f"配图构思卡生成完成：{len(cards)} 张（{n_pass} 张八项全过）；简报 {p}")
        return

    if args.cmd == "illustrate":
        kind, payload = illustrate_note(
            args.note, mode=args.mode, points=args.points, out=args.out)
        log(f"配图简报（{args.mode}）-> {payload}")
        return

    if args.cmd == "research":
        kind, payload = research_note(args.note, max_claims=args.max, out=args.out)
        log(f"研究简报 -> {payload}")
        return

    if args.cmd == "slides":
        kind, payload = slides_note(args.note, minutes=args.minutes, out=args.out)
        log(f"PPT 简报 -> {payload}")
        return

    if args.cmd == "batch":
        items = load_manifest(manifest=args.manifest, dir_=args.dir_,
                              glob_=args.glob_, sources=args.sources)
        if not items:
            log("batch：没有可处理的源（请用 位置参数 / --manifest / --dir 提供）")
            return
        then = [t.strip() for t in args.then.split(",") if t.strip()]
        opts = {"style": args.style, "targets": args.targets, "preset": args.preset,
                "mode": args.mode,
                "points": args.points, "max_claims": args.max,
                "minutes": args.minutes, "lang": args.lang,
                "clean": not args.no_clean, "review": args.review,
                "export_formats": [f for f in args.export.split(",") if f]}
        records = run_batch(items, op=args.op, out_dir=args.out_dir,
                            jobs=args.jobs, then=then,
                            stop_on_error=args.stop_on_error, opts=opts)
        out_base = Path(args.out_dir)
        out_base.mkdir(parents=True, exist_ok=True)
        (out_base / "batch-report.md").write_text(
            format_batch_report(records, args.op), encoding="utf-8")
        (out_base / "batch-report.json").write_text(
            json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
        ok = sum(1 for r in records if r["status"] == "ok")
        log(f"批处理完成：{ok}/{len(records)} 成功 -> "
            f"{out_base / 'batch-report.md'}")
        return


def preview_style(style: str = "") -> str:
    """Plan D：渲染某个风格注入笔记时的完整提示词，便于产出前预览与选型。

    传入具体 style 键名 → 返回该风格的一句话说明 + 详细规范（若有补丁）；
    留空 → 返回全部 11 种风格的清单（键名 + 标签 + 是否含详细补丁）。
    与真实注入逻辑共享 STYLES / _STYLE_PROMPTS_CACHE，所见即所得。"""
    style = (style or "").strip()
    if not style:
        lines = [f"可用风格共 {len(STYLES)} 种（* = 含详细规范补丁）：", ""]
        for k, (label, instr) in STYLES.items():
            mark = "*" if get_style_spec(k) else " "
            lines.append(f" {mark} {k:<12} {label} — {instr.split('：', 1)[-1][:40]}…")
        lines += ["", f"默认风格：{DEFAULT_STYLE}",
                  "查看单个风格完整提示词：style-preview <风格键名>"]
        return "\n".join(lines)
    if style not in STYLES:
        return (f"未知风格「{style}」。可用：{', '.join(STYLES.keys())}\n"
                f"（留空运行 style-preview 可列出全部风格）")
    label, instruction = STYLES[style]
    out = [f"# 风格预览：{style}（{label}）", "",
           "## 报告风格要求（注入所有模式）", instruction]
    detail = get_style_spec(style)
    if detail:
        out += ["", f"## 平台/通用详细规范（{label}，注入 API 与 native 两种模式）",
                "", detail]
    else:
        out += ["", "（该风格暂无 references/style_prompts.md 详细补丁，"
                "仅注入上方风格要求。）"]
    return "\n".join(out)


# --------------------------------------------------------------------------
# Plan B: 一源多产内容工厂（produce）
#   以 note.md 为唯一事实源（SSOT），并行派生多份风格化/多平台产物，
#   不重新跑 summarize，避免重复转写与事实漂移。
# --------------------------------------------------------------------------


def build_produce_prompt(note_md: str, style: str) -> str:
    """拼接 produce 改写 user prompt：源笔记 + 目标风格规范。"""
    label, instruction = STYLES.get(style, STYLES[DEFAULT_STYLE])
    detail = get_style_spec(style)
    parts = [
        f"# 改写任务：将下面笔记改写为「{label}」风格\n",
        "## 改写要求（必须严格遵循）\n",
        f"1. 目标风格：{label}。{instruction}\n",
    ]
    if detail:
        parts.append(
            f"2. 平台/通用详细规范（{label}），直接套用结构化模板：\n\n"
            f"{detail}\n")
    parts.append(
        "3. 事实守恒：必须完整保留原文所有关键论点、数据、案例、结论、建议；"
        "不得增删事实、不得编造、不得改变原意，只能改变表达形式。\n"
        "4. 仅输出改写后的 Markdown 正文，不包裹代码块，不重复原文。\n"
        "---\n\n"
        "## 源笔记（单一事实源）\n---\n\n"
        f"{note_md}\n"
    )
    return "\n".join(parts)


def _write_produce_brief(note: str, target_style: str, out_path: Path) -> str:
    """agent-native 模式：为某风格改写写一份自足简报，交 agent 当前 LLM。"""
    label, instruction = STYLES.get(target_style, STYLES[DEFAULT_STYLE])
    detail = get_style_spec(target_style)
    rules = [
        "仅返回改写后的 Markdown 正文，不要包裹在代码块中。",
        "使用中文撰写，专有名词/术语保留英文。",
        f"**目标风格 = {label}**：{instruction}",
    ]
    if detail:
        rules.append(
            f"**目标风格详细规范（{label}）**：严格按以下模板改写：\n\n{detail}")
    rules.append(
        "**事实守恒**：必须完整保留源笔记所有关键论点/数据/案例/结论/建议；"
        "不得增删事实、不得编造、不得改变原意，只改表达/结构/排版/语气。")
    rules.append("不要重新生成内容，只基于下方源笔记做风格化改写。")
    rules_text = ("## 改写规则（必须严格遵循）\n"
                  + "\n".join(f"{i+1}. {r}" for i, r in enumerate(rules)) + "\n")
    content = (
        "# BiliNote 风格改写简报（agent 原生模式）\n\n"
        f"请使用**当前对话的 LLM** 将下方【源笔记】改写为「{label}」风格，"
        f"无需任何外部 API Key。\n\n"
        "## 元数据\n"
        f"- 目标风格: {target_style}（{label}）\n"
        f"- 输出文件: {out_path}\n\n"
        f"{rules_text}\n"
        "---\n\n"
        "## 源笔记（单一事实源，请勿增删事实）\n---\n\n"
        f"{note}\n"
        "---\n\n"
        f"改写后写入：{out_path}\n"
    )
    return content


# ---- Plan B 扩展：成品包预设（produce --preset） ----
# 每个预设 = 一组 targets（风格键/gzh/mindmap/pdf/docx）+ 一组 extensions
# （humanize/illustrate/research/slides，复用同名子命令的既有实现）。
# 用法：produce note.md --preset knowledge-ip  等价于自动展开 targets + 跑扩展简报。
PRODUCE_PRESETS = {
    "knowledge-ip": {
        "desc": "知识IP专栏包：detailed底 + 公众号 + 思维导图 + PDF + 去AI味 + 信息图 + 研究",
        "targets": "detailed,wechat,gzh,mindmap,pdf",
        "extensions": ["humanize", "illustrate", "research"],
    },
    "courseware": {
        "desc": "课程课件包：tutorial底 + 公众号 + 思维导图 + PDF + 手绘 + PPT",
        "targets": "tutorial,gzh,mindmap,pdf",
        "extensions": ["illustrate", "slides"],
    },
    "social-matrix": {
        "desc": "社媒矩阵包：小红书 + B站 + 知乎 + 信息图",
        "targets": "xiaohongshu,bilibili,zhihu",
        "extensions": ["illustrate"],
    },
}


def _run_produce_extension(name: str, note_path: str) -> tuple:
    """跑单个扩展（复用 humanize/illustrate/research/slides 既有实现）。
    返回 (ext_name, kind, payload)。失败安全：异常返回 ('error', ...)。"""
    try:
        if name == "humanize":
            k, p, _ = humanize_note(note_path)
            return (name, k, p)
        if name == "illustrate":
            k, p = illustrate_note(note_path, mode="infographic")
            return (name, k, p)
        if name == "research":
            k, p = research_note(note_path)
            return (name, k, p)
        if name == "slides":
            k, p = slides_note(note_path)
            return (name, k, p)
        return (name, "skip", f"未知扩展 {name}")
    except Exception as e:  # noqa: BLE001
        return (name, "error", str(e))


def produce_from_note(note_path: str, targets: str = "",
                      out_dir: str = "", preset: str = "") -> list:
    """Plan B 一源多产编排器（agent 原生模式，唯一模式）。

    以 note.md 为唯一事实源，派生多份产物。返回 [(kind, payload), ...]，
    kind ∈ {'brief','gzh','mindmap','export','ext'}：
      - 风格键 (xiaohongshu/wechat/...): 改写简报 -> <stem>.<style>.produce-brief.md
      - gzh:    公众号内联 HTML  -> <stem>.gzh.html
      - mindmap: Mermaid 思维导图 -> <stem>.mindmap.md
      - pdf/docx: 导出            -> <stem>.pdf / <stem>.docx
    风格改写走原生简报（交 agent 当前 LLM 改写）；其余为本地确定性产物。
    preset（成品包）：选中后展开为预设的 targets + extensions，复用同名子命令实现。
    """
    src = Path(note_path)
    note = src.read_text(encoding="utf-8")
    stem = src.stem
    out_base = Path(out_dir) if out_dir else src.parent
    out_base.mkdir(parents=True, exist_ok=True)
    # 成品包预设展开
    ext_list = []
    if preset:
        ps = PRODUCE_PRESETS.get(preset)
        if ps:
            if not targets:
                targets = ps["targets"]
            ext_list = ps.get("extensions", [])
            log(f"produce 预设 {preset}：{ps['desc']}")
        else:
            log(f"produce 未知预设 {preset}，忽略；可用：{list(PRODUCE_PRESETS)}")
    targets_list = [t.strip() for t in (targets or "").split(",") if t.strip()] \
        or ["xiaohongshu", "wechat", "bilibili"]
    results = []
    for t in targets_list:
        if t in STYLES:
            brief_path = out_base / f"{stem}.{t}.produce-brief.md"
            brief_path.write_text(
                _write_produce_brief(note, t, brief_path), encoding="utf-8")
            results.append(("brief", str(brief_path)))
        elif t == "gzh":
            try:
                import gzh_format
            except ImportError:
                sys.path.insert(0, str(Path(__file__).resolve().parent))
                import gzh_format
            out_html = out_base / f"{stem}.gzh.html"
            title = _gzh_guess_title(note, stem)
            gzh_format.convert_file(str(src), str(out_html), title=title)
            results.append(("gzh", str(out_html)))
        elif t == "mindmap":
            out_mm = out_base / f"{stem}.mindmap.md"
            out_mm.write_text(generate_mindmap(note), encoding="utf-8")
            results.append(("mindmap", str(out_mm)))
        elif t in ("pdf", "docx"):
            export(str(src), out_base / stem, [t])
            results.append(("export", str((out_base / stem).with_suffix("." + t))))
        else:
            log(f"produce 跳过未知目标：{t}")
    # 成品包扩展（targets 之后再跑后处理简报）
    for name in ext_list:
        ext_name, kind, payload = _run_produce_extension(name, str(src))
        results.append(("ext", f"{ext_name}:{kind}:{payload}"))
        log(f"produce 扩展 {ext_name} -> {kind}")
    record_metric("produce", targets=",".join(targets_list),
                  preset=preset, count=len(results))
    return results


# --------------------------------------------------------------------------
# Plan A: 扩展能力闭环固化（E1–E7 代码化）
#   把原本仅文档化、靠 agent 临场发挥的后处理扩展，固化为「确定性脚手架 +
#   native 简报」：本地能算的（AI 痕迹扫描、主张抽取、锚点抽取、大纲、SPA 检测、
#   顺序编排）用纯 Python 算死；需要 agent 工具的（ImageGen/WebSearch）产出
#   自足简报交当前 LLM/工具执行。让后处理链可脚本触发、可测试、顺序可控。
# --------------------------------------------------------------------------

# ---- E7 扩展调用顺序（规范后处理链） ----
EXT_PIPELINE_ORDER = ["E1", "E2", "E3", "E6", "E4", "E5", "E8"]


def order_extensions(codes) -> list:
    """E7：把一组扩展代码按规范后处理链顺序排列。

    顺序：E1(JS渲染) -> E2(补料) -> E3(配图嵌入) -> [note 定稿]
          -> E6(去AI味) -> E4/E5(信息图/PPT) -> E8(公众号排版, 永远最后)。
    未知代码忽略。"""
    want = {str(c).strip().upper() for c in (codes or []) if str(c).strip()}
    return [c for c in EXT_PIPELINE_ORDER if c in want]


# ---- E1 浏览器自动化：SPA / 需 JS 渲染检测 ----
_SPA_MARKERS = [
    r'<div[^>]*\bid=["\']root["\']', r'<div[^>]*\bid=["\']app["\']',
    r'window\.__NUXT__', r'window\.__NEXT_DATA__', r'__INITIAL_STATE__',
    r'ng-version', r'data-reactroot', r'v-cloak',
]


def needs_js_render(html: str, extracted_len: int = 0) -> bool:
    """E1：判断页面是否需要浏览器 JS 渲染（SPA / 动态加载）。

    命中任一 SPA 标志且已提取正文过短（<200 字）即判定需要浏览器渲染，
    供 agent 决定是否自动切到 browser-use 流程。"""
    if not html:
        return False
    has_spa = any(re.search(p, html, re.I) for p in _SPA_MARKERS)
    return has_spa and extracted_len < 200


# ---- E6 AI 痕迹去除：本地确定性扫描 + native 改写 ----
# 每类给出若干正则；命中即计数并取样，作为改写提示的优先级信号。
# E6 · 24 类 AI 写作痕迹（正则可检子集）
# 前 10 类继承自 v2.1；新增 14 类（语义类由 LLM 在改写时自查，见 SKILL.md §E6）。
AI_PATTERNS = [
    ("过度强调意义", [r"标志着[^。！]{0,24}(里程碑|新时代|新篇章|新纪元|转折)",
                r"具有(重要|深远|重大)(意义|价值|影响)", r"意义重大", r"至关重要"]),
    ("过度强调知名度", [r"广受好评", r"备受瞩目", r"享誉[^，。]{0,10}",
                 r"深受[^，。]{0,8}(喜爱|欢迎)", r"闻名(遐迩|中外)"]),
    ("肤浅分析铺垫", [r"值得(注意|一提|关注)的是", r"需要(指出|注意)的是",
                r"不难(发现|看出)", r"显而易见(的是)?", r"众所周知"]),
    ("宣传性语言", [r"革命性", r"颠覆性", r"前所未有", r"史无前例",
               r"完美(地|的)", r"极致(的)?", r"堪称", r"惊艳"]),
    ("模糊归因", [r"专家(表示|认为|指出|建议)", r"业内人士", r"有人(认为|说|指出)",
              r"据(悉|了解|报道)", r"相关人士"]),
    ("破折号滥用", [r"——", r"—"]),
    ("三段式法则", [r"第一[，、].{0,60}第二[，、].{0,60}第三",
               r"首先[，、].{0,60}其次[，、].{0,60}(最后|再者|此外)"]),
    ("AI英文词汇", [r"\bdelve\b", r"\bleverage\b", r"\bnavigate\b",
              r"\brobust\b", r"\bseamless\b", r"\bunlock\b", r"\btapestry\b"]),
    ("否定式排比", [r"不是[^，。]{1,20}而是[^。]{1,40}不是[^，。]{1,20}而是"]),
    ("连接短语堆砌", [r"此外", r"然而", r"因此", r"与此同时"]),
    # ---- v3.0 新增 14 类 ----
    ("系动词回避", [r"这意味着", r"这表明", r"这标志着", r"说明(了|：)", r"体现出",
               r"反映出", r"彰显(出)?"]),
    ("刻意换词同义循环", [r"换言之", r"换句话说", r"即[，、：]", r"也就是(说)?"]),
    ("虚假范围", [r"从[^，。]{1,12}到[^，。]{1,12}(的)?(跨越|转变|升级|演进)"]),
    ("协作交流痕迹", [r"希望这对您(有帮助|有所启发)", r"如有(任何|疑问|需要)",
                 r"如果(您|您有).{0,12}(问题|疑问)", r"希望(以上|我的).{0,8}(对您|能帮)"]),
    ("知识截止免责", [r"我的(知识|训练)(截止|截至)", r"作为(一个)?AI", r"我无法(获取|访问|提供).{0,8}(实时|最新)",
                 r"我(的)?(训练|知识)数据"]),
    ("谄媚语气", [r"您(真|非常|实在)(聪明|敏锐|有眼光|专业)", r"很高兴(为您|能帮|有机会)",
               r"您(提的|问的).{0,8}(问题)?(非常)?(好|棒|深刻)"]),
    ("标记符号滥用", [r"\*\*[^*]{1,40}\*\*", r"[\U0001F300-\U0001FAFF]", r"[“”]"]),
    ("模板化开场", [r"在当今", r"随着[^，。]{1,20}(发展|普及|兴起|到来)",
               r"在[^，。]{1,12}(时代|背景|环境)下", r"在[^，。]{1,10}的(今天|当下)"]),
    ("绝对化论断", [r"一定(会|要|能)", r"必然", r"毫无疑问", r"毋庸置疑",
               r"无可辩驳", r"亘古不变"]),
    ("冗余程度副词", [r"非常(地)?", r"极其", r"十分", r"格外", r"尤为", r"颇为"]),
    ("伪深刻反问", [r"难道[^？]{1,20}\？", r"何尝[^？]{1,20}\？", r"岂(能|可|不)[^？]{1,16}\？"]),
    ("总结词套话", [r"综上所述", r"总而言之", r"由此可见", r"总的来看", r"归根结底"]),
    ("空洞量词", [r"众多", r"大量(的)?", r"一系列", r"多种(的)?", r"各类", r"方方面面"]),
    ("数据夸张", [r"数倍", r"翻(倍|番)", r"暴涨", r"飙升", r"井喷", r"爆炸(式)?(增长|发展)"]),
]


def scan_ai_patterns(text: str) -> list:
    """E6：本地确定性扫描 24 类 AI 写作痕迹（正则可检子集）。

    返回 [(模式名, 命中数, [样例…]), ...]，仅保留有命中的模式。
    语义类（谄媚/模板开场等）正则可覆盖大部分，剩余由 LLM 改写时自查（见 SKILL.md §E6）。"""
    findings = []
    for name, patterns in AI_PATTERNS:
        count = 0
        examples = []
        for pat in patterns:
            for m in re.finditer(pat, text):
                count += 1
                s = m.group(0).strip()
                if s and s not in examples and len(examples) < 5:
                    examples.append(s)
        if count:
            findings.append((name, count, examples))
    return findings


def format_scan_report(findings: list) -> str:
    """把 scan_ai_patterns 结果渲染成可读报告。"""
    if not findings:
        return "未检测到明显 AI 写作痕迹（24 类模式均未命中）。"
    total = sum(c for _, c, _ in findings)
    lines = [f"检测到 {len(findings)} 类 AI 写作模式，共 {total} 处：", ""]
    for name, count, examples in findings:
        ex = "；".join(examples)
        lines.append(f"- **{name}** ×{count}" + (f"　例：{ex}" if ex else ""))
    return "\n".join(lines)


HUMANIZE_FIX_TABLE = (
    "| AI 模式 | 修复方法 |\n"
    "|---|---|\n"
    "| 过度强调意义 | 删除或替换为具体事实 |\n"
    "| 过度强调知名度 | 删掉修饰语，让事实说话 |\n"
    "| 肤浅分析铺垫 | 直接给结论，去掉「值得注意的是」类铺垫 |\n"
    "| 宣传性语言 | 用具体数据替代「革命性/颠覆性」 |\n"
    "| 模糊归因 | 补具体来源或删除 |\n"
    "| 破折号滥用 | 减少约 70%，只在真正补充说明时用（正文用 `--`） |\n"
    "| 三段式法则 | 打破为自然段落，混合句长 |\n"
    "| AI 英文词汇 | 用日常表达替代 delve/leverage/navigate/robust |\n"
    "| 否定式排比 | 「不是…而是…」改为肯定句 |\n"
    "| 连接短语堆砌 | 减少约 50%，用标点代替 |\n"
    "| 系动词回避 | 删「这意味着/表明」，直接说结论 |\n"
    "| 刻意换词同义循环 | 去掉「换言之」，保留一处即可 |\n"
    "| 虚假范围 | 用具体区间替代「从 X 到 Y」空泛表述 |\n"
    "| 协作交流痕迹 | 删除「希望这对您有帮助」等客套尾句 |\n"
    "| 知识截止免责 | 删除「作为 AI/我的知识截止」类声明 |\n"
    "| 谄媚语气 | 去掉「您真聪明」等奉承，保持平视 |\n"
    "| 标记符号滥用 | 去掉无意义的加粗/emoji/弯引号 |\n"
    "| 模板化开场 | 删「在当今/随着…发展」空泛起手 |\n"
    "| 绝对化论断 | 「一定/必然」改为有依据的限定表述 |\n"
    "| 冗余程度副词 | 删「非常/极其」，用事实强度本身 |\n"
    "| 伪深刻反问 | 反问改为陈述判断 |\n"
    "| 总结词套话 | 用一句话收束替代「综上所述」 |\n"
    "| 空洞量词 | 「众多/大量」改为可数或具体比例 |\n"
    "| 数据夸张 | 「数倍/井喷」改为精确数字 |\n")

# 50 分制五维评分卡模板（LLM 改写时自评填写，见 SKILL.md §E6）
HUMANIZE_SCORECARD = (
    "## 去 AI 味评分卡（改写后自评，满分 50）\n"
    "| 维度 | 分值 | 你的评分 | 说明 |\n"
    "|---|---|---|---|\n"
    "| 节奏（长短句交替） | 10 | /10 | |\n"
    "| 真实性（有观点不骑墙） | 10 | /10 | |\n"
    "| 信任度（无谄媚/无免责声明） | 10 | /10 | |\n"
    "| 直接性（无铺垫/无模板开场） | 10 | /10 | |\n"
    "| 精炼度（无冗余副词/空洞量词） | 10 | /10 | |\n"
    "| **合计** | **50** | **/50** | 45-50 优秀 / 35-44 良好 / <35 需重写 |\n")

HUMANIZE_SOUL = (
    "注入「人味」：有观点不骑墙；长短句交替换节奏；承认复杂性而非硬凑「综上所述」；"
    "合适场合可用「我认为/我觉得」；保留少量口语（「其实」「说白了」）。"
    "但须保持核心信息完整、匹配原文语调、绝不引入新的 AI 痕迹。")


def build_humanize_prompt(note_md: str, findings: list) -> str:
    """E6：拼接去 AI 味改写 user prompt（含本地扫描结果作为优先级信号）。"""
    return (
        "# 去 AI 味改写任务\n\n"
        "请去除下面这篇笔记里的 AI 写作痕迹，让它读起来自然、像真人写的。\n\n"
        "## 本地扫描已发现的 AI 模式（请优先修复这些）\n"
        f"{format_scan_report(findings)}\n\n"
        "## 修复对照表\n"
        f"{HUMANIZE_FIX_TABLE}\n"
        "## 注入灵魂\n"
        f"{HUMANIZE_SOUL}\n\n"
        f"{HUMANIZE_SCORECARD}\n"
        "## 输出要求\n"
        "- 保持全部关键信息 / 数据 / 结论不变，只改表达。\n"
        "- 仅输出改写后的 Markdown 正文，不要代码块、不要解释。\n"
        "- 改写后填写上方评分卡（合计需 ≥40 才算达标）。\n\n"
        "---\n\n## 待改写笔记\n---\n\n"
        f"{note_md}\n")


def _write_humanize_brief(note: str, findings: list, out_path: Path) -> str:
    """E6 native 模式：写一份自足去 AI 味简报，交 agent 当前 LLM 改写。"""
    return (
        "# BiliNote 去 AI 味简报（agent 原生模式）\n\n"
        "请使用**当前对话的 LLM** 将下方【待改写笔记】去除 AI 痕迹后改写，"
        "无需任何外部 API Key。\n\n"
        "## 本地已扫描出的 AI 模式（优先修复）\n"
        f"{format_scan_report(findings)}\n\n"
        "## 修复对照表\n"
        f"{HUMANIZE_FIX_TABLE}\n"
        "## 注入灵魂\n"
        f"{HUMANIZE_SOUL}\n\n"
        f"{HUMANIZE_SCORECARD}\n"
        "## 规则\n"
        "1. 保持全部关键信息 / 数据 / 结论不变，只改表达。\n"
        "2. 不得引入新的 AI 痕迹；匹配原文语调（学术风格别改成口语）。\n"
        "3. 改写后把评分卡（合计行必须保留）附在笔记末尾，连同正文一起写入目标文件。\n"
        "4. 仅输出改写后的 Markdown 正文 + 末尾评分卡，不要代码块包裹。\n\n"
        f"改写后写入：{out_path}\n\n"
        "---\n\n## 待改写笔记\n---\n\n"
        f"{note}\n")


def humanize_note(note_path: str, out: str = "") -> tuple:
    """E6：去 AI 味（agent 原生模式，唯一模式）。本地先扫描，再写简报。

    返回 (kind, payload, findings)：
      - ('brief', <stem>.humanize-brief.md, findings) 交 agent 当前 LLM 改写
    """
    src = Path(note_path)
    note = src.read_text(encoding="utf-8")
    findings = scan_ai_patterns(note)
    out_path = Path(out) if out else src.with_name(src.stem + ".humanize-brief.md")
    out_path.write_text(_write_humanize_brief(note, findings, out_path),
                        encoding="utf-8")
    return ("brief", str(out_path), findings)


# ---- 迭代2·质量收口：确定性质量门禁（quality_gate） ----
def _parse_humanize_score(text: str):
    """从改写后笔记/评分卡解析合计分（0-50）。返回 int / None。"""
    m = re.search(r"合计[*\s]*\*?\*?[\s*]*(\d{1,2})\s*/\s*50", text)
    if not m:
        m = re.search(r"(\d{1,2})\s*/\s*50", text)
    if m:
        v = int(m.group(1))
        return v if 0 <= v <= 50 else None
    return None


def _count_source_tags(report: str) -> dict:
    """统计质检报告里的来源标签 [原文]/[检索]/[推断]。"""
    return {
        "原文": len(re.findall(r"\[原文\]", report)),
        "检索": len(re.findall(r"\[检索\]", report)),
        "推断": len(re.findall(r"\[推断\]", report)),
    }


def _print_quality_gate(g: dict) -> None:
    score = g["humanize_score"]
    if score is None:
        sc_str, sc_pass = "未评估（未应用E6）", "—"
    else:
        sc_str, sc_pass = f"{score}/50", ("✅达标(≥40)" if g["score_pass"] else "❌未达标(<40)")
    print("📋 质量门禁（迭代2·质量收口）")
    print(f"  去AI味评分 : {sc_str}  -> {sc_pass}")
    print(f"  残留AI痕迹 : {g['residual_classes']} 类 / 共 {g['residual_total']} 处"
          f"  -> {'⚠️偏机器腔' if g['residual_flag'] else '✅可接受'}")
    if g["source_tags"]:
        s = g["source_tags"]
        print(f"  来源标签   : [原文]{s.get('原文', 0)} [检索]{s.get('检索', 0)} "
              f"[推断]{s.get('推断', 0)}"
              f"  -> {'🔴CHECKPOINT6触发' if g['checkpoint6'] else '✅'}")
    print(f"  → 综合判定 : {'✅ 通过，可交付' if g['passed'] else '❌ 不交付（先修复）'}")


def quality_gate(note_path: str, humanized_path: str = "", review_path: str = "",
                 as_json: bool = False) -> dict:
    """迭代2·质量收口：确定性门禁判定。

    - 去AI味评分：解析 humanized 笔记里的评分卡合计（≥40 达标）；未应用E6时为 None（不拦截）。
    - 残留痕迹：重新扫描 note，统计 AI 痕迹命中类别数（正则确定性）。
    - 来源可信度（可选，给 review 报告时）：统计 [原文]/[检索]/[推断]。

    返回结构化判定 dict；非 json 模式打印可读结论。"""
    note_md = Path(note_path).read_text(encoding="utf-8")
    # 1. 评分卡（优先解析 humanized 笔记，其次回退到 note 自身）
    score = None
    cand = humanized_path if humanized_path else note_path
    if Path(cand).exists():
        score = _parse_humanize_score(Path(cand).read_text(encoding="utf-8"))
    if score is None:
        score_pass = None
    else:
        score_pass = bool(score >= 40)
    # 2. 残留痕迹（确定性，复用 E6 扫描）
    findings = scan_ai_patterns(note_md)
    residual_classes = len(findings)
    residual_total = sum(c for _, c, _ in findings)
    # 3. 来源标签
    src = {}
    if review_path and Path(review_path).exists():
        src = _count_source_tags(Path(review_path).read_text(encoding="utf-8"))
    gate = {
        "humanize_score": score,
        "score_pass": score_pass,
        "residual_classes": residual_classes,
        "residual_total": residual_total,
        "residual_flag": bool(residual_classes >= 6 or residual_total >= 12),
        "source_tags": src,
        "checkpoint6": bool(src.get("推断", 0) >= 3),
    }
    # 综合 PASS：评分未未达标 且 残留未超标（score_pass=None 视为未评，不拦）
    gate["passed"] = bool((score_pass in (True, None)) and not gate["residual_flag"])
    if as_json:
        print(json.dumps(gate, ensure_ascii=False, indent=2))
    else:
        _print_quality_gate(gate)
    return gate


# ---- 通用：笔记结构抽取（供 E2/E3/E4/E5 复用） ----
def _note_title(note_md: str, fallback: str = "笔记") -> str:
    m = re.search(r"^#\s+(.+)$", note_md, re.M)
    return m.group(1).strip() if m else fallback


def _note_headings(note_md: str) -> list:
    """抽取 ## / ### 小节标题（去掉编号）。"""
    hs = re.findall(r"^#{2,3}\s+(.+)$", note_md, re.M)
    out = []
    for h in hs:
        h = re.sub(r"^\s*[\d.、]+\s*", "", h.strip())
        if h and h not in out:
            out.append(h)
    return out


def _note_bullets(note_md: str) -> list:
    """抽取要点行（-/*/数字列表），去掉标记。"""
    out = []
    for m in re.finditer(r"^\s*(?:[-*]|\d+[.、])\s+(.+)$", note_md, re.M):
        s = m.group(1).strip()
        if s:
            out.append(s)
    return out


# ---- E3/E4 手绘配图 / 信息图：抽取锚点 -> ImageGen 生图简报 ----
def extract_illustration_points(note_md: str, max_points: int = 4) -> list:
    """E3/E4：抽取适合配图的「认知锚点」——优先小节标题，其次要点。"""
    points = _note_headings(note_md)
    if len(points) < max_points:
        for b in _note_bullets(note_md):
            if b not in points:
                points.append(b)
            if len(points) >= max_points:
                break
    return points[:max_points]


def build_illustration_brief(note_md: str, points: list, mode: str,
                             out_path: Path) -> str:
    """E3(手绘)/E4(信息图) native 简报：列出每张图的 ImageGen 提示词与约束。"""
    title = _note_title(note_md)
    if mode == "infographic":
        head = (f"# BiliNote 信息图简报（E4）\n\n"
                f"用 **ImageGen 工具**为《{title}》生成一张一页图解信息图，"
                f"尺寸 1536x1024、高质量、手绘卡通风格、16:9 横向。\n\n"
                "## 关键点（3–7 个，每点配独立视觉隐喻，勿纯文字罗列）\n")
        body = "\n".join(
            f"{i+1}. {p} — 视觉隐喻：<为该点想一个具体隐喻，如缓存→快递柜>"
            for i, p in enumerate(points))
        rules = ("\n\n## 提示词模板\n"
                 f"主题：{title}\n标题：{title}\n分点：见上（附隐喻）\n"
                 "风格约束：手绘卡通、16:9 横向、克制配色、留白充足\n\n"
                 "## 规则\n- 关键点 3–7 个优先奇数；语言匹配笔记（中文→中文图）\n"
                 "- 迭代最多 2 次；生成后存为 note-infographic.png（与 note 同目录）\n")
    else:  # illustration = 小黑 IP 手绘
        head = (f"# BiliNote 手绘配图简报（E3）\n\n"
                f"用 **ImageGen 工具**为《{title}》逐张生成小黑 IP 手绘配图，"
                f"共 {len(points)} 张（短文 1–2，长文 3–4，最多 8）。\n\n"
                "## 配图锚点与逐张提示词\n")
        body = "\n".join(
            f"### {i+1}. {p}\n"
            "- 提示词：16:9、纯白背景、黑色线稿、少量红/橙/蓝批注、大量留白；"
            f"小黑（简笔小人）参与「{p}」的核心动作\n"
            "- 禁止：PPT风格 / 商业插画 / 可爱风 / 复杂架构图 / 左上角标题 / 水印\n"
            f"- 保存：assets/<note-slug>-illustrations/{i+1:02d}-topic.png，"
            f"并在笔记对应位置插入 ![{p}](assets/.../{i+1:02d}-topic.png)"
            for i, p in enumerate(points))
        rules = ("\n\n## QA 检查\n- 小黑是否装饰化？画面是否太满？是否像 PPT/流程图？"
                 "中文是否过多？背景是否干净白底？\n")
    return head + body + rules + "\n---\n\n## 源笔记\n---\n\n" + note_md + "\n"


def illustrate_note(note_path: str, mode: str = "illustration",
                    points: int = 4, out: str = "") -> tuple:
    """E3/E4：抽取配图锚点，产出 ImageGen 生图简报（native，交 agent 生图）。"""
    src = Path(note_path)
    note = src.read_text(encoding="utf-8")
    pts = extract_illustration_points(note, max_points=points)
    suffix = ".infographic-brief.md" if mode == "infographic" \
        else ".illustration-brief.md"
    out_path = Path(out) if out else src.with_name(src.stem + suffix)
    out_path.write_text(build_illustration_brief(note, pts, mode, out_path),
                        encoding="utf-8")
    return ("brief", str(out_path))


# ---- 迭代3·E3 配图科学化：构思卡 + 八项测试（illustration_plan） ----
# 6 图型关键词信号：用于脚本化识别候选图型
_FIGURE_SIGNALS = {
    "流程": [r"步骤", r"流程", r"阶段", r"首先[，、].{0,30}其次", r"第一[，、].{0,30}第二"],
    "架构": [r"架构", r"模块", r"系统", r"组成", r"分层", r"组件"],
    "对比": [r"对比", r"\bvs\.?\b", r"区别", r"差异", r"优于", r"不如", r"相比"],
    "关系": [r"因果", r"导致", r"影响", r"关联", r"推动", r"取决于", r"带来"],
    "结构": [r"结构", r"框架", r"层级", r"包含", r"分类", r"维度"],
    "概念": [r"概念", r"定义", r"本质", r"核心", r"是什么", r"意味着"],
}
# 泛化词（换文测试：含过多则太泛，放进别篇也成立）
_VAGUE_WORDS = ["众多", "大量", "一系列", "多种", "各类", "方方面面",
                "一些", "某些", "通常", "一般", "许多"]


def _detect_figure_type(text: str) -> str:
    """按关键词信号给文本猜图型（6 选 1，默认概念）。"""
    best, best_n = "概念", 0
    for fig, pats in _FIGURE_SIGNALS.items():
        n = sum(1 for p in pats if re.search(p, text))
        if n > best_n:
            best, best_n = fig, n
    return best


def _extract_entities(text: str, max_n: int = 5) -> list:
    """抽中文实体（2-4 字名词短语），作为信息单位 / 原文指纹。"""
    cands = re.findall(
        r"[\u4e00-\u9fa5]{2,4}(?:系统|框架|模式|机制|方法|理论|模型|"
        r"策略|结构|流程|数据|风险|收益|成本|效果|能力|需求|问题|"
        r"方案|标准|原则|陷阱|信号|周期|率|度|法)", text)
    seen, out = set(), []
    for c in cands:
        if c not in seen:
            seen.add(c)
            out.append(c)
        if len(out) >= max_n:
            break
    return out


def _paragraph_for_heading(note_md: str, heading: str) -> str:
    """取某标题下的正文段（到下一个同级/上级标题前）。"""
    lines = note_md.splitlines()
    buf, capturing = [], False
    for ln in lines:
        if re.match(r"^#{2,3}\s+", ln):
            h = re.sub(r"^\s*[\d.、]+\s*", "", ln).strip("# ").strip()
            if h == heading:
                capturing = True
                continue
            elif capturing:
                break
        if capturing:
            buf.append(ln)
    return "\n".join(buf).strip()


def _illustration_tests(card: dict) -> dict:
    """八项测试（脚本可判部分）：返回 {test: (status, reason)}。
    status: pass / warn / fail。主观项（视觉重量/阅读顺序）留由 LLM 在简报中补。"""
    res = {}
    prop = card["single_prop"]
    fig = card["figure"]
    units = card["info_units"]
    vague = card["vague_hits"]
    res["一秒"] = (("pass" if len(prop) >= 6 else "fail"),
                   "单一命题可一句话说清" if len(prop) >= 6 else "单一命题过短/空")
    res["三秒"] = (("pass" if re.search(r"[是成为导致影响决定属于]", prop) else "warn"),
                   "命题含判断动词更佳")
    res["换文"] = (("pass" if vague == 0 else "fail"),
                   (f"含 {vague} 个泛化词，放进别篇也成立→太泛" if vague
                    else "具体，不易泛化"))
    res["删减"] = (("pass" if len(units) <= 5 else "warn"),
                   (f"信息单位 {len(units)} 个，建议删减至 ≤5" if len(units) > 5
                    else "信息密度合适"))
    res["图型"] = (("pass" if fig != "概念" or card["signal_hits"] > 0 else "warn"),
                   f"识别为「{fig}」图型")
    res["无文"] = (("pass" if len(units) <= 5 else "warn"),
                   ("结构可脱离长标签理解" if len(units) <= 5 else "依赖长标签"))
    res["重复审美"] = (("pass" if (len(set(units)) == len(units) or len(units) <= 4)
                        else "warn"), "重复元素未超预算")
    res["眯眼平衡"] = (("pass" if 1 <= len(units) <= 6 else "warn"),
                       ("视觉重量较均衡" if 1 <= len(units) <= 6 else "元素过少/过满"))
    return res


def _write_illustration_plan(note: str, cards: list) -> str:
    """把构思卡 + 八项测试渲染为 agent 可读简报。"""
    title = _note_title(note)
    lines = [f"# BiliNote 配图构思卡（迭代3·E3 科学化）\n",
             f"> 为《{title}》生成的 {len(cards)} 张配图构思卡 + 八项测试。"
             "agent 据此调 ImageGen；**全过才进生图**，未过项先回退。\n"]
    for c in cards:
        lines.append(f"\n## 构思卡 {c['idx']} — {c['anchor']}\n")
        lines.append("| 字段 | 值 |")
        lines.append("|---|---|")
        lines.append(f"| 插入位置 | 标题「{c['insert_after']}」末尾 |")
        lines.append(f"| 图型 | {c['figure']} |")
        lines.append(f"| 配图任务 | {c['task']} |")
        lines.append(f"| 单一命题 | {c['single_prop']} |")
        lines.append(f"| 原文指纹 | {c['fingerprint']} |")
        lines.append(f"| 信息单位 | {'、'.join(c['info_units']) or '—'} |")
        lines.append(f"| 重复元素预算 | {c['budget']} |")
        lines.append(f"| 图型信号命中 | {c['signal_hits']} |")
        lines.append("\n**八项测试**：\n")
        for t, (st, why) in c["tests"].items():
            icon = "✅" if st == "pass" else ("⚠️" if st == "warn" else "❌")
            lines.append(f"- {icon} {t}：{why}")
        if c["fallbacks"]:
            lines.append(f"\n**回退建议**：未过 {', '.join(c['fallbacks'])}"
                         " → 重选锚点 / 换图型 / 精简信息单位后重测。")
        else:
            lines.append("\n**✅ 八项全过，可进 ImageGen**（路由：小黑手绘 / 极简强线条）。")
    all_pass = all(c["passed"] for c in cards)
    advice = ("全部构思卡通过八项测试，可直接生图。"
              if all_pass else
              "存在未过构思卡，建议先回退修正后再生图（避免凭运气出图）。")
    lines.append(f"\n---\n\n## 进 ImageGen 建议\n{advice}")
    return "\n".join(lines) + "\n"


def illustration_plan(note_path: str, max_cards: int = 4, out: str = "") -> tuple:
    """迭代3·E3 配图科学化：脚本化生成构思卡 + 八项测试。

    流程：扫描 note.md 建段落地图 → 识别图型 → 选 2–4 候选配图位 →
    每张出构思卡关键字段（插入位置/图型/单一命题/原文指纹/信息单位/重复预算）
    → 跑八项测试（脚本可判部分）→ 写 <note>.illustration-plan.md 简报交 agent。
    全过才建议进 ImageGen；否则标注未过项与回退建议。"""
    src = Path(note_path)
    note = src.read_text(encoding="utf-8")
    headings = _note_headings(note)
    skip = ("前言", "引言", "结语", "总结", "目录", "参考", "附录", "后记")
    candidates = [h for h in headings
                  if not any(s in h for s in skip) and len(h) >= 4][:max_cards] \
        or headings[:max_cards]
    cards = []
    for i, h in enumerate(candidates):
        para = _paragraph_for_heading(note, h)
        fig = _detect_figure_type(h + " " + para)
        entities = _extract_entities(para or h)
        vague = sum(1 for w in _VAGUE_WORDS if w in (para or h))
        card = {
            "idx": i + 1,
            "anchor": h,
            "insert_after": h,
            "figure": fig,
            "task": "解释",
            "single_prop": h,
            "fingerprint": entities[0] if entities else h[:6],
            "info_units": entities,
            "vague_hits": vague,
            "signal_hits": sum(1 for p in _FIGURE_SIGNALS[fig]
                               if re.search(p, h + " " + para)),
            "budget": f"同类元素≤{min(4, len(entities) or 4)}",
        }
        card["tests"] = _illustration_tests(card)
        fails = [t for t, (st, _) in card["tests"].items() if st == "fail"]
        card["passed"] = not fails
        card["fallbacks"] = fails
        cards.append(card)
    out_path = Path(out) if out else src.with_name(src.stem + ".illustration-plan.md")
    out_path.write_text(_write_illustration_plan(note, cards), encoding="utf-8")
    return ("brief", str(out_path), cards)


# ---- E2 研究写作辅助：抽取待证据主张 -> WebSearch 研究简报 ----
_CLAIM_SIGNALS = {
    "数据": r"\d+(?:\.\d+)?\s*(?:%|％|倍|万|亿|千|百|元|美元|年|月|天|个|次|条|人|款)",
    "比较": r"(?:最[高低多少快慢好差大小]|更[高低多少快慢好差]|"
          r"比[^，。]{1,10}(?:高|低|多|少|快|慢|好|差)|领先|超过|不如)",
    "因果": r"(?:因为|由于|导致|因此|所以|使得|带来|意味着)",
    "绝对": r"(?:一定|必然|所有|从不|绝不|唯一|完全|彻底|必须)",
}


def extract_claims(note_md: str, max_claims: int = 8) -> list:
    """E2：抽取最需要外部佐证的主张句（含数据/比较/因果/绝对断言）。

    返回 [(句子, [命中信号类别…]), ...]，按命中类别数降序。"""
    body = re.sub(r"^#.*$", "", note_md, flags=re.M)  # 去标题行
    body = re.sub(r"^\s*(?:[-*]|\d+[.、])\s*", "", body, flags=re.M)
    sentences = re.split(r"[。！？\n]", body)
    scored = []
    for s in sentences:
        s = s.strip()
        if len(s) < 8:
            continue
        cats = [k for k, pat in _CLAIM_SIGNALS.items() if re.search(pat, s)]
        if cats:
            scored.append((s, cats))
    scored.sort(key=lambda x: len(x[1]), reverse=True)
    return scored[:max_claims]


def build_research_brief(note_md: str, claims: list, out_path: Path) -> str:
    """E2 native 简报：把待证据主张转成 WebSearch 查询清单 + 引用规范 + 置信度表（Plan J）。"""
    title = _note_title(note_md)
    lines = [
        "# BiliNote 研究补充简报（E2 + Plan J 声明级事实核查）\n",
        f"为《{title}》补充权威来源。请用 **WebSearch 工具**逐条检索下列主张，"
        "找到可信来源后按引用格式回填到 note.md 末尾「延伸阅读」章节，"
        "并在「事实置信度表」中标注每条主张的核验结果。\n",
        "## 待证据主张与建议检索词\n",
    ]
    if not claims:
        lines.append("（未自动识别到含数据/断言的主张，可由你判断需补证之处。）\n")
    for i, (s, cats) in enumerate(claims):
        kw = re.sub(r"\s+", " ", s)[:40]
        lines.append(f"{i+1}. 主张（{'/'.join(cats)}）：{s}\n"
                     f"   - 建议检索：`{kw}` 出处 / 数据来源\n")
    lines += [
        "\n## 引用格式\n",
        "> 来源标题 (作者, 年份)。URL\n",
        "\n## 事实置信度表（Plan J · 必填）\n",
        "对每条主张，检索后在下表标注核验结果。标注规则：\n"
        "- ✅ 已证：找到权威来源支持该主张\n"
        "- ⚠️ 存疑：来源不一致 / 仅找到间接证据 / 数据口径不同\n"
        "- ❌ 冲突：权威来源明确反驳该主张（须在 note.md 标注更正）\n\n"
        "| # | 主张(摘要) | 置信度 | 来源/说明 |\n"
        "|---|----------|--------|----------|\n",
    ]
    for i, (s, _) in enumerate(claims):
        lines.append(f"| {i+1} | {re.sub(r'[|\\n]', ' ', s)[:30]} | ⚠️存疑 | （待检索） |\n")
    lines += [
        "\n## 规则\n",
        "- 只用权威来源（学术期刊 / 官方报告 / 知名媒体），每条须含可访问 URL。\n"
        "- 检索不到就如实说明，**不得编造来源**。\n"
        "- 标 ❌冲突 的主张，须在 note.md 对应处加 `> ⚠️ 事实核查：此说法与权威来源冲突（见延伸阅读）`。\n"
        "- 补充内容与正文用 `---` 分隔，标注「延伸阅读」。\n",
        "\n---\n\n## 源笔记\n---\n\n", note_md, "\n",
    ]
    return "".join(lines)


def research_note(note_path: str, max_claims: int = 8, out: str = "") -> tuple:
    """E2：抽取待证据主张，产出研究简报（native，交 agent 跑 WebSearch）。"""
    src = Path(note_path)
    note = src.read_text(encoding="utf-8")
    claims = extract_claims(note, max_claims=max_claims)
    out_path = Path(out) if out else src.with_name(src.stem + ".research-brief.md")
    out_path.write_text(build_research_brief(note, claims, out_path),
                        encoding="utf-8")
    return ("brief", str(out_path))


# ---- E5 网页PPT：抽取叙事弧大纲 -> PPT 简报 ----
def extract_slide_outline(note_md: str, minutes: int = 15) -> list:
    """E5：按「钩子→定调→主体→转折→收束」搭建页骨架。

    页数约束：15min≈10 页、30min≈20 页、45min≈25–30 页。"""
    title = _note_title(note_md)
    headings = _note_headings(note_md)
    target_pages = max(6, min(30, round(minutes / 1.5)))
    body_budget = max(1, target_pages - 4)  # 去掉封面/钩子/收束/致谢
    slides = [
        ("封面", f"《{title}》"),
        ("钩子", "用一个问题 / 反差 / 数据抓注意力"),
        ("定调", "本次要解决什么、为什么值得听"),
    ]
    for h in headings[:body_budget]:
        slides.append(("主体", h))
    slides.append(("收束", "一句话总结 + 行动建议"))
    return slides[:target_pages]


def build_slides_brief(note_md: str, outline: list, minutes: int,
                       out_path: Path) -> str:
    """E5 native 简报：把大纲交 agent 生成单文件 HTML 横向翻页 PPT。"""
    title = _note_title(note_md)
    pages = "\n".join(f"{i+1}. [{tag}] {txt}" for i, (tag, txt) in enumerate(outline))
    return (
        "# BiliNote 网页 PPT 简报（E5）\n\n"
        f"为《{title}》生成**单文件 HTML 横向翻页 PPT**（约 {minutes} 分钟 / "
        f"{len(outline)} 页），CSS/JS/图片全部内联。\n\n"
        "## 叙事弧大纲（钩子→定调→主体→转折→收束）\n"
        f"{pages}\n\n"
        "## 风格二选一\n"
        "- A. 电子杂志 × 电子墨水（衬线 + 流体背景 + 暖色）\n"
        "- B. 瑞士国际主义（无衬线 + 网格点阵 + 高亮色）\n\n"
        "## 规则\n"
        "- 单文件、无外部依赖；左右箭头 / 触摸滑动翻页。\n"
        "- 无图片素材时用纯文字 + 排版 + 色彩传达；生成前先对齐叙事弧与页数节奏。\n"
        f"- 保存为 note-presentation.html（与 note 同目录）。\n\n"
        "---\n\n## 源笔记（内容来源）\n---\n\n"
        f"{note_md}\n")


def slides_note(note_path: str, minutes: int = 15, out: str = "") -> tuple:
    """E5：抽取叙事弧大纲，产出 PPT 简报（native，交 agent 生成 HTML）。"""
    src = Path(note_path)
    note = src.read_text(encoding="utf-8")
    outline = extract_slide_outline(note, minutes=minutes)
    out_path = Path(out) if out else src.with_name(src.stem + ".slides-brief.md")
    out_path.write_text(build_slides_brief(note, outline, minutes, out_path),
                        encoding="utf-8")
    return ("brief", str(out_path))


# --------------------------------------------------------------------------
# Plan E: 批处理与系列化（batch）
#   对「一批源」（清单文件 / 目录扫描 / 多个位置参数）批量执行既有能力
#   （run / produce / humanize / illustrate / research / slides），
#   支持串行或线程并发、失败隔离（单项报错不影响其余）、统一输出目录、
#   run 模式下的链式后处理（--then produce,humanize），并产出汇总报告。
#   复用全部已有函数，不重复实现流水线；确定性部分（清单解析 / slug /
#   报告）为纯函数，可离线测试。
# --------------------------------------------------------------------------
BATCH_OPS = ("run", "produce", "humanize", "illustrate", "research", "slides")


def parse_manifest_text(text: str) -> list:
    """解析文本清单：每行一个源；空行与 # 注释忽略；支持 '源 | 标题'。"""
    items = []
    for line in (text or "").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if "|" in line:
            src, _, title = line.partition("|")
            items.append({"source": src.strip(), "title": title.strip()})
        else:
            items.append({"source": line})
    return items


def parse_manifest_json(obj) -> list:
    """解析 JSON 清单：数组（字符串或对象）或 {items/sources, defaults}。

    对象项支持 source/url、title、out、style、lang 等字段；defaults 会被
    合并进每一项（项内同名字段优先）。"""
    if isinstance(obj, dict):
        defaults = obj.get("defaults", {}) or {}
        raw = obj.get("items") or obj.get("sources") or []
    else:
        defaults, raw = {}, obj
    items = []
    for it in raw or []:
        if isinstance(it, str):
            d = {"source": it}
        elif isinstance(it, dict):
            d = dict(it)
            if "source" not in d and "url" in d:
                d["source"] = d["url"]
        else:
            continue
        merged = dict(defaults)
        merged.update(d)
        items.append(merged)
    return items


def load_manifest(manifest: str = "", dir_: str = "", glob_: str = "*.md",
                  sources: list = None) -> list:
    """汇聚三种来源为统一的 item 列表：位置参数 + 清单文件 + 目录扫描。

    清单文件按后缀/内容自动识别 JSON 或纯文本。返回去掉空源后的 item 列表。"""
    items = []
    if sources:
        items += [{"source": s} for s in sources if str(s).strip()]
    if manifest:
        p = Path(manifest)
        txt = p.read_text(encoding="utf-8")
        if p.suffix.lower() == ".json" or txt.lstrip().startswith(("[", "{")):
            items += parse_manifest_json(json.loads(txt))
        else:
            items += parse_manifest_text(txt)
    if dir_:
        for f in sorted(Path(dir_).glob(glob_)):
            if f.is_file():
                items.append({"source": str(f)})
    return [it for it in items if str(it.get("source", "")).strip()]


def slugify(text: str, maxlen: int = 40) -> str:
    """把源/标题压成安全的文件名词干：URL 取 BV/av 号或末段，路径取 stem。"""
    s = str(text or "").strip()
    m = re.search(r"(BV[0-9A-Za-z]{10}|av\d+)", s)
    if m:
        s = m.group(1)
    elif re.match(r"^https?://", s, re.I):
        s = re.split(r"[?#]", s)[0]          # 先去掉查询串/锚点
        seg = [x for x in s.split("/") if x]
        s = seg[-1] if seg else s
    elif re.search(r"[\\/]", s) or s.lower().endswith(
            (".md", ".txt", ".vtt", ".srt", ".ass", ".mp4", ".mkv", ".webm")):
        s = Path(s).stem
    s = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", s, flags=re.U)
    s = re.sub(r"-{2,}", "-", s).strip("-")
    return (s[:maxlen].strip("-") or "item")


def dedup_slug(slug: str, used: set) -> str:
    """保证 slug 在一批内唯一：冲突则追加 -2 / -3 …"""
    base, n, out = slug, 2, slug
    while out in used:
        out = f"{base}-{n}"
        n += 1
    used.add(out)
    return out


def _batch_item_slug(item: dict, index: int) -> str:
    return slugify(item.get("title") or item.get("source") or f"item{index}")


def _first_real_note(paths: list) -> str:
    """从产物路径里找出「真实笔记」（.md 且非 brief 且存在），用于链式后处理。"""
    for p in paths:
        s = str(p)
        name = Path(s).name.lower()
        if s.endswith(".md") and "brief" not in name and Path(s).exists():
            return s
    return ""


def _batch_dispatch(op: str, item: dict, out_base: Path, slug: str,
                    opts: dict) -> list:
    """按 op 分派到既有能力，返回 [(kind, payload), ...]。"""
    if op == "run":
        out_md = out_base / f"{slug}.md"
        path = run(item["source"], out_md,
                   lang=item.get("lang") or opts.get("lang"),
                   style=item.get("style") or opts.get("style") or DEFAULT_STYLE,
                   export_formats=opts.get("export_formats") or [],
                   review=opts.get("review", False),
                   clean=opts.get("clean", True))
        return [("note", str(path) if path else str(out_md))]
    note = item["source"]
    if op == "produce":
        return produce_from_note(
            note, targets=opts.get("targets", ""),
            out_dir=str(out_base),
            preset=opts.get("preset", ""))
    if op == "humanize":
        kind, payload, _ = humanize_note(note, out="")
        return [(kind, payload)]
    if op == "illustrate":
        return [illustrate_note(note, mode=opts.get("mode", "illustration"),
                                points=opts.get("points", 4))]
    if op == "research":
        return [research_note(note, max_claims=opts.get("max_claims", 8))]
    if op == "slides":
        return [slides_note(note, minutes=opts.get("minutes", 15))]
    raise ValueError(f"未知 batch op: {op}")


def run_batch(items: list, op: str = "run", out_dir: str = "",
              jobs: int = 1, then: list = None,
              stop_on_error: bool = False, opts: dict = None) -> list:
    """批处理编排器：对 items 逐项执行 op，失败隔离，返回结果记录列表。

    - jobs>1 时用线程池并发（适合 run 的下载/LLM I/O）；否则串行。
    - op=='run' 且 then 非空时，对每篇产出的「真实笔记」链式跑后处理。
    - 单项异常被捕获记入 record.error，不影响其余项；stop_on_error 时提前止损。
    """
    opts = opts or {}
    then = then or []
    out_base = Path(out_dir) if out_dir else Path(".")
    out_base.mkdir(parents=True, exist_ok=True)
    used, plan = set(), []
    for i, item in enumerate(items, 1):
        plan.append((i, dedup_slug(_batch_item_slug(item, i), used), item))
    records = [None] * len(plan)

    def work(idx0, i, slug, item):
        rec = {"index": i, "source": item.get("source"), "slug": slug,
               "op": op, "status": "ok", "outputs": [], "then": [], "error": ""}
        try:
            outs = _batch_dispatch(op, item, out_base, slug, opts)
            rec["outputs"] = [str(x) for _, x in outs]
            if op == "run" and then:
                note_path = _first_real_note(rec["outputs"])
                if note_path:
                    for f in then:
                        fo = _batch_dispatch(f, {"source": note_path},
                                             out_base, slug, opts)
                        rec["then"] += [f"{f}:{x}" for _, x in fo]
                else:
                    rec["then"] = ["<deferred: native brief，需先据简报补全笔记>"]
        except Exception as e:  # 失败隔离
            rec["status"] = "failed"
            rec["error"] = str(e)
        return idx0, rec

    if jobs and jobs > 1:
        from concurrent.futures import ThreadPoolExecutor, as_completed
        with ThreadPoolExecutor(max_workers=jobs) as ex:
            futs = [ex.submit(work, k, i, slug, item)
                    for k, (i, slug, item) in enumerate(plan)]
            for fut in as_completed(futs):
                k, rec = fut.result()
                records[k] = rec
                log(f"[batch {rec['index']}/{len(plan)}] "
                    f"{rec['status']} {rec['source']}")
                if rec["status"] == "failed" and stop_on_error:
                    break
    else:
        for k, (i, slug, item) in enumerate(plan):
            _, rec = work(k, i, slug, item)
            records[k] = rec
            log(f"[batch {i}/{len(plan)}] {rec['status']} {rec['source']}")
            if rec["status"] == "failed" and stop_on_error:
                break
    return [r for r in records if r]


def format_batch_report(records: list, op: str) -> str:
    """把批处理结果渲染成 Markdown 汇总报告。"""
    total = len(records)
    ok = sum(1 for r in records if r["status"] == "ok")
    failed = total - ok
    lines = [f"# 批处理报告（op={op}）", "",
             f"- 共 **{total}** 项：✅ {ok} 成功 / ❌ {failed} 失败", "",
             "| # | 源 | 状态 | 产物 | 备注 |",
             "|---|---|---|---|---|"]
    for r in records:
        outs = "<br>".join(r["outputs"]) or "-"
        if r["status"] == "failed":
            note = r["error"]
        else:
            note = ", ".join(r.get("then") or [])
        st = "✅" if r["status"] == "ok" else "❌"
        lines.append(f"| {r['index']} | {r['source']} | {st} | {outs} | {note} |")
    return "\n".join(lines)


# --------------------------------------------------------------------------
# Plan H — 跨笔记聚合 / 系列专题资产
#   读一批 note.md → 合并系列总思维导图 + 专题索引页 + 跨集知识图谱。
# --------------------------------------------------------------------------
def series_aggregate(note_paths: list, out_dir: str = "",
                      title: str = "系列专题") -> dict:
    """聚合多份 note.md 产出系列资产。返回 {mindmap, index, graph} 路径。

    - <title>.series.mindmap.md：系列总 Mermaid 思维导图（每篇一个分支）
    - <title>.series.index.md：专题索引页（每篇标题+摘要+跳转）
    - <title>.series.graph.md：跨集知识图谱（共享概念共现）
    """
    out_base = Path(out_dir) if out_dir else Path(note_paths[0]).parent
    out_base.mkdir(parents=True, exist_ok=True)
    stem = slugify(title) or "series"
    entries = []
    all_headings = []
    for p in note_paths:
        try:
            md = Path(p).read_text(encoding="utf-8")
        except Exception:  # noqa: BLE001
            continue
        t = _note_title(md, Path(p).stem)
        hs = _note_headings(md)
        # 摘要：首段非标题文本前 80 字
        body = re.sub(r"^#.*$", "", md, flags=re.M).strip()
        summary = re.sub(r"\s+", " ", body)[:80]
        entries.append({"path": str(p), "title": t, "headings": hs,
                        "summary": summary})
        all_headings.extend(hs)

    # 系列总思维导图
    mm = ["```mermaid", f"mindmap", f"  ({title})"]
    for e in entries:
        mm.append(f"  :::{e['title']}")
        for h in e["headings"][:6]:
            mm.append(f"    {h}")
    mm.append("```")
    mm_path = out_base / f"{stem}.series.mindmap.md"
    mm_path.write_text(f"# {title} · 系列总思维导图\n\n" + "\n".join(mm),
                       encoding="utf-8")

    # 专题索引页
    idx = [f"# {title} · 专题索引\n", f"共 {len(entries)} 篇。\n"]
    for i, e in enumerate(entries):
        idx.append(f"## {i+1}. {e['title']}\n")
        idx.append(f"{e['summary']}...\n")
        idx.append(f"- 📄 笔记：`{e['path']}`\n")
        if e["headings"]:
            idx.append("- 小节：" + " / ".join(e["headings"][:5]) + "\n")
    idx_path = out_base / f"{stem}.series.index.md"
    idx_path.write_text("\n".join(idx), encoding="utf-8")

    # 跨集知识图谱（共享概念共现）
    from collections import Counter
    freq = Counter(all_headings)
    shared = [h for h, n in freq.items() if n >= 2]
    graph = [f"# {title} · 跨集知识图谱\n",
             f"共现概念 {len(shared)} 个（出现在 ≥2 篇中）。\n\n```mermaid\ngraph LR\n"]
    for h in shared[:20]:
        safe = re.sub(r"[^\w\u4e00-\u9fa5]", "", h)[:12]
        graph.append(f'  {safe}["{h} ({freq[h]})"]')
    graph.append("```")
    graph_path = out_base / f"{stem}.series.graph.md"
    graph_path.write_text("\n".join(graph), encoding="utf-8")

    log(f"系列聚合：{len(entries)} 篇 → 总导图/索引/图谱 已写入 {out_base}")
    return {"mindmap": str(mm_path), "index": str(idx_path),
            "graph": str(graph_path), "count": len(entries)}


def _gzh_guess_title(md: str, fallback: str) -> str:
    for line in md.splitlines():
        s = line.strip()
        if s.startswith("# "):
            return s[2:].strip()
    return fallback


if __name__ == "__main__":
    main()
